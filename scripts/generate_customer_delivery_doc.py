from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = PROJECT_ROOT / "dist" / "ESG_Agentic_RAG_Copilot_客户交付说明_2026-04-07.docx"


def _set_default_font(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name in ("Title", "Heading 1", "Heading 2"):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def build_document() -> Document:
    document = Document()
    _set_default_font(document)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("ESG Agentic RAG Copilot 客户交付说明")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run(f"交付日期：{date(2026, 4, 7).isoformat()}")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10.5)

    summary = document.add_paragraph()
    summary.add_run("交付定位：").bold = True
    summary.add_run(
        "本版本用于客户交付、验收演示和实施部署。当前默认运行链路为“本地模型优先，"
        "失败后回退 DeepSeek，最终由 OpenAI 兜底”；远端 GPU / 5090 不作为本次交付前置条件。"
    )

    document.add_heading("1. 本次交付范围", level=1)
    _add_bullet_list(
        document,
        [
            "后端 API 服务与智能分析链路",
            "前端静态页面与运营展示入口",
            "RAG 检索、报告生成、数据同步与运维检查能力",
            "本地优先部署方案，以及后续远端 GPU 扩展接口保留",
        ],
    )

    document.add_heading("2. 本次交付基线", level=1)
    _add_bullet_list(
        document,
        [
            "应用模式：APP_MODE=local",
            "LLM 策略：LLM_BACKEND_MODE=auto",
            "默认回答链路：Local -> DeepSeek -> OpenAI",
            "当前交付不强制依赖 5090 或远端 LoRA 服务",
        ],
    )

    document.add_heading("3. 最终交付物清单", level=1)
    _add_bullet_list(
        document,
        [
            "项目源码与运行目录结构",
            "环境变量模板文件：.env.example",
            "后端启动入口：scripts/run_local_first_windows.bat",
            "环境初始化脚本：scripts/bootstrap_local_windows.bat",
            "运行时诊断脚本：scripts/runtime_doctor.py",
            "部署预检脚本：scripts/staging_check.py",
            "交付部署与验收清单：docs/PRODUCT_DELIVERY_CHECKLIST_ZH.md",
            "前端静态资源目录：frontend/",
            "数据库迁移脚本目录：gateway/db/migrations/",
        ],
    )

    document.add_heading("4. 部署与启动入口", level=1)
    _add_bullet_list(
        document,
        [
            "步骤 1：执行 scripts/bootstrap_local_windows.bat 初始化环境",
            "步骤 2：确认 .env 中 Supabase 与云端兜底密钥已配置",
            "步骤 3：启动 Qdrant 服务",
            "步骤 4：执行 scripts/run_local_first_windows.bat 启动 API",
            "步骤 5：执行 scripts/runtime_doctor.py 与 scripts/staging_check.py preflight 进行体检",
        ],
    )

    document.add_heading("5. 客户侧环境要求", level=1)
    _add_bullet_list(
        document,
        [
            "Windows 环境下建议使用 Python 3.11 及以上版本",
            "Docker Desktop 或 Docker Engine 已安装并可正常启动",
            "Qdrant 服务可访问",
            "Supabase 项目与访问密钥已准备完成",
            "至少配置一个云端兜底密钥：DEEPSEEK_API_KEY 或 OPENAI_API_KEY",
        ],
    )

    document.add_heading("6. 标准验收项", level=1)
    _add_bullet_list(
        document,
        [
            "/health 返回 200，且状态为 ok",
            "/dashboard/overview 返回 200，页面总览可访问",
            "/agent/analyze 可正常返回 answer 与 confidence",
            "/admin/reports/generate 可生成报告并返回 report_id",
            "/admin/data-sources/sync 可正常创建同步任务",
        ],
    )

    document.add_heading("7. 上线前确认项", level=1)
    _add_bullet_list(
        document,
        [
            "Docker Desktop / Docker Engine 已启动",
            "Qdrant 已可连接，避免首次分析时触发长时间内存重建",
            "本地 API 启动后已通过一次真实业务烟测",
            "Supabase 连接与客户密钥配置已完成",
            "正式环境日志、联系人与运维路径已确认",
        ],
    )

    document.add_heading("8. 交付结论", level=1)
    conclusion = document.add_paragraph()
    conclusion.add_run("当前版本已完成代码交付、测试通过与部署预检，可作为客户交付与验收版本。").bold = True
    conclusion.add_run(
        "正式生产上线前，请先确认 Docker 与 Qdrant 运行条件已就绪，并完成一次业务烟测通过。"
    )

    support = document.add_paragraph()
    support.add_run("建议运维联系人在交接时同步说明：").bold = True
    support.add_run(
        "默认链路为本地优先；若未来恢复远端 GPU，再切换到远端增强方案，不影响本次交付基线。"
    )

    return document


def main() -> int:
    output_path = DEFAULT_OUTPUT
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(output_path)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
