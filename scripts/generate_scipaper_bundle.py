from __future__ import annotations

import argparse
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "scipaper"
HISTORICAL_REFERENCE_FILES = ["论文1初稿.docx", "实验操作手册.docx"]
COLPALI_SCAN_DIRS = ["gateway", "training", "docs", "analysis", "rag"]
TEXT_EXTENSIONS = {".py", ".md", ".txt", ".json", ".toml", ".yml", ".yaml"}


def _utc_timestamp() -> str:
    return datetime.now(timezone.utc).isoformat()


def _relative_path(path: Path) -> str:
    return path.relative_to(PROJECT_ROOT).as_posix()


def _load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _load_historical_paper_title(path: Path) -> str:
    if not path.exists():
        return "Historical paper not found"
    document = Document(path)
    candidates = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    if not candidates:
        return "Historical paper title unavailable"

    generic_headers = {"SCI PAPER DRAFT", "EXPERIMENT GUIDE", "DRAFT"}
    for text in candidates:
        if text.upper() in generic_headers:
            continue
        if len(text.split()) >= 5 or "-" in text:
            return text
    return candidates[0]


def _resolve_historical_reference_path(output_dir: Path, filename: str) -> Path:
    direct = output_dir / filename
    if direct.exists():
        return direct
    folder_05 = next((path for path in output_dir.iterdir() if path.is_dir() and path.name.startswith("05_")), None)
    if folder_05 is not None:
        candidate = folder_05 / filename
        if candidate.exists():
            return candidate
    recursive = next((path for path in output_dir.rglob(filename) if path.is_file()), None)
    return recursive or direct


def _scan_term_hits(term: str, roots: list[str]) -> list[str]:
    matches: list[str] = []
    lowered_term = term.lower()
    for root_name in roots:
        root = PROJECT_ROOT / root_name
        if not root.exists():
            continue
        for path in root.rglob("*"):
            if not path.is_file() or path.suffix.lower() not in TEXT_EXTENSIONS:
                continue
            try:
                content = path.read_text(encoding="utf-8", errors="ignore")
            except Exception:
                continue
            if lowered_term in content.lower():
                matches.append(_relative_path(path))
    return sorted(set(matches))


def _load_lora_eval_summary() -> dict[str, Any]:
    eval_report_path = PROJECT_ROOT / "data" / "rag_eval" / "eval_report.json"
    if not eval_report_path.exists():
        return {"num_samples": None, "avg_rougeL": None, "path": _relative_path(eval_report_path)}
    payload = _load_json(eval_report_path)
    return {
        "num_samples": payload.get("num_samples"),
        "avg_rougeL": payload.get("avg_rougeL"),
        "path": _relative_path(eval_report_path),
    }


def _configure_document_fonts(document: Document, *, latin_font: str, east_asia_font: str, base_size: float = 11.0) -> None:
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "List Bullet"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        style.font.name = latin_font
        style.font.size = Pt(base_size)
        if style._element.rPr is None:
            style._element.get_or_add_rPr()
        style._element.rPr.rFonts.set(qn("w:ascii"), latin_font)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), latin_font)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def _add_title_block(document: Document, title: str, subtitle: str | None = None) -> None:
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)

    if subtitle:
        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(10.5)


def _add_paragraphs(document: Document, paragraphs: list[str]) -> None:
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)


def _add_bullets(document: Document, bullets: list[str]) -> None:
    for bullet in bullets:
        document.add_paragraph(bullet, style="List Bullet")


def _add_numbered_items(document: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        document.add_paragraph(f"[{index}] {item}")


def _add_tables(document: Document, tables: list[dict[str, Any]]) -> None:
    for table_spec in tables:
        caption = table_spec.get("caption")
        if caption:
            document.add_paragraph(caption)
        headers = [str(item) for item in table_spec.get("headers", [])]
        rows = [[str(cell) for cell in row] for row in table_spec.get("rows", [])]
        if not headers:
            continue

        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            table.cell(0, index).text = header

        for row in rows:
            row_cells = table.add_row().cells
            for index, cell in enumerate(row):
                if index < len(row_cells):
                    row_cells[index].text = cell


def _build_document(
    *,
    title: str,
    subtitle: str | None,
    sections: list[dict[str, Any]],
    latin_font: str,
    east_asia_font: str,
) -> Document:
    document = Document()
    _configure_document_fonts(document, latin_font=latin_font, east_asia_font=east_asia_font)
    _add_title_block(document, title, subtitle)

    for section in sections:
        document.add_heading(section["heading"], level=section.get("level", 1))
        if section.get("paragraphs"):
            _add_paragraphs(document, section["paragraphs"])
        if section.get("bullets"):
            _add_bullets(document, section["bullets"])
        if section.get("numbered_items"):
            _add_numbered_items(document, section["numbered_items"])
        if section.get("tables"):
            _add_tables(document, section["tables"])
    return document


def _write_docx(path: Path, document: Document) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def _write_json(path: Path, payload: dict[str, Any]) -> None:
    _write_text(path, json.dumps(payload, ensure_ascii=False, indent=2))


def _append_optional_suffix(base: str, suffix: str | None) -> str:
    if not suffix:
        return base
    if base.endswith("."):
        base = base[:-1]
    return f"{base}, {suffix}."


def _format_ieee_reference(reference: dict[str, Any]) -> str:
    authors = reference["authors"]["ieee"]
    title = reference["title"]
    source_type = reference["source_type"]
    venue = reference["venue"]["ieee"]
    year = reference["year"]
    doi = reference.get("doi")
    pages = reference.get("pages")
    volume = reference.get("volume")
    issue = reference.get("issue")
    arxiv_id = reference.get("arxiv_id")

    if source_type == "journal":
        entry = f'{authors}, "{title}," {venue}'
        if volume:
            entry += f", vol. {volume}"
        if issue:
            entry += f", no. {issue}"
        if pages:
            entry += f", pp. {pages}"
        entry += f", {year}"
        if doi:
            entry += f", doi: {doi}"
        return entry + "."

    if source_type == "conference":
        entry = f'{authors}, "{title}," in {venue}'
        if pages:
            entry += f", pp. {pages}"
        entry += f", {year}"
        if doi:
            entry += f", doi: {doi}"
        return entry + "."

    entry = f'{authors}, "{title}," arXiv preprint arXiv:{arxiv_id}, {year}'
    if doi:
        entry += f", doi: {doi}"
    return entry + "."


def _format_springer_reference(reference: dict[str, Any]) -> str:
    authors = reference["authors"]["springer"]
    title = reference["title"]
    source_type = reference["source_type"]
    venue = reference["venue"]["springer"]
    year = reference["year"]
    doi = reference.get("doi")
    doi_url = f"https://doi.org/{doi}" if doi else None
    pages = reference.get("pages")
    volume = reference.get("volume")
    issue = reference.get("issue")
    arxiv_id = reference.get("arxiv_id")

    if source_type == "journal":
        entry = f"{authors} ({year}) {title}. {venue}"
        if volume:
            entry += f" {volume}"
        if issue:
            entry += f"({issue})"
        if pages:
            entry += f":{pages}"
        entry += "."
        return _append_optional_suffix(entry, doi_url)

    if source_type == "conference":
        entry = f"{authors} ({year}) {title}. In: {venue}"
        if pages:
            entry += f", pp {pages}"
        entry += "."
        return _append_optional_suffix(entry, doi_url)

    return f"{authors} ({year}) {title}. arXiv preprint arXiv:{arxiv_id}."


def _format_mdpi_reference(reference: dict[str, Any]) -> str:
    authors = reference["authors"]["mdpi"]
    title = reference["title"]
    source_type = reference["source_type"]
    venue = reference["venue"]["mdpi"]
    year = reference["year"]
    doi = reference.get("doi")
    doi_url = f"https://doi.org/{doi}" if doi else None
    pages = reference.get("pages")
    volume = reference.get("volume")
    issue = reference.get("issue")
    arxiv_id = reference.get("arxiv_id")

    if source_type == "journal":
        entry = f"{authors} {title}. {venue} {year}"
        if volume:
            entry += f", {volume}"
            if issue:
                entry += f"({issue})"
        if pages:
            entry += f", {pages}"
        entry += "."
        return _append_optional_suffix(entry, doi_url)

    if source_type == "conference":
        entry = f"{authors} {title}. In {venue}; {year}"
        if pages:
            entry += f"; pp. {pages}"
        entry += "."
        return _append_optional_suffix(entry, doi_url)

    return f"{authors} {title}. arXiv {year}, arXiv:{arxiv_id}."


def _format_elsevier_reference(reference: dict[str, Any]) -> str:
    authors = reference["authors"]["ieee"]
    title = reference["title"]
    source_type = reference["source_type"]
    venue = reference["venue"]["ieee"]
    year = reference["year"]
    doi = reference.get("doi")
    pages = reference.get("pages")
    volume = reference.get("volume")
    issue = reference.get("issue")
    arxiv_id = reference.get("arxiv_id")

    if source_type == "journal":
        entry = f"{authors}, {title}, {venue}"
        if volume:
            entry += f" {volume}"
        if issue:
            entry += f" ({issue})"
        if pages:
            entry += f" ({year}) {pages}"
        else:
            entry += f" ({year})"
        if doi:
            entry += f". https://doi.org/{doi}"
        return entry + "."

    if source_type == "conference":
        entry = f"{authors}, {title}, in {venue}, {year}"
        if pages:
            entry += f", pp. {pages}"
        if doi:
            entry += f". https://doi.org/{doi}"
        return entry + "."

    return f"{authors}, {title}, arXiv:{arxiv_id}, {year}."


def _render_references(references: list[dict[str, Any]], style: str) -> list[str]:
    formatter_map = {
        "ieee": _format_ieee_reference,
        "springer": _format_springer_reference,
        "mdpi": _format_mdpi_reference,
        "elsevier": _format_elsevier_reference,
    }
    formatter = formatter_map[style]
    return [formatter(reference) for reference in references]


def _extend_section_paragraphs(sections: list[dict[str, Any]], heading_keyword: str, paragraphs: list[str]) -> None:
    target = heading_keyword.lower()
    for section in sections:
        if target in section["heading"].lower():
            section.setdefault("paragraphs", [])
            section["paragraphs"].extend(paragraphs)
            return
    raise ValueError(f"Could not find section containing heading keyword: {heading_keyword}")


def _append_manual_review_sections(
    sections: list[dict[str, Any]],
    *,
    reviewer_focus: list[str],
    response_points: list[str],
    top_experiments: list[str],
) -> None:
    start_number = len(sections) + 1
    sections.extend(
        [
            {
                "heading": f"{start_number}. 审稿人真正关心什么",
                "bullets": reviewer_focus,
            },
            {
                "heading": f"{start_number + 1}. 你应该准备什么回应",
                "bullets": response_points,
            },
            {
                "heading": f"{start_number + 2}. 最值当先补的实验（Top 3）",
                "bullets": top_experiments,
            },
        ]
    )


def _append_experiment_priority_section(sections: list[dict[str, Any]], bullets: list[str]) -> None:
    sections.append(
        {
            "heading": f"{len(sections) + 1}. 为什么先做",
            "bullets": bullets,
        }
    )


def _replace_section_content(
    sections: list[dict[str, Any]],
    heading_keyword: str,
    *,
    paragraphs: list[str] | None = None,
    bullets: list[str] | None = None,
    tables: list[dict[str, Any]] | None = None,
) -> None:
    target = heading_keyword.lower()
    for section in sections:
        if target in section["heading"].lower():
            if paragraphs is not None:
                section["paragraphs"] = paragraphs
            if bullets is not None:
                section["bullets"] = bullets
            if tables is not None:
                section["tables"] = tables
            return
    raise ValueError(f"Could not find section containing heading keyword: {heading_keyword}")


def _fmt_decimal(value: Any, *, digits: int = 4, pending: str = "[pending]") -> str:
    try:
        if value is None:
            return pending
        return f"{float(value):.{digits}f}"
    except Exception:
        return pending


def _fmt_percent(value: Any, *, digits: int = 2, pending: str = "[pending]") -> str:
    try:
        if value is None:
            return pending
        return f"{float(value) * 100:.{digits}f}%"
    except Exception:
        return pending


def _fmt_interval(values: Any, *, digits: int = 4, percent: bool = False, pending: str = "[pending]") -> str:
    if not isinstance(values, (list, tuple)) or len(values) != 2:
        return pending
    formatter = _fmt_percent if percent else _fmt_decimal
    left = formatter(values[0], digits=digits, pending=pending)
    right = formatter(values[1], digits=digits, pending=pending)
    if pending in {left, right}:
        return pending
    return f"[{left}, {right}]"


def _maybe_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    try:
        return _load_json(path)
    except Exception:
        return {}


def _parse_paper_run_result_path(root: Path, path: Path) -> dict[str, Any]:
    details = {"formula": None, "sample": None, "group": None, "seed": None}
    try:
        relative = path.relative_to(root)
    except ValueError:
        return details
    parts = list(relative.parts)
    for part in parts:
        if part.startswith("formula_"):
            details["formula"] = part.replace("formula_", "", 1)
        elif part.startswith("sample_"):
            details["sample"] = part.replace("sample_", "", 1)
    if "results" in parts:
        idx = parts.index("results")
        if idx + 1 < len(parts):
            details["group"] = parts[idx + 1]
        if idx + 2 < len(parts) and parts[idx + 2].startswith("seed"):
            details["seed"] = parts[idx + 2]
    return details


def _load_paper_run_context() -> dict[str, Any]:
    root = PROJECT_ROOT / "storage" / "quant" / "rl-experiments" / "paper-run"
    protocol_dir = root / "protocol"
    summary_dir = root / "summary"
    manifest_path = protocol_dir / "expected_run_manifest.json"
    root_summary_path = summary_dir / "esg_contribution_report.json"
    frozen_input_paths = {
        "latest": protocol_dir / "frozen_inputs_latest.json",
        "full": protocol_dir / "frozen_inputs_full_2022_2025.json",
        "post": protocol_dir / "frozen_inputs_post_esg_effective.json",
    }
    expected_manifest = _maybe_json(manifest_path)
    expected_run_count = int(expected_manifest.get("expected_run_count") or 0)

    metrics_paths = sorted(root.rglob("metrics.json"))
    run_status_paths = sorted(root.rglob("run_status.json"))
    equity_paths = sorted(root.rglob("equity_curve.csv"))
    metric_runs = [_parse_paper_run_result_path(root, path) for path in metrics_paths]

    available_groups = sorted({item["group"] for item in metric_runs if item.get("group")})
    available_formulas = sorted({item["formula"] for item in metric_runs if item.get("formula")})
    available_samples = sorted({item["sample"] for item in metric_runs if item.get("sample")})
    latest_metrics = []
    for path in sorted(metrics_paths, key=lambda current: current.stat().st_mtime, reverse=True)[:5]:
        item = _parse_paper_run_result_path(root, path)
        latest_metrics.append(
            "/".join(
                part
                for part in (
                    item.get("formula"),
                    item.get("sample"),
                    item.get("group"),
                    item.get("seed"),
                )
                if part
            )
        )

    verification: dict[str, Any] | None = None
    verification_error: str | None = None
    completed_run_count = 0
    if manifest_path.exists():
        try:
            from scripts.quant_rl_expected_run_manifest import verify_expected_manifest

            verification = verify_expected_manifest(
                manifest_path=manifest_path,
                report_path=None,
                require_completed_status=True,
            )
            completed_run_count = max(
                0,
                int(verification.get("expected_run_count") or expected_run_count)
                - int(verification.get("missing_count") or 0)
                - int(verification.get("failed_status_count") or 0),
            )
        except Exception as exc:
            verification_error = str(exc)

    progress_run_count = completed_run_count or len(metrics_paths)
    progress_pct = round((progress_run_count / expected_run_count) * 100, 1) if expected_run_count else 0.0
    if expected_run_count and completed_run_count >= expected_run_count > 0:
        result_state = "complete_results"
    elif metrics_paths or run_status_paths or equity_paths:
        result_state = "partial_results"
    else:
        result_state = "no_results"

    reports: dict[str, dict[str, Any]] = {}
    report_refresh_error: str | None = None
    if result_state != "no_results":
        try:
            from scripts.quant_rl_esg_contribution_report import (
                _load_equity_curves as _quant_load_equity_curves,
                _paired as _quant_paired,
                _paired_equity_bootstrap as _quant_paired_equity_bootstrap,
                build_report as _quant_build_report,
            )

            def _build_report(formula: str, sample: str, key: str) -> None:
                results_root = root / f"formula_{formula}" / f"sample_{sample}" / "results"
                if not results_root.exists() or not any(results_root.rglob("metrics.json")):
                    return
                output_dir = root / f"formula_{formula}" / f"sample_{sample}" / "summary"
                payload = _quant_build_report(
                    results_root,
                    output_dir,
                    metadata={
                        "run_namespace": "paper-run",
                        "sample": sample,
                        "formula_mode": formula,
                    },
                )
                curves = _quant_load_equity_curves(results_root)
                rows = payload.get("rows") or []
                payload["generated_comparisons"] = {
                    "ours_vs_b3_sharpe": _quant_paired(rows, "B3_sac_noesg", "OURS_full", "sharpe_ratio"),
                    "ours_vs_b4_sharpe": _quant_paired(rows, "B4_sac_esg", "OURS_full", "sharpe_ratio"),
                    "b4_vs_b3_sharpe": _quant_paired(rows, "B3_sac_noesg", "B4_sac_esg", "sharpe_ratio"),
                    "6a_vs_ours_sharpe": _quant_paired(rows, "6a_no_esg_obs", "OURS_full", "sharpe_ratio"),
                    "6b_vs_ours_sharpe": _quant_paired(rows, "6b_no_esg_reward", "OURS_full", "sharpe_ratio"),
                    "6c_vs_ours_sharpe": _quant_paired(rows, "6c_no_regime", "OURS_full", "sharpe_ratio"),
                    "6a_vs_ours_mdd": _quant_paired(rows, "6a_no_esg_obs", "OURS_full", "max_drawdown"),
                    "6b_vs_ours_mdd": _quant_paired(rows, "6b_no_esg_reward", "OURS_full", "max_drawdown"),
                    "6c_vs_ours_mdd": _quant_paired(rows, "6c_no_regime", "OURS_full", "max_drawdown"),
                }
                payload["generated_equity_comparisons"] = {
                    "ours_vs_b3_curve": _quant_paired_equity_bootstrap(curves, "B3_sac_noesg", "OURS_full"),
                    "ours_vs_b4_curve": _quant_paired_equity_bootstrap(curves, "B4_sac_esg", "OURS_full"),
                    "b4_vs_b3_curve": _quant_paired_equity_bootstrap(curves, "B3_sac_noesg", "B4_sac_esg"),
                }
                reports[key] = payload

            _build_report("v2_1", "full_2022_2025", "main")
            _build_report("v2", "full_2022_2025", "robust_formula_v2")
            _build_report("v2_1", "post_esg_effective", "robust_post_effective")
        except Exception as exc:
            report_refresh_error = str(exc)

    return {
        "root": _relative_path(root),
        "manifest_path": _relative_path(manifest_path) if manifest_path.exists() else None,
        "root_summary_path": _relative_path(root_summary_path) if root_summary_path.exists() else None,
        "summary_dir": _relative_path(summary_dir) if summary_dir.exists() else _relative_path(summary_dir),
        "expected_run_count": expected_run_count,
        "metrics_count": len(metrics_paths),
        "run_status_count": len(run_status_paths),
        "equity_curve_count": len(equity_paths),
        "completed_run_count": completed_run_count,
        "progress_run_count": progress_run_count,
        "progress_pct": progress_pct,
        "progress_label": f"{progress_run_count}/{expected_run_count}" if expected_run_count else str(progress_run_count),
        "result_state": result_state,
        "available_groups": available_groups,
        "available_formulas": available_formulas,
        "available_samples": available_samples,
        "latest_metrics": [item for item in latest_metrics if item],
        "verification": verification,
        "verification_error": verification_error,
        "reports": reports,
        "report_refresh_error": report_refresh_error,
        "root_summary": _maybe_json(root_summary_path),
        "frozen_input_paths": {
            key: (_relative_path(path) if path.exists() else None) for key, path in frozen_input_paths.items()
        },
        "frozen_inputs": {
            key: _maybe_json(path) for key, path in frozen_input_paths.items()
        },
    }


def _formal_bundle_specs(context: dict[str, Any]) -> list[dict[str, Any]]:
    historical_title = context["historical_paper_title"]
    lora_summary = context["lora_summary"]
    colpali_hits = context["colpali_hits"]
    colpali_hit_count = len(colpali_hits)
    paper_run_context = context["paper_run_context"]
    paper_01_reference_meta = [
        {
            "authors": {
                "ieee": "P. Lewis et al.",
                "springer": "Lewis P, Perez E, Piktus A, et al",
                "mdpi": "Lewis, P.; Perez, E.; Piktus, A.; et al.",
            },
            "title": "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. Adv. Neural Inf. Process. Syst.",
                "springer": "Advances in Neural Information Processing Systems",
                "mdpi": "Adv. Neural Inf. Process. Syst.",
            },
            "year": "2020",
            "pages": "9459-9474",
        },
        {
            "authors": {
                "ieee": "S. Robertson and H. Zaragoza",
                "springer": "Robertson S, Zaragoza H",
                "mdpi": "Robertson, S.; Zaragoza, H.",
            },
            "title": "The probabilistic relevance framework: BM25 and beyond",
            "source_type": "journal",
            "venue": {
                "ieee": "Found. Trends Inf. Retr.",
                "springer": "Foundations and Trends in Information Retrieval",
                "mdpi": "Found. Trends Inf. Retr.",
            },
            "year": "2009",
            "volume": "3",
            "issue": "4",
            "pages": "333-389",
            "doi": "10.1561/1500000019",
        },
        {
            "authors": {
                "ieee": "G. V. Cormack, C. L. A. Clarke, and S. Buettcher",
                "springer": "Cormack GV, Clarke CLA, Buettcher S",
                "mdpi": "Cormack, G.V.; Clarke, C.L.A.; Buettcher, S.",
            },
            "title": "Reciprocal rank fusion outperforms Condorcet and individual rank learning methods",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. 32nd Int. ACM SIGIR Conf. Res. Develop. Inf. Retr.",
                "springer": "Proceedings of the 32nd International ACM SIGIR Conference on Research and Development in Information Retrieval",
                "mdpi": "Proc. 32nd Int. ACM SIGIR Conf. Res. Dev. Inf. Retr.",
            },
            "year": "2009",
            "pages": "758-759",
            "doi": "10.1145/1571941.1572114",
        },
        {
            "authors": {
                "ieee": "V. Karpukhin et al.",
                "springer": "Karpukhin V, Oguz B, Min S, et al",
                "mdpi": "Karpukhin, V.; Oguz, B.; Min, S.; et al.",
            },
            "title": "Dense passage retrieval for open-domain question answering",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. Conf. Empirical Methods Natural Language Process.",
                "springer": "Proceedings of the 2020 Conference on Empirical Methods in Natural Language Processing",
                "mdpi": "Proc. Conf. Empir. Methods Nat. Lang. Process.",
            },
            "year": "2020",
            "pages": "6769-6781",
            "doi": "10.18653/v1/2020.emnlp-main.550",
        },
        {
            "authors": {
                "ieee": "G. Izacard and E. Grave",
                "springer": "Izacard G, Grave E",
                "mdpi": "Izacard, G.; Grave, E.",
            },
            "title": "Leveraging passage retrieval with generative models for open domain question answering",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. 16th Conf. Eur. Chapter Assoc. Comput. Linguistics",
                "springer": "Proceedings of the 16th Conference of the European Chapter of the Association for Computational Linguistics",
                "mdpi": "Proc. Conf. Eur. Chapter Assoc. Comput. Linguist.",
            },
            "year": "2021",
        },
        {
            "authors": {
                "ieee": "S. Borgeaud et al.",
                "springer": "Borgeaud S, Mensch A, Hoffmann J, et al",
                "mdpi": "Borgeaud, S.; Mensch, A.; Hoffmann, J.; et al.",
            },
            "title": "Improving language models by retrieving from trillions of tokens",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. 39th Int. Conf. Mach. Learn.",
                "springer": "Proceedings of the 39th International Conference on Machine Learning",
                "mdpi": "Proc. Int. Conf. Mach. Learn.",
            },
            "year": "2022",
        },
        {
            "authors": {
                "ieee": "N. Thakur et al.",
                "springer": "Thakur N, Reimers N, Ruckle A, et al",
                "mdpi": "Thakur, N.; Reimers, N.; Ruckle, A.; et al.",
            },
            "title": "BEIR: A heterogeneous benchmark for zero-shot evaluation of information retrieval models",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. NeurIPS Datasets Benchmarks Track",
                "springer": "Proceedings of the NeurIPS Datasets and Benchmarks Track",
                "mdpi": "Proc. NeurIPS Datasets Benchmarks Track",
            },
            "year": "2021",
        },
        {
            "authors": {
                "ieee": "Y. Gao et al.",
                "springer": "Gao Y, Xiong Y, Gao X, et al",
                "mdpi": "Gao, Y.; Xiong, Y.; Gao, X.; et al.",
            },
            "title": "Retrieval-augmented generation for large language models: A survey",
            "source_type": "arxiv",
            "venue": {
                "ieee": "arXiv",
                "springer": "arXiv",
                "mdpi": "arXiv",
            },
            "year": "2023",
            "arxiv_id": "2312.10997",
        },
        {
            "authors": {
                "ieee": "Y. Huang and J. Huang",
                "springer": "Huang Y, Huang J",
                "mdpi": "Huang, Y.; Huang, J.",
            },
            "title": "A survey on retrieval-augmented text generation for large language models",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2024",
            "arxiv_id": "2404.10981",
        },
        {
            "authors": {
                "ieee": "A. Asai, Z. Wu, Y. Wang, A. Sil, and H. Hajishirzi",
                "springer": "Asai A, Wu Z, Wang Y, Sil A, Hajishirzi H",
                "mdpi": "Asai, A.; Wu, Z.; Wang, Y.; Sil, A.; Hajishirzi, H.",
            },
            "title": "Self-RAG: Learning to retrieve, generate, and critique through self-reflection",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2023",
            "arxiv_id": "2310.11511",
        },
        {
            "authors": {
                "ieee": "S. Yan, C. Pang, X. Jia, and D. Zeng",
                "springer": "Yan S, Pang C, Jia X, Zeng D",
                "mdpi": "Yan, S.; Pang, C.; Jia, X.; Zeng, D.",
            },
            "title": "Corrective retrieval augmented generation",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2024",
            "arxiv_id": "2401.15884",
        },
        {
            "authors": {
                "ieee": "C. He et al.",
                "springer": "He C, Zhou X, Wu Y, et al",
                "mdpi": "He, C.; Zhou, X.; Wu, Y.; et al.",
            },
            "title": "ESGenius: Benchmarking LLMs on environmental, social, and governance (ESG) and sustainability knowledge",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2025",
            "arxiv_id": "2506.01646",
        },
    ]
    paper_02_reference_meta = [
        {
            "authors": {
                "ieee": "E. J. Hu et al.",
                "springer": "Hu EJ, Shen Y, Wallis P, Allen-Zhu Z, Li Y, Wang S, Wang L, Chen W",
                "mdpi": "Hu, E.J.; Shen, Y.; Wallis, P.; Allen-Zhu, Z.; Li, Y.; Wang, S.; Wang, L.; Chen, W.",
            },
            "title": "LoRA: Low-rank adaptation of large language models",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. Int. Conf. Learn. Represent.",
                "springer": "International Conference on Learning Representations",
                "mdpi": "Int. Conf. Learn. Represent.",
            },
            "year": "2022",
        },
        {
            "authors": {
                "ieee": "T. Dettmers, A. Pagnoni, A. Holtzman, and L. Zettlemoyer",
                "springer": "Dettmers T, Pagnoni A, Holtzman A, Zettlemoyer L",
                "mdpi": "Dettmers, T.; Pagnoni, A.; Holtzman, A.; Zettlemoyer, L.",
            },
            "title": "QLoRA: Efficient finetuning of quantized LLMs",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. Adv. Neural Inf. Process. Syst.",
                "springer": "Advances in Neural Information Processing Systems",
                "mdpi": "Adv. Neural Inf. Process. Syst.",
            },
            "year": "2023",
        },
        {
            "authors": {
                "ieee": "X. L. Li and P. Liang",
                "springer": "Li XL, Liang P",
                "mdpi": "Li, X.L.; Liang, P.",
            },
            "title": "Prefix-tuning: Optimizing continuous prompts for generation",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. 59th Annu. Meeting Assoc. Comput. Linguistics",
                "springer": "Proceedings of the 59th Annual Meeting of the Association for Computational Linguistics",
                "mdpi": "Proc. Annu. Meet. Assoc. Comput. Linguist.",
            },
            "year": "2021",
            "pages": "4582-4597",
            "doi": "10.18653/v1/2021.acl-long.353",
        },
        {
            "authors": {
                "ieee": "L. Ouyang et al.",
                "springer": "Ouyang L, Wu J, Jiang X, et al",
                "mdpi": "Ouyang, L.; Wu, J.; Jiang, X.; et al.",
            },
            "title": "Training language models to follow instructions with human feedback",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2022",
            "arxiv_id": "2203.02155",
        },
        {
            "authors": {
                "ieee": "Y. Wang et al.",
                "springer": "Wang Y, Kordi Y, Mishra S, Liu A, Smith NA, Khashabi D, Hajishirzi H",
                "mdpi": "Wang, Y.; Kordi, Y.; Mishra, S.; Liu, A.; Smith, N.A.; Khashabi, D.; Hajishirzi, H.",
            },
            "title": "Self-Instruct: Aligning language models with self-generated instructions",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. 61st Annu. Meeting Assoc. Comput. Linguistics",
                "springer": "Proceedings of the 61st Annual Meeting of the Association for Computational Linguistics",
                "mdpi": "Proc. Annu. Meet. Assoc. Comput. Linguist.",
            },
            "year": "2023",
            "pages": "13484-13508",
            "doi": "10.18653/v1/2023.acl-long.754",
        },
        {
            "authors": {
                "ieee": "A. Yang et al.",
                "springer": "Yang A, Yang B, Zhang B, et al",
                "mdpi": "Yang, A.; Yang, B.; Zhang, B.; et al.",
            },
            "title": "Qwen2.5 technical report",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2024",
            "arxiv_id": "2412.15115",
        },
        {
            "authors": {
                "ieee": "C.-Y. Lin",
                "springer": "Lin CY",
                "mdpi": "Lin, C.-Y.",
            },
            "title": "ROUGE: A package for automatic evaluation of summaries",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. ACL Workshop Text Summarization Branches Out",
                "springer": "Proceedings of the ACL Workshop on Text Summarization Branches Out",
                "mdpi": "Proc. ACL Workshop Text Summ. Branches Out",
            },
            "year": "2004",
            "pages": "74-81",
        },
        {
            "authors": {
                "ieee": "S. Min et al.",
                "springer": "Min S, Krishna K, Lyu X, et al",
                "mdpi": "Min, S.; Krishna, K.; Lyu, X.; et al.",
            },
            "title": "FActScore: Fine-grained atomic evaluation of factual precision in long form text generation",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. Conf. Empirical Methods Natural Language Process.",
                "springer": "Proceedings of the 2023 Conference on Empirical Methods in Natural Language Processing",
                "mdpi": "Proc. Conf. Empir. Methods Nat. Lang. Process.",
            },
            "year": "2023",
        },
        {
            "authors": {
                "ieee": "Y. Liu, D. Iter, Y. Xu, S. Wang, R. Xu, and C. Zhu",
                "springer": "Liu Y, Iter D, Xu Y, Wang S, Xu R, Zhu C",
                "mdpi": "Liu, Y.; Iter, D.; Xu, Y.; Wang, S.; Xu, R.; Zhu, C.",
            },
            "title": "G-Eval: NLG evaluation using GPT-4 with better human alignment",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. Conf. Empirical Methods Natural Language Process.",
                "springer": "Proceedings of the 2023 Conference on Empirical Methods in Natural Language Processing",
                "mdpi": "Proc. Conf. Empir. Methods Nat. Lang. Process.",
            },
            "year": "2023",
        },
        {
            "authors": {
                "ieee": "S. Wu et al.",
                "springer": "Wu S, Irsoy O, Lu S, et al",
                "mdpi": "Wu, S.; Irsoy, O.; Lu, S.; et al.",
            },
            "title": "BloombergGPT: A large language model for finance",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2023",
            "arxiv_id": "2303.17564",
        },
        {
            "authors": {
                "ieee": "X.-Y. Liu, G. Wang, H. Yang, and D. Zha",
                "springer": "Liu XY, Wang G, Yang H, Zha D",
                "mdpi": "Liu, X.-Y.; Wang, G.; Yang, H.; Zha, D.",
            },
            "title": "FinGPT: Democratizing internet-scale data for financial large language models",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2023",
            "arxiv_id": "2307.10485",
        },
        {
            "authors": {
                "ieee": "C. He et al.",
                "springer": "He C, Zhou X, Wu Y, et al",
                "mdpi": "He, C.; Zhou, X.; Wu, Y.; et al.",
            },
            "title": "ESGenius: Benchmarking LLMs on environmental, social, and governance (ESG) and sustainability knowledge",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2025",
            "arxiv_id": "2506.01646",
        },
    ]
    paper_03_reference_meta = [
        {
            "authors": {
                "ieee": "J. Wei et al.",
                "springer": "Wei J, Wang X, Schuurmans D, et al",
                "mdpi": "Wei, J.; Wang, X.; Schuurmans, D.; et al.",
            },
            "title": "Chain-of-thought prompting elicits reasoning in large language models",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. Adv. Neural Inf. Process. Syst.",
                "springer": "Advances in Neural Information Processing Systems",
                "mdpi": "Adv. Neural Inf. Process. Syst.",
            },
            "year": "2022",
        },
        {
            "authors": {
                "ieee": "L. Wang et al.",
                "springer": "Wang L, Ma C, Feng X, et al",
                "mdpi": "Wang, L.; Ma, C.; Feng, X.; et al.",
            },
            "title": "Plan-and-solve prompting: Improving zero-shot chain-of-thought reasoning by large language models",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. 61st Annu. Meeting Assoc. Comput. Linguistics",
                "springer": "Proceedings of the 61st Annual Meeting of the Association for Computational Linguistics",
                "mdpi": "Proc. Annu. Meet. Assoc. Comput. Linguist.",
            },
            "year": "2023",
        },
        {
            "authors": {
                "ieee": "S. Yao et al.",
                "springer": "Yao S, Yu D, Zhao J, et al",
                "mdpi": "Yao, S.; Yu, D.; Zhao, J.; et al.",
            },
            "title": "Tree of thoughts: Deliberate problem solving with large language models",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. Adv. Neural Inf. Process. Syst.",
                "springer": "Advances in Neural Information Processing Systems",
                "mdpi": "Adv. Neural Inf. Process. Syst.",
            },
            "year": "2023",
        },
        {
            "authors": {
                "ieee": "S. Yao et al.",
                "springer": "Yao S, Zhao J, Yu D, et al",
                "mdpi": "Yao, S.; Zhao, J.; Yu, D.; et al.",
            },
            "title": "ReAct: Synergizing reasoning and acting in language models",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. Int. Conf. Learn. Represent.",
                "springer": "International Conference on Learning Representations",
                "mdpi": "Int. Conf. Learn. Represent.",
            },
            "year": "2023",
        },
        {
            "authors": {
                "ieee": "T. Schick et al.",
                "springer": "Schick T, Dwivedi-Yu J, Dessi R, et al",
                "mdpi": "Schick, T.; Dwivedi-Yu, J.; Dessi, R.; et al.",
            },
            "title": "Toolformer: Language models can teach themselves to use tools",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. Adv. Neural Inf. Process. Syst.",
                "springer": "Advances in Neural Information Processing Systems",
                "mdpi": "Adv. Neural Inf. Process. Syst.",
            },
            "year": "2023",
        },
        {
            "authors": {
                "ieee": "A. Madaan et al.",
                "springer": "Madaan A, Tandon N, Clark P, et al",
                "mdpi": "Madaan, A.; Tandon, N.; Clark, P.; et al.",
            },
            "title": "Self-Refine: Iterative refinement with self-feedback",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. Adv. Neural Inf. Process. Syst.",
                "springer": "Advances in Neural Information Processing Systems",
                "mdpi": "Adv. Neural Inf. Process. Syst.",
            },
            "year": "2023",
        },
        {
            "authors": {
                "ieee": "N. Shinn et al.",
                "springer": "Shinn N, Cassano F, Gopinath A, et al",
                "mdpi": "Shinn, N.; Cassano, F.; Gopinath, A.; et al.",
            },
            "title": "Reflexion: Language agents with verbal reinforcement learning",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. Adv. Neural Inf. Process. Syst.",
                "springer": "Advances in Neural Information Processing Systems",
                "mdpi": "Adv. Neural Inf. Process. Syst.",
            },
            "year": "2023",
        },
        {
            "authors": {
                "ieee": "Q. Wu et al.",
                "springer": "Wu Q, Bansal G, Zhang J, et al",
                "mdpi": "Wu, Q.; Bansal, G.; Zhang, J.; et al.",
            },
            "title": "AutoGen: Enabling next-gen LLM applications via multi-agent conversation",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2023",
            "arxiv_id": "2308.08155",
        },
        {
            "authors": {
                "ieee": "G. Li et al.",
                "springer": "Li G, Hammoud HAAK, Itani H, et al",
                "mdpi": "Li, G.; Hammoud, H.A.A.K.; Itani, H.; et al.",
            },
            "title": "CAMEL: Communicative agents for mind exploration of large language model society",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2023",
            "arxiv_id": "2303.17760",
        },
        {
            "authors": {
                "ieee": "S. Hong et al.",
                "springer": "Hong S, Zhuge M, Chen J, et al",
                "mdpi": "Hong, S.; Zhuge, M.; Chen, J.; et al.",
            },
            "title": "MetaGPT: Meta programming for a multi-agent collaborative framework",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2023",
            "arxiv_id": "2308.00352",
        },
        {
            "authors": {
                "ieee": "G. Chen et al.",
                "springer": "Chen G, Li S, Chang L, et al",
                "mdpi": "Chen, G.; Li, S.; Chang, L.; et al.",
            },
            "title": "AgentVerse: Facilitating multi-agent collaboration and exploring emergent behaviors",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2023",
            "arxiv_id": "2308.10848",
        },
        {
            "authors": {
                "ieee": "C. He et al.",
                "springer": "He C, Zhou X, Wu Y, et al",
                "mdpi": "He, C.; Zhou, X.; Wu, Y.; et al.",
            },
            "title": "ESGenius: Benchmarking LLMs on environmental, social, and governance (ESG) and sustainability knowledge",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2025",
            "arxiv_id": "2506.01646",
        },
    ]
    paper_05_reference_meta = [
        {
            "authors": {
                "ieee": "V. Mnih et al.",
                "springer": "Mnih V, Kavukcuoglu K, Silver D, et al",
                "mdpi": "Mnih, V.; Kavukcuoglu, K.; Silver, D.; et al.",
            },
            "title": "Human-level control through deep reinforcement learning",
            "source_type": "journal",
            "venue": {"ieee": "Nature", "springer": "Nature", "mdpi": "Nature"},
            "year": "2015",
            "volume": "518",
            "issue": "7540",
            "pages": "529-533",
            "doi": "10.1038/nature14236",
        },
        {
            "authors": {
                "ieee": "T. Haarnoja et al.",
                "springer": "Haarnoja T, Zhou A, Abbeel P, Levine S",
                "mdpi": "Haarnoja, T.; Zhou, A.; Abbeel, P.; Levine, S.",
            },
            "title": "Soft actor-critic: Off-policy maximum entropy deep reinforcement learning with a stochastic actor",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. 35th Int. Conf. Mach. Learn.",
                "springer": "Proceedings of the 35th International Conference on Machine Learning",
                "mdpi": "Proceedings of the 35th International Conference on Machine Learning",
            },
            "year": "2018",
            "pages": "1861-1870",
        },
        {
            "authors": {
                "ieee": "X. Y. Liu et al.",
                "springer": "Liu X-Y, Bao W, Cao H, et al",
                "mdpi": "Liu, X.-Y.; Bao, W.; Cao, H.; et al.",
            },
            "title": "FinRL: A deep reinforcement learning library for automated stock trading in quantitative finance",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2020",
            "arxiv_id": "2011.09607",
        },
        {
            "authors": {
                "ieee": "T. Theate and D. Ernst",
                "springer": "Theate T, Ernst D",
                "mdpi": "Théate, T.; Ernst, D.",
            },
            "title": "An application of deep reinforcement learning to algorithmic trading",
            "source_type": "journal",
            "venue": {
                "ieee": "Expert Syst. Appl.",
                "springer": "Expert Systems with Applications",
                "mdpi": "Expert Syst. Appl.",
            },
            "year": "2021",
            "volume": "173",
            "pages": "114632",
            "doi": "10.1016/j.eswa.2021.114632",
        },
        {
            "authors": {
                "ieee": "A. Shavandi and M. Khedmati",
                "springer": "Shavandi A, Khedmati M",
                "mdpi": "Shavandi, A.; Khedmati, M.",
            },
            "title": "A multi-agent deep reinforcement learning framework for algorithmic trading in financial markets",
            "source_type": "journal",
            "venue": {
                "ieee": "Expert Syst. Appl.",
                "springer": "Expert Systems with Applications",
                "mdpi": "Expert Syst. Appl.",
            },
            "year": "2022",
            "volume": "208",
            "pages": "118124",
            "doi": "10.1016/j.eswa.2022.118124",
        },
        {
            "authors": {
                "ieee": "J. Moody and M. Saffell",
                "springer": "Moody J, Saffell M",
                "mdpi": "Moody, J.; Saffell, M.",
            },
            "title": "Learning to trade via direct reinforcement",
            "source_type": "journal",
            "venue": {
                "ieee": "IEEE Trans. Neural Netw.",
                "springer": "IEEE Transactions on Neural Networks",
                "mdpi": "IEEE Trans. Neural Netw.",
            },
            "year": "2001",
            "volume": "12",
            "issue": "4",
            "pages": "875-889",
            "doi": "10.1109/72.935097",
        },
        {
            "authors": {
                "ieee": "J. D. Hamilton",
                "springer": "Hamilton JD",
                "mdpi": "Hamilton, J.D.",
            },
            "title": "A new approach to the economic analysis of nonstationary time series and the business cycle",
            "source_type": "journal",
            "venue": {"ieee": "Econometrica", "springer": "Econometrica", "mdpi": "Econometrica"},
            "year": "1989",
            "volume": "57",
            "issue": "2",
            "pages": "357-384",
            "doi": "10.2307/1912559",
        },
        {
            "authors": {
                "ieee": "R. A. Jacobs et al.",
                "springer": "Jacobs RA, Jordan MI, Nowlan SJ, Hinton GE",
                "mdpi": "Jacobs, R.A.; Jordan, M.I.; Nowlan, S.J.; Hinton, G.E.",
            },
            "title": "Adaptive mixtures of local experts",
            "source_type": "journal",
            "venue": {"ieee": "Neural Comput.", "springer": "Neural Computation", "mdpi": "Neural Comput."},
            "year": "1991",
            "volume": "3",
            "issue": "1",
            "pages": "79-87",
            "doi": "10.1162/neco.1991.3.1.79",
        },
        {
            "authors": {
                "ieee": "G. Ke et al.",
                "springer": "Ke G, Meng Q, Finley T, et al",
                "mdpi": "Ke, G.; Meng, Q.; Finley, T.; et al.",
            },
            "title": "LightGBM: A highly efficient gradient boosting decision tree",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. Adv. Neural Inf. Process. Syst.",
                "springer": "Advances in Neural Information Processing Systems",
                "mdpi": "Advances in Neural Information Processing Systems",
            },
            "year": "2017",
            "pages": "3146-3154",
        },
        {
            "authors": {
                "ieee": "T. Akiba et al.",
                "springer": "Akiba T, Sano S, Yanase T, Ohta T, Koyama M",
                "mdpi": "Akiba, T.; Sano, S.; Yanase, T.; Ohta, T.; Koyama, M.",
            },
            "title": "Optuna: A next-generation hyperparameter optimization framework",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. 25th ACM SIGKDD Int. Conf. Knowl. Discovery Data Mining",
                "springer": "Proceedings of the 25th ACM SIGKDD International Conference on Knowledge Discovery and Data Mining",
                "mdpi": "Proceedings of the 25th ACM SIGKDD International Conference on Knowledge Discovery and Data Mining",
            },
            "year": "2019",
            "pages": "2623-2631",
            "doi": "10.1145/3292500.3330701",
        },
        {
            "authors": {
                "ieee": "G. Friede et al.",
                "springer": "Friede G, Busch T, Bassen A",
                "mdpi": "Friede, G.; Busch, T.; Bassen, A.",
            },
            "title": "ESG and financial performance: Aggregated evidence from more than 2000 empirical studies",
            "source_type": "journal",
            "venue": {
                "ieee": "J. Sustainable Finance Invest.",
                "springer": "Journal of Sustainable Finance & Investment",
                "mdpi": "J. Sustain. Finance Invest.",
            },
            "year": "2015",
            "volume": "5",
            "issue": "4",
            "pages": "210-233",
            "doi": "10.1080/20430795.2015.1118917",
        },
        {
            "authors": {
                "ieee": "E. F. Fama and K. R. French",
                "springer": "Fama EF, French KR",
                "mdpi": "Fama, E.F.; French, K.R.",
            },
            "title": "A five-factor asset pricing model",
            "source_type": "journal",
            "venue": {
                "ieee": "J. Financial Econ.",
                "springer": "Journal of Financial Economics",
                "mdpi": "J. Financ. Econ.",
            },
            "year": "2015",
            "volume": "116",
            "issue": "1",
            "pages": "1-22",
            "doi": "10.1016/j.jfineco.2014.10.010",
        },
        {
            "authors": {
                "ieee": "L. H. Pedersen et al.",
                "springer": "Pedersen LH, Fitzgibbons S, Pomorski L",
                "mdpi": "Pedersen, L.H.; Fitzgibbons, S.; Pomorski, L.",
            },
            "title": "Responsible investing: The ESG-efficient frontier",
            "source_type": "journal",
            "venue": {
                "ieee": "J. Financial Econ.",
                "springer": "Journal of Financial Economics",
                "mdpi": "J. Financ. Econ.",
            },
            "year": "2021",
            "volume": "142",
            "issue": "2",
            "pages": "572-597",
            "doi": "10.1016/j.jfineco.2020.11.001",
        },
        {
            "authors": {
                "ieee": "F. Berg et al.",
                "springer": "Berg F, Koelbel JF, Rigobon R",
                "mdpi": "Berg, F.; Koelbel, J.F.; Rigobon, R.",
            },
            "title": "Aggregate confusion: The divergence of ESG ratings",
            "source_type": "journal",
            "venue": {"ieee": "Rev. Finance", "springer": "Review of Finance", "mdpi": "Rev. Finance"},
            "year": "2022",
            "volume": "26",
            "issue": "6",
            "pages": "1315-1344",
            "doi": "10.1093/rof/rfac033",
        },
        {
            "authors": {
                "ieee": "P. Lewis et al.",
                "springer": "Lewis P, Perez E, Piktus A, et al",
                "mdpi": "Lewis, P.; Perez, E.; Piktus, A.; et al.",
            },
            "title": "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. Adv. Neural Inf. Process. Syst.",
                "springer": "Advances in Neural Information Processing Systems",
                "mdpi": "Advances in Neural Information Processing Systems",
            },
            "year": "2020",
            "pages": "9459-9474",
        },
    ]
    paper_04_reference_meta = [
        {
            "authors": {
                "ieee": "M. Faysse et al.",
                "springer": "Faysse M, Sibille H, Wu T, Omrani B, Viaud G, Hudelot C, Colombo P",
                "mdpi": "Faysse, M.; Sibille, H.; Wu, T.; Omrani, B.; Viaud, G.; Hudelot, C.; Colombo, P.",
            },
            "title": "ColPali: Efficient document retrieval with vision language models",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2024",
            "arxiv_id": "2407.01449",
        },
        {
            "authors": {
                "ieee": "L. Beyer et al.",
                "springer": "Beyer L, Steiner A, Pinto AS, et al",
                "mdpi": "Beyer, L.; Steiner, A.; Pinto, A.S.; et al.",
            },
            "title": "PaliGemma: A versatile 3B VLM for transfer",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2024",
            "arxiv_id": "2407.07726",
        },
        {
            "authors": {
                "ieee": "O. Khattab and M. Zaharia",
                "springer": "Khattab O, Zaharia M",
                "mdpi": "Khattab, O.; Zaharia, M.",
            },
            "title": "ColBERT: Efficient and effective passage search via contextualized late interaction over BERT",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. 43rd Int. ACM SIGIR Conf. Res. Develop. Inf. Retr.",
                "springer": "Proceedings of the 43rd International ACM SIGIR Conference on Research and Development in Information Retrieval",
                "mdpi": "Proc. Int. ACM SIGIR Conf. Res. Dev. Inf. Retr.",
            },
            "year": "2020",
            "pages": "39-48",
            "doi": "10.1145/3397271.3401075",
        },
        {
            "authors": {
                "ieee": "S. Robertson and H. Zaragoza",
                "springer": "Robertson S, Zaragoza H",
                "mdpi": "Robertson, S.; Zaragoza, H.",
            },
            "title": "The probabilistic relevance framework: BM25 and beyond",
            "source_type": "journal",
            "venue": {
                "ieee": "Found. Trends Inf. Retr.",
                "springer": "Foundations and Trends in Information Retrieval",
                "mdpi": "Found. Trends Inf. Retr.",
            },
            "year": "2009",
            "volume": "3",
            "issue": "4",
            "pages": "333-389",
            "doi": "10.1561/1500000019",
        },
        {
            "authors": {
                "ieee": "N. Thakur et al.",
                "springer": "Thakur N, Reimers N, Rücklé A, et al",
                "mdpi": "Thakur, N.; Reimers, N.; Rücklé, A.; et al.",
            },
            "title": "BEIR: A heterogeneous benchmark for zero-shot evaluation of information retrieval models",
            "source_type": "conference",
            "venue": {
                "ieee": "Proc. NeurIPS Datasets Benchmarks Track",
                "springer": "Proceedings of the NeurIPS Datasets and Benchmarks Track",
                "mdpi": "Proc. NeurIPS Datasets Benchmarks Track",
            },
            "year": "2021",
        },
        {
            "authors": {
                "ieee": "G. Kim et al.",
                "springer": "Kim G, Hong T, Yim M, et al",
                "mdpi": "Kim, G.; Hong, T.; Yim, M.; et al.",
            },
            "title": "OCR-free document understanding transformer",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2021",
            "arxiv_id": "2111.15664",
        },
        {
            "authors": {
                "ieee": "M. Mathew, D. Karatzas, and C. V. Jawahar",
                "springer": "Mathew M, Karatzas D, Jawahar CV",
                "mdpi": "Mathew, M.; Karatzas, D.; Jawahar, C.V.",
            },
            "title": "DocVQA: A dataset for VQA on document images",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2020",
            "arxiv_id": "2007.00398",
        },
        {
            "authors": {
                "ieee": "M. Mathew et al.",
                "springer": "Mathew M, Bagal V, Tito RP, Karatzas D, Valveny E, Jawahar CV",
                "mdpi": "Mathew, M.; Bagal, V.; Tito, R.P.; Karatzas, D.; Valveny, E.; Jawahar, C.V.",
            },
            "title": "InfographicVQA",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2021",
            "arxiv_id": "2104.12756",
        },
        {
            "authors": {
                "ieee": "A. Masry et al.",
                "springer": "Masry A, Long DX, Tan JQ, Joty S, Hoque E",
                "mdpi": "Masry, A.; Long, D.X.; Tan, J.Q.; Joty, S.; Hoque, E.",
            },
            "title": "ChartQA: A benchmark for question answering about charts with visual and logical reasoning",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2022",
            "arxiv_id": "2203.10244",
        },
        {
            "authors": {
                "ieee": "R. Tito, D. Karatzas, and E. Valveny",
                "springer": "Tito R, Karatzas D, Valveny E",
                "mdpi": "Tito, R.; Karatzas, D.; Valveny, E.",
            },
            "title": "Hierarchical multimodal transformers for Multi-Page DocVQA",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2022",
            "arxiv_id": "2212.05935",
        },
        {
            "authors": {
                "ieee": "L. Kang et al.",
                "springer": "Kang L, Tito R, Valveny E, Karatzas D",
                "mdpi": "Kang, L.; Tito, R.; Valveny, E.; Karatzas, D.",
            },
            "title": "Multi-Page Document Visual Question Answering using Self-Attention Scoring Mechanism",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2024",
            "arxiv_id": "2404.19024",
        },
        {
            "authors": {
                "ieee": "G. Baechler et al.",
                "springer": "Baechler G, Sunkara S, Wang M, et al",
                "mdpi": "Baechler, G.; Sunkara, S.; Wang, M.; et al.",
            },
            "title": "ScreenAI: A vision-language model for UI and infographics understanding",
            "source_type": "arxiv",
            "venue": {"ieee": "arXiv", "springer": "arXiv", "mdpi": "arXiv"},
            "year": "2024",
            "arxiv_id": "2402.04615",
        },
    ]

    paper_01_references = _render_references(paper_01_reference_meta, "ieee")
    paper_02_references = _render_references(paper_02_reference_meta, "springer")
    paper_03_references = _render_references(paper_03_reference_meta, "mdpi")
    paper_04_references = _render_references(paper_04_reference_meta, "springer")
    paper_05_references = _render_references(paper_05_reference_meta, "elsevier")

    paper_01_sections = [
        {
            "heading": "Abstract",
            "paragraphs": [
                "This paper presents an ESG question answering system that combines hybrid retrieval with service-level validation for deployment-oriented use. The architecture integrates dense retrieval, BM25 lexical matching, reciprocal rank fusion, parent-node context expansion, API-level service composition, and Railway-ready packaging within one executable stack. Rather than optimizing retrieval quality in isolation, the system is designed to preserve inspectability across retrieval, inference, and delivery boundaries. The current artifact already supports route-level contract tests, retriever robustness checks for noisy PDF-derived text, and deployment descriptors that make runtime assumptions explicit. It also keeps application behavior, retrieval diagnostics, and deployment readiness within one verifiable engineering surface. These properties establish a practical contribution: an ESG QA pipeline whose evidence path remains testable, reproducible, deployment-aware, and ready for later benchmark-based and groundedness-oriented evaluation in production-like settings under practical service constraints."
            ],
        },
        {
            "heading": "Keywords",
            "paragraphs": [
                "ESG question answering, hybrid retrieval, service-level validation, retrieval-augmented generation, deployment-ready systems"
            ],
        },
        {
            "heading": "Highlights",
            "bullets": [
                "The system combines dense retrieval, BM25, reciprocal rank fusion, and parent-node expansion within one deployable ESG QA stack.",
                "Service-level validation is treated as part of the contribution through contract tests, retriever robustness checks, and deployment descriptors.",
                "The artifact keeps retrieval behavior, application wiring, and deployment assumptions reproducible in one engineering-facing package.",
            ],
        },
        {
            "heading": "1. Introduction",
            "paragraphs": [
                "Retrieval-augmented generation has become a practical way to connect large language models with external evidence, particularly for knowledge-intensive tasks in which direct parametric recall is insufficient [1], [4], [5]. Yet many deployed ESG analysis workflows still fail at the boundary between retrieval quality, answer grounding, and runtime integration, because production concerns are handled after the model pipeline is designed rather than as part of the core system architecture.",
                "This issue is amplified for ESG reports, which are long, repetitive, and often contaminated by PDF extraction noise. Zero-shot retrieval quality alone is therefore not enough; robust systems also require strong validation surfaces, explicit service contracts, and evaluation hooks that survive deployment changes [7]-[12]. The present paper studies ESG question answering from that systems perspective and treats hybrid retrieval plus service-level validation as the primary engineering contribution.",
            ],
        },
        {
            "heading": "2. Related Work",
            "paragraphs": [
                "Early retrieval-augmented generation work established the value of combining external evidence with generative models for knowledge-intensive tasks [1]. Dense retrievers and retrieval-conditioned generators then improved open-domain QA by strengthening evidence access and synthesis under realistic corpus settings [4], [5]. Large-scale retrieval systems such as RETRO further showed that retrieval can be treated as a persistent systems component rather than as a lightweight prompt-time add-on [6].",
                "A second line of research concerns hybrid retrieval and retrieval robustness. BM25 remains a strong lexical baseline when collections contain repeated terminology and formulaic disclosures [2], while reciprocal rank fusion provides an effective way to combine heterogeneous rankers without complex score calibration [3]. BEIR highlighted the importance of robust zero-shot retrieval evaluation across domains [7], and recent surveys describe hybrid and corrective retrieval as practical responses to failure modes in deployed RAG pipelines [8], [9], [11].",
                "More recent work has turned from raw retrieval quality to end-to-end evidence use. Self-RAG introduces explicit retrieval diagnosis and self-critique into the answer pipeline [10], while corrective retrieval generation emphasizes recovery from retrieval errors during generation [11]. Emerging ESG-focused benchmarks extend this discussion into sustainability reporting, where evidence traceability and explanation quality matter as much as answer fluency [12]. Our system paper adopts that evidence-grounded perspective, but emphasizes deployability and service-level validation rather than answer-generation strategy alone.",
            ],
        },
        {
            "heading": "3. System Overview and Contribution Boundary",
            "paragraphs": [
                "The contribution boundary is intentionally system-centric. The manuscript does not introduce a new ESG scoring model, a new financial prediction algorithm, or a new theory of agents. Instead, it examines how an ESG QA application can combine hybrid retrieval, service composition, and deployment packaging without losing traceability or validation coverage.",
                "The three main contributions are summarized as follows.",
            ],
            "bullets": [
                "Contribution 1: A deployable ESG QA stack that combines dense retrieval, BM25, reciprocal rank fusion, parent-node expansion, and document-quality filtering.",
                "Contribution 2: A service-oriented application design that links retrieval, evidence handling, API routes, and frontend delivery within one executable architecture.",
                "Contribution 3: A validation protocol that uses interface contracts, retriever checks, and deployment descriptors as first-class evidence for system readiness.",
            ],
        },
        {
            "heading": "4. Hybrid Retrieval and Platform Integration",
            "paragraphs": [
                "At the application layer, the system is exposed through a FastAPI entrypoint that organizes health checks, ESG analysis routes, report-generation routes, and quant-oriented integration surfaces. This arrangement supports the engineering expectation that retrieval systems should be observable and testable at runtime rather than treated as notebook-only assets.",
                "At the retrieval layer, the implementation combines vector retrieval with BM25 lexical retrieval and merges ranked results through reciprocal rank fusion, a design motivated by the complementary strengths of dense and lexical signals in heterogeneous corpora [2]-[5]. Parent-node expansion widens local context before answer synthesis, while a quality filter removes duplicated or PDF-noisy chunks, echoing current RAG design patterns that prioritize robustness and retrieval correction over single-signal ranking [8]-[11].",
                "Figure 1 placeholder. End-to-end architecture showing ingestion, hybrid retrieval, API routes, service orchestration as a subsystem, and Railway deployment packaging.",
                "Figure 2 placeholder. Hybrid retrieval pipeline showing dense retrieval, BM25 retrieval, reciprocal rank fusion, parent-node expansion, and document-quality filtering.",
            ],
        },
        {
            "heading": "5. Validation Protocol and Current Findings",
            "paragraphs": [
                "The present evidence is organized as a validation protocol rather than a benchmark league table. Current validation covers API contract behavior, retriever stability under noisy ESG text, and deployment readiness through a concrete Railway descriptor. This framing is consistent with the system-level claim that an ESG QA stack should remain executable and inspectable across application boundaries.",
                "The strongest immediate finding is that the stack is already verifiable across multiple boundaries: health and analysis routes are covered by contract tests, retriever post-processing removes duplicates and PDF noise, and deployment assumptions are codified through a service descriptor instead of remaining implicit.",
            ],
            "tables": [
                {
                    "caption": "Table 1. Current validation evidence supporting the system-level claims.",
                    "headers": ["Validation Dimension", "Current Evidence", "Current Observation"],
                    "rows": [
                        ["API contract behavior", "Endpoint tests for health, analyze, and quant routes", "Stable route-level validation is already available"],
                        ["Retriever robustness", "Retriever runtime tests with duplicate and PDF-noisy text", "Hybrid retrieval pipeline includes explicit quality checks"],
                        ["Deployment readiness", "Railway deployment descriptor and health check configuration", "Service packaging assumptions are codified"],
                        ["Operational integration", "FastAPI entrypoint and application factory", "Retrieval and delivery layers are wired into one application"],
                    ],
                }
            ],
        },
        {
            "heading": "6. Discussion and Threats to Validity",
            "paragraphs": [
                "The current validation scope is sufficient for a submission-oriented systems draft, but not yet for a final claim of benchmark superiority. A stronger camera-ready version should add a curated ESG QA benchmark, latency measurements across deployment modes, and groundedness scoring over retrieved evidence, in line with current RAG evaluation practice and ESG-specific explainability expectations [7]-[12].",
                "Another threat to validity is that interface tests and retriever checks do not directly measure user-perceived answer quality. They establish system readiness and robustness, but they must eventually be complemented by task-level evaluation. The paper should therefore continue to present current evidence as integration validation rather than as a downstream performance leaderboard [8]-[11].",
            ],
        },
        {
            "heading": "7. Reproducibility and Artifact Availability",
            "paragraphs": [
                "The current artifact package includes the application entrypoint, application factory, hybrid retriever implementation, Qdrant persistence logic, deployment descriptor, and targeted regression tests. These assets are cataloged in the per-paper evidence manifest so that later revisions can separate scientific claims from implementation anchors.",
            ],
        },
        {
            "heading": "8. Conclusion",
            "paragraphs": [
                "A strong ESG QA systems paper should explain not only how evidence is retrieved, but also how the application remains testable, deployable, and reproducible under realistic service constraints. The current stack already supports that argument. With benchmark supplementation and task-level groundedness evaluation, the manuscript can mature into a submission-ready engineering systems paper without changing its central design claims.",
            ],
        },
        {
            "heading": "Author Contributions",
            "paragraphs": [
                "The author contributed to conceptualization, system design, software integration, validation, visualization planning, and writing of the manuscript."
            ],
        },
        {
            "heading": "Data Availability Statement",
            "paragraphs": [
                "The code, configuration, deployment descriptors, targeted tests, and manuscript-generation assets that support this study are available within the project repository. System components for application entry, hybrid retrieval, Qdrant persistence, and deployment packaging can be inspected directly in the repository artifact set. External ESG reports or source documents used to motivate downstream use cases remain subject to their original access conditions and licensing terms."
            ],
        },
        {
            "heading": "Conflict of Interest",
            "paragraphs": [
                "The author declares no conflict of interest."
            ],
        },
        {
            "heading": "References",
            "numbered_items": paper_01_references,
        },
    ]

    manual_01_sections = [
        {
            "heading": "1. 论文定位",
            "paragraphs": [
                "这篇稿子要稳定定位成系统论文，而不是白皮书，也不是多智能体论文。主线是混合检索、服务化集成、部署准备度和端到端验证。",
            ],
        },
        {
            "heading": "2. 现有可直接写入结果",
            "bullets": [
                "可直接写入系统结构：FastAPI 入口、应用工厂、混合检索、Qdrant 持久化、Railway 部署描述。",
                "可直接写入验证方式：API contract tests、retriever runtime tests、deployment descriptor。",
                "可直接写入图表占位：系统架构图、混合检索流程图、验证证据表、部署检查表。",
            ],
        },
        {
            "heading": "3. 投稿前必须补的实验",
            "bullets": [
                "补一个外部 ESG QA benchmark 或人工评测集。",
                "补系统延迟、吞吐和降级行为的量化结果。",
                "补回答 groundedness 评价，而不只是接口与检索层验证。",
                "补至少一个与纯 dense retrieval 或纯 lexical retrieval 的对照实验。",
            ],
        },
        {
            "heading": "4. 建议复现命令",
            "bullets": [
                "python -m pytest tests/test_api_contracts.py tests/test_quant_api.py -q",
                "python -m pytest tests/test_rag_retriever_runtime.py -q",
                "python -m uvicorn gateway.main:app --host 0.0.0.0 --port 8012",
            ],
        },
        {
            "heading": "5. 图表与写作提醒",
            "bullets": [
                "图1 用系统总架构图，不要再单独突出四智能体。",
                "图2 用混合检索流程图，强调 fusion 和 quality filter。",
                "表1 用当前验证证据表，表2 用部署与接口验证表。",
                "正文避免与论文03重复，不要把 router、retriever、analyst、verifier 写成主贡献列表。",
            ],
        },
    ]

    paper_02_sections = [
        {
            "heading": "Abstract",
            "paragraphs": [
                f"This study examines LoRA-based domain adaptation of Qwen2.5-7B for ESG question answering under a reproducible fine-tuning workflow. The method aligns chat-formatted supervision, parameter-efficient adaptation, and downstream answer generation within one consistent training and evaluation pipeline. The current repository already contains a validated snapshot over {lora_summary['num_samples']} samples with an average ROUGE-L of {lora_summary['avg_rougeL']}, showing that the adapted model can be evaluated reliably in a domain-specific ESG setting. This snapshot also exposes a stable evaluation boundary for prompt-formatted supervision, adapter-based updating, and ESG-specific response generation. The present evidence does not yet support a superiority claim because matched base-model baselines and groundedness-oriented evaluation are still missing. Nevertheless, the artifact already provides a transparent adaptation study in which the model configuration, data format, and evaluation pathway are reproducible, and the resulting answers can be analyzed under a clearly documented experimental boundary. The resulting study serves as a reproducible benchmark scaffold for later baseline-matched and faithfulness-aware comparison."
            ],
        },
        {
            "heading": "Keywords",
            "paragraphs": [
                "LoRA, Qwen2.5-7B, domain adaptation, ESG question answering, parameter-efficient fine-tuning"
            ],
        },
        {
            "heading": "Highlights",
            "bullets": [
                "The study documents a reproducible LoRA adaptation workflow for ESG question answering with Qwen2.5-7B-Instruct.",
                f"The current repository snapshot reports {lora_summary['num_samples']} validation samples and an average ROUGE-L of {lora_summary['avg_rougeL']}.",
                "The paper frames the current result as a reproducible evaluation boundary that can be extended with matched baselines and faithfulness checks.",
            ],
        },
        {
            "heading": "1. Introduction",
            "paragraphs": [
                "Domain-specific ESG question answering requires more precise evidence phrasing, instruction following, and disclosure-oriented language than general-purpose models typically provide without adaptation. Parameter-efficient fine-tuning has therefore become a practical path for specializing strong open models while preserving feasible training cost and deployment complexity [1]-[6].",
                "This manuscript studies a LoRA-based adaptation pipeline built around Qwen2.5-7B-Instruct and chat-formatted ESG QA supervision. The goal is not to claim final benchmark dominance, but to present a reproducible adaptation workflow whose current evaluation snapshot can already support a methodologically credible paper and whose remaining gaps are clearly bounded by prior work on automatic and factuality-aware evaluation [7]-[12].",
            ],
        },
        {
            "heading": "2. Related Work",
            "paragraphs": [
                "Parameter-efficient transfer learning is now a standard route for adapting large language models under practical hardware limits. LoRA and QLoRA reduce the cost of specialization by constraining trainable parameters, while prompt-based alternatives such as prefix-tuning offer a complementary view of lightweight adaptation [1]-[3].",
                "Instruction tuning provides the second major ingredient for this line of work. When supervision mirrors downstream user intent and output structure, response quality improves in ways that matter for conversational QA tasks [4], [5]. Our pipeline follows that logic by matching the training format to the serving format instead of treating prompt structure as a post-training concern [6].",
                "Financial and ESG-oriented language models add a stronger domain shift. BloombergGPT and FinGPT show that finance-centric corpora and tasks create adaptation pressures that differ from generic open-domain tuning [10], [11]. ESG-specific benchmark work further emphasizes that explainability and evidence sensitivity are central to sustainability-oriented QA, not optional evaluation extras [12].",
                "A final line of work concerns evaluation. ROUGE remains a common overlap-based metric for generation [7], but factuality-oriented methods such as FActScore and LLM-assisted evaluators such as G-Eval target different aspects of answer reliability [8], [9]. This paper therefore treats reproducible ROUGE reporting as a starting point rather than a complete evaluation story.",
            ],
        },
        {
            "heading": "3. Task Definition and Dataset Description",
            "paragraphs": [
                "The target task is ESG question answering formulated as chat-style answer generation. Each instance is stored as a structured message sequence that mirrors the downstream application format, allowing the same conversational schema to be reused during training and inference, which is consistent with modern instruction-tuning practice [4]-[6].",
                "Training and validation samples are stored under data/rag_training_data/train.jsonl and data/rag_training_data/val.jsonl. This choice prioritizes alignment between supervision format, serving format, and evaluation prompts rather than benchmark portability. The paper should make that design choice explicit and connect it to the broader discussion of domain data pipelines in finance-oriented LLM work [10], [11].",
                "A careful submission version should further document how the QA pairs were constructed, how the split was created, and what protections were used to reduce leakage between template variants or repeated evidence snippets. Those disclosure details are particularly important for ESG-oriented datasets because benchmark realism and answer support are closely tied [11], [12].",
            ],
        },
        {
            "heading": "4. Training Pipeline",
            "paragraphs": [
                "The training pipeline loads Qwen/Qwen2.5-7B-Instruct with low-precision weights, prepares the backbone for k-bit adaptation, and applies LoRA modules to attention and MLP projection layers. The repository-default configuration uses rank 16, alpha 32, and dropout 0.05, which follows the practical efficiency rationale established in LoRA and QLoRA-style adaptation work [1], [2].",
                "The dataset builder formats each sample through a chat template and masks prompt tokens in the label sequence, thereby focusing optimization on assistant-side answer generation. This is a suitable choice for a domain QA task in which answer faithfulness and instruction compliance matter more than unconstrained continuation quality [3]-[6].",
                "Figure 1 placeholder. Training and evaluation workflow covering chat-formatted data construction, LoRA attachment, validation-time generation, and metric export.",
            ],
        },
        {
            "heading": "5. Experimental Setup and Current Results",
            "paragraphs": [
                f"The current evaluation artifact at {lora_summary['path']} reports {lora_summary['num_samples']} validation samples and an average ROUGE-L of {lora_summary['avg_rougeL']}. We treat this result as the current reproducible evaluation snapshot rather than as a final experimental claim, because summary-overlap metrics alone cannot establish factual support or superiority [7]-[9].",
                "The evidence is strong enough to support a reproducible methods paper draft, but not yet strong enough to claim that LoRA adaptation materially outperforms the untuned base model or alternative adaptation strategies. Table 1 therefore separates the currently validated row from the baseline and faithfulness comparisons that must be added before submission [7]-[12].",
            ],
            "tables": [
                {
                    "caption": "Table 1. Current evaluation snapshot and missing comparison rows for the final paper version.",
                    "headers": ["Model Setting", "Validation Samples", "Average ROUGE-L", "Faithfulness Metric", "Status"],
                    "rows": [
                        ["LoRA-tuned Qwen2.5-7B-Instruct", str(lora_summary["num_samples"]), str(lora_summary["avg_rougeL"]), "Not yet reported", "Available in repository"],
                        ["Untuned base model", "TBD", "TBD", "TBD", "Must be added before submission"],
                        ["Human or groundedness review", "TBD", "N/A", "TBD", "Must be added before submission"],
                    ],
                }
            ],
        },
        {
            "heading": "6. Baseline and Evaluation Gap",
            "paragraphs": [
                "The most important missing experiment is a direct base-model comparison under the same prompts, decoding settings, and validation split. Without that row, the present evidence shows that the adapted model can be evaluated, but not yet that adaptation improves task performance in a controlled sense [1], [2], [10], [11].",
                "A second gap is the absence of answer-faithfulness or groundedness evaluation. ROUGE captures lexical overlap [7], whereas factuality-oriented measures such as FActScore and LLM-assisted evaluators such as G-Eval target different dimensions of answer reliability [8], [9]. A final version of the paper should therefore add either human review or evidence-grounded automatic scoring instead of relying on overlap metrics alone.",
            ],
        },
        {
            "heading": "7. Threats to Validity",
            "paragraphs": [
                "The current evidence is vulnerable to three validity threats. First, ROUGE alone may underrepresent factual correctness and grounding [7]-[9]. Second, the dataset construction process may encode template regularities that favor stylistic matching over genuine generalization. Third, without a base-model comparison, observed performance cannot yet be attributed cleanly to LoRA adaptation rather than to prompt format or data curation effects [1]-[6].",
                "These limitations do not invalidate the current manuscript, but they do define a strict boundary for the claims that can be made at this stage. The paper should therefore continue to frame current results as a reproducible evaluation snapshot rather than a final leaderboard result, while using finance- and ESG-specific prior work to justify the next experimental additions [10]-[12].",
            ],
        },
        {
            "heading": "8. Reproducibility and Artifact Availability",
            "paragraphs": [
                "The current artifact package includes the training entrypoint, evaluation entrypoint, chat-formatted train and validation data, LoRA adapter configuration, tokenizer files, and metric exports. These assets are sufficient to reproduce the currently reported LoRA row and to extend the manuscript with stronger baselines.",
            ],
        },
        {
            "heading": "9. Conclusion",
            "paragraphs": [
                "The present Qwen2.5 LoRA pipeline already supports a credible academic manuscript on ESG domain adaptation. The strongest next step is not to broaden the method story, but to tighten it with matched baselines, groundedness evaluation, and explicit dataset-lineage reporting so that the current reproducible snapshot can become a submission-ready results section.",
            ],
        },
        {
            "heading": "Author Contributions",
            "paragraphs": [
                "The author contributed to conceptualization, data curation, fine-tuning workflow design, evaluation design, result interpretation, and writing of the manuscript."
            ],
        },
        {
            "heading": "Data Availability Statement",
            "paragraphs": [
                "The training and validation JSONL files, evaluation exports, model-adaptation configuration, and manuscript-generation assets that support this study are available within the project repository. In particular, the current supervised data reside under data/rag_training_data/, and the reported evaluation snapshot is derived from artifacts under data/rag_eval/. Any external ESG source documents referenced by the dataset remain subject to their original access conditions and licensing terms."
            ],
        },
        {
            "heading": "Conflict of Interest",
            "paragraphs": [
                "The author declares no conflict of interest."
            ],
        },
        {
            "heading": "References",
            "numbered_items": paper_02_references,
        },
    ]

    manual_02_sections = [
        {
            "heading": "1. 论文定位",
            "paragraphs": [
                "这篇稿子要稳稳地定位成领域适配论文，而不是系统论文。主线是 LoRA、数据格式、训练流程和当前可用的评估结果。",
            ],
        },
        {
            "heading": "2. 现有可直接写入结果",
            "bullets": [
                f"当前可直接写入：验证样本数 {lora_summary['num_samples']}。",
                f"当前可直接写入：平均 ROUGE-L = {lora_summary['avg_rougeL']}。",
                "当前可直接写入：LoRA r=16，alpha=32，dropout=0.05。",
                "当前可直接写入：训练和评估脚本、checkpoint 和可视化文件都已存在。",
            ],
        },
        {
            "heading": "3. 投稿前必须补的实验",
            "bullets": [
                "补 untuned base model 在同一验证集上的结果。",
                "补 groundedness 或人工 factuality 评估，而不只看 ROUGE。",
                "补数据卡，明确样本来源、构造方式、切分方式和潜在泄漏风险。",
                "补至少一个错误案例分析表，解释 LoRA 目前失败在哪里。",
            ],
        },
        {
            "heading": "4. 建议复现命令",
            "bullets": [
                "python training/finetune.py --model_name Qwen/Qwen2.5-7B-Instruct --output_dir model-serving/checkpoint",
                "python training/evaluate_model.py --base_model Qwen/Qwen2.5-7B-Instruct --checkpoint model-serving/checkpoint",
                "python -m pytest tests/test_training_finetune.py tests/test_training_evaluate_model.py -q",
            ],
        },
        {
            "heading": "5. 图表与写作提醒",
            "bullets": [
                "图1 用训练与评估流程图。",
                "图2 用 ROUGE 分布或案例对比图。",
                "表1 必须保留当前 LoRA 结果和 base model 待补行。",
                "正文不要大篇幅展开系统架构和混合检索，以免和论文01重复。",
            ],
        },
    ]

    paper_03_sections = [
        {
            "heading": "Abstract",
            "paragraphs": [
                "This paper presents a typed-state multi-agent orchestration workflow for ESG question answering, with retry control and fallback execution treated as core design elements. The method emphasizes control-plane robustness rather than retrieval novelty: shared state is explicitly typed, node transitions are conditionally routed, verifier feedback can trigger corrective passes, and a sequential fallback executor preserves execution continuity when the graph runtime is unavailable. The current implementation already supports state propagation, task-sensitive routing, verifier-triggered retry behavior, and dependency-tolerant fallback paths. These properties make the workflow inspectable and reproducible under realistic runtime conditions while enabling auditable execution traces and graceful degradation during service operation. The resulting contribution is a workflow architecture that is narrow in scope but methodologically complete enough to support controlled ablations on verification, retry, and query rewriting in an ESG QA setting."
            ],
        },
        {
            "heading": "Keywords",
            "paragraphs": [
                "multi-agent orchestration, typed state, retry control, fallback execution, ESG question answering"
            ],
        },
        {
            "heading": "Highlights",
            "bullets": [
                "The workflow treats typed shared state as the core interface for routing, retrieval, analysis, and verification stages.",
                "Retry logic and fallback execution are modeled as explicit control mechanisms rather than implicit prompt behavior.",
                "The current implementation already supports auditable execution traces and graceful degradation under runtime dependency loss.",
            ],
        },
        {
            "heading": "1. Introduction",
            "paragraphs": [
                "Large-language-model workflows are often assembled as opaque prompt chains in which routing, evidence access, reasoning, and verification are compressed into a single sequential interaction. While this pattern is convenient for prototyping, it obscures control flow and makes failure analysis difficult, especially in tasks where reasoning steps and tool calls materially affect reliability [1]-[4].",
                "For ESG question answering, unsupported claims and weakly justified summaries carry outsized downstream cost. Control-plane design therefore matters as much as prompt quality. The present manuscript focuses on a multi-agent workflow in which routing, evidence acquisition, analysis, and verification are decomposed into explicit stages connected through typed shared state, retry semantics, and fallback execution logic [4]-[12].",
            ],
        },
        {
            "heading": "2. Related Work",
            "paragraphs": [
                "Reasoning-control prompting introduced the idea that intermediate reasoning steps can be made explicit rather than hidden inside a single opaque generation. Chain-of-thought prompting and plan-and-solve prompting illustrate how structured reasoning can improve transparency and controllability under complex tasks [1], [2]. Tree of Thoughts extends that direction by treating reasoning as an explicit search problem over intermediate states [3].",
                "A second line of work connects reasoning with action. ReAct couples reasoning and action selection in one loop [4], while Toolformer shows that models can learn to invoke tools when external actions are useful [5]. These studies motivate workflows in which control flow and tool use are first-class design choices rather than side effects of prompt wording.",
                "Reflective correction adds another layer of control. Self-Refine and Reflexion demonstrate that iterative self-critique and verbal reinforcement can improve outputs after an initial pass [6], [7]. Our retry mechanism borrows the intuition of corrective refinement, but implements it as explicit workflow control instead of free-form prompt iteration.",
                "Multi-agent orchestration frameworks provide the closest systems-level precedent. AutoGen, CAMEL, MetaGPT, and AgentVerse all treat agent interaction as a controllable software architecture [8]-[11]. Our work adopts that perspective but narrows the design target to ESG QA, where typed state, retry behavior, fallback continuity, and evidence-sensitive execution matter more than open-ended autonomy [12].",
            ],
        },
        {
            "heading": "3. Workflow Definition and Typed State",
            "paragraphs": [
                "The workflow is defined over a typed shared state that carries user input, query-rewrite artifacts, retrieved context, intermediate analytical outputs, confidence values, and retry counters. This design makes it possible to inspect what each node consumes and produces, rather than hiding cross-stage information inside prompts alone. Such explicit state design aligns with the broader movement from single-chain prompting toward decomposed, inspectable control structures [2]-[4], [8]-[11].",
                "The four roles are organized as router, retriever, analyst, and verifier. The router determines task class, the retriever acquires and stores evidence in shared state, the analyst produces structured interpretation when needed, and the verifier decides whether the current answer should be accepted or sent back for another pass. The main claim is not that any one role is novel, but that typed-state coordination creates a workflow whose control semantics are easier to analyze and ablate [6]-[11].",
            ],
            "bullets": [
                "Typed shared state is the core interface between nodes.",
                "Conditional routing determines whether analysis is required before verification.",
                "Verification can terminate the workflow or trigger another analytical pass.",
            ],
        },
        {
            "heading": "4. Control Flow, Retry, and Fallback Execution",
            "paragraphs": [
                "A key property of the workflow is that verifier feedback can alter control flow. When the verifier determines that an answer is insufficiently grounded or otherwise unacceptable, the process can return to the analytical stage instead of failing silently or returning an unqualified result. This design resonates with recent work on self-refinement and reflective correction, but implements the mechanism as explicit workflow logic rather than as an informal prompt convention [6]-[8].",
                "A second property is local fallback execution. When the graph runtime is unavailable, the workflow can still be executed through a sequential fallback executor that preserves the same logical ordering. This lets the paper discuss robustness under dependency degradation instead of assuming ideal runtime conditions only, and aligns with multi-agent systems work that treats orchestration infrastructure as part of the method rather than an implementation footnote [8]-[11].",
                "Figure 1 placeholder. Control-flow diagram for router, retriever, analyst, verifier, retry loops, and fallback execution paths.",
            ],
        },
        {
            "heading": "5. Validation Protocol and Planned Ablations",
            "paragraphs": [
                "Current validation evidence focuses on workflow correctness and state-handling behavior. Existing tests confirm that fallback graph execution works when the graph runtime is missing and that retriever-side state updates can be cached and reused consistently across runs. The present evidence therefore supports the control-plane claim, even though it does not yet establish large-scale answer-quality gains.",
                "For submission, the paper should add explicit ablation studies that compare the current workflow with simpler alternatives. The most useful first set would remove the verifier, remove retry behavior, or remove query rewriting while holding the rest of the application stable, which would create a direct bridge from workflow design to measurable impact [4], [6]-[8].",
            ],
            "tables": [
                {
                    "caption": "Table 1. Current validation evidence and planned ablation comparisons for the orchestration paper.",
                    "headers": ["Comparison or Check", "Current Status", "Purpose"],
                    "rows": [
                        ["Fallback graph execution", "Available", "Shows workflow continuity without the graph runtime"],
                        ["Retriever state caching", "Available", "Shows cross-stage state persistence"],
                        ["Without verifier", "Planned", "Measures the value of verification"],
                        ["Without retry", "Planned", "Measures the value of verifier-triggered correction"],
                        ["Without query rewrite", "Planned", "Measures the value of pre-retrieval normalization"],
                    ],
                }
            ],
        },
        {
            "heading": "6. Discussion and Threats to Validity",
            "paragraphs": [
                "The strongest current evidence concerns workflow structure, not final answer-quality gains. The paper should therefore avoid overstating benchmark superiority and instead frame its contribution around robustness, inspectability, and controllable execution semantics. That framing is consistent with the broader literature on reasoning control, multi-agent cooperation, and self-corrective workflows [1]-[11].",
                "A second threat to validity is that evidence acquisition remains part of the end-to-end workflow, which can blur the contribution boundary with the systems paper if described too broadly. The manuscript should continue to mention retrieval only as supporting context and keep the main emphasis on typed state, control flow, retry behavior, and fallback execution, while using ESG-domain literature only to justify task relevance [8]-[12].",
            ],
        },
        {
            "heading": "7. Artifact Availability",
            "paragraphs": [
                "The current artifact package includes the shared-state graph definition, node implementations, fallback execution logic, and focused runtime tests. These assets are sufficient for a method paper draft centered on orchestration mechanics and can later support ablation-focused evaluation without changing the workflow definition.",
            ],
        },
        {
            "heading": "8. Conclusion",
            "paragraphs": [
                "The present workflow already supports a credible manuscript on orchestration design for ESG question answering. The strongest next step is to quantify the value of its control mechanisms through targeted ablations and node-level traces, while preserving the paper's narrow focus on typed state, retry control, and fallback execution.",
            ],
        },
        {
            "heading": "Author Contributions",
            "paragraphs": [
                "The author contributed to conceptualization, workflow design, orchestration implementation, validation planning, runtime analysis, and writing of the manuscript."
            ],
        },
        {
            "heading": "Data Availability Statement",
            "paragraphs": [
                "The graph definition, node implementations, fallback execution logic, focused runtime tests, and manuscript-generation assets that support this study are available within the project repository. These materials are sufficient to inspect the typed-state workflow and reproduce the current orchestration-focused artifact package. Any external ESG documents referenced as application context remain subject to their original access conditions and licensing terms."
            ],
        },
        {
            "heading": "Conflict of Interest",
            "paragraphs": [
                "The author declares no conflict of interest."
            ],
        },
        {
            "heading": "References",
            "numbered_items": paper_03_references,
        },
    ]

    manual_03_sections = [
        {
            "heading": "1. 论文定位",
            "paragraphs": [
                "这篇稿子要收窄成编排论文，主线是 typed state、conditional routing、retry 和 fallback。不要再把它写成泛化的金融文档智能框架。",
            ],
        },
        {
            "heading": "2. 现有可直接写入结果",
            "bullets": [
                "可直接写入：四节点流程定义、共享状态字段、条件路由逻辑、fallback executor。",
                "可直接写入：fallback graph test 和 retriever state caching test。",
                "可直接写入：未来 ablation 方案的设计理由。",
            ],
        },
        {
            "heading": "3. 投稿前必须补的实验",
            "bullets": [
                "补单链路或去 verifier 的对照实验。",
                "补去 retry、去 query rewrite 的 ablation。",
                "补节点级 trace 和耗时表。",
                "补 10 到 20 个案例分析，展示 verifier 或 retry 的实际价值。",
            ],
        },
        {
            "heading": "4. 建议复现命令",
            "bullets": [
                "python -m pytest tests/test_graph_runtime.py tests/test_retriever_agent_runtime.py -q",
                "python -m uvicorn gateway.main:app --host 0.0.0.0 --port 8012",
                "调用 /agent/analyze 接口并保存一组真实 trace。",
            ],
        },
        {
            "heading": "5. 图表与去重提醒",
            "bullets": [
                "图1 用控制流图，图2 用节点 trace 图。",
                "表1 用共享状态字段定义，表2 用 ablation 计划表。",
                "正文不要展开 BM25、RRF、Qdrant 或部署细节。",
                "这篇只盯住 orchestration、state、retry、fallback，避免和论文01混淆。",
            ],
        },
    ]

    paper_01_sections.insert(
        8,
        {
            "heading": "6. Submission-Grade Experimental Design",
            "paragraphs": [
                "A submission-grade evaluation for this systems paper should combine answer quality, retrieval quality, runtime behavior, and deployment resilience in one protocol. The recommended benchmark contains 300 to 500 ESG questions spanning policy disclosures, metric extraction, trend comparison, controversy follow-up, and cross-document evidence synthesis. Queries should be split by company and document year so that repeated report language does not leak between development and held-out evaluation.",
                "The core comparison set should include dense-only retrieval, lexical-only retrieval, the deployed hybrid retriever, and one degraded variant without parent-node expansion. Beyond answer correctness, the paper should report retrieval Recall@k, grounded answer rate, citation support quality, p95 latency, end-to-end throughput, and degraded-mode service availability. Table 2 should hold the quality metrics, while Table 3 should summarize latency and deployment-resilience measurements under repeated route-level calls.",
            ],
            "tables": [
                {
                    "caption": "Table 2. Recommended benchmark, baselines, and metrics for the system paper.",
                    "headers": ["Dimension", "Required Design Choice", "Reporting Target"],
                    "rows": [
                        ["Evaluation set", "300-500 ESG questions across five task families", "Held-out report years and company split"],
                        ["Retriever baselines", "Dense only, lexical only, hybrid, hybrid minus parent expansion", "Quality and latency comparison"],
                        ["Answer evaluation", "Exactness, grounded answer rate, human support review", "Submission main table"],
                        ["System evaluation", "p95 latency, throughput, degraded-mode success rate", "Deployment table"],
                    ],
                }
            ],
        },
    )
    paper_01_sections[9]["heading"] = "7. Discussion and Threats to Validity"
    paper_01_sections[10]["heading"] = "8. Reproducibility and Artifact Availability"
    paper_01_sections[11]["heading"] = "9. Conclusion"

    paper_02_sections.insert(
        9,
        {
            "heading": "7. Submission-Grade Experimental Design",
            "paragraphs": [
                "The final submission should evaluate the LoRA adapter against a matched untuned base model, at least one reduced-capacity adapter setting, and one prompt-only baseline under identical prompts and decoding parameters. The validation split should be frozen before training, and all comparisons should report ROUGE-L, lexical overlap variance across seeds, answer-faithfulness scores, and expert or semi-expert preference judgments on a stratified ESG QA subset.",
                "A high-spec protocol should include three complementary views: automatic overlap metrics, groundedness-oriented factual evaluation, and structured error analysis. The recommended analysis sheet separates numeric extraction errors, unsupported policy claims, chronology mistakes, and template overfitting. Table 2 should report the main metric matrix, while Table 3 should summarize faithfulness and failure taxonomy so the paper can support stronger claims without changing its contribution boundary.",
            ],
            "tables": [
                {
                    "caption": "Table 2. Recommended final comparison matrix for the LoRA adaptation paper.",
                    "headers": ["System", "ROUGE-L", "Faithfulness", "Expert Preference", "Status"],
                    "rows": [
                        ["Untuned Qwen2.5-7B-Instruct", "[run and fill]", "[run and fill]", "[run and fill]", "Required baseline"],
                        ["LoRA r=16 alpha=32", "[repository row exists]", "[run and fill]", "[run and fill]", "Main system"],
                        ["Reduced LoRA capacity", "[run and fill]", "[run and fill]", "[run and fill]", "Ablation"],
                        ["Prompt-only adaptation", "[run and fill]", "[run and fill]", "[run and fill]", "Prompt baseline"],
                    ],
                }
            ],
        },
    )
    paper_02_sections[10]["heading"] = "8. Threats to Validity"
    paper_02_sections[11]["heading"] = "9. Reproducibility and Artifact Availability"
    paper_02_sections[12]["heading"] = "10. Conclusion"

    paper_03_sections.insert(
        8,
        {
            "heading": "6. Submission-Grade Experimental Design",
            "paragraphs": [
                "The orchestration paper should be evaluated on a fixed ESG QA task suite that stresses routing decisions, retrieval dependency, verification failures, and runtime degradation. A practical submission set would include at least 150 to 250 prompts across direct QA, comparative analysis, multi-hop evidence synthesis, and low-support failure cases. The main comparisons should remove the verifier, remove retry, remove query rewrite, and replace the graph workflow with a single-pass chain while preserving the same retriever and LLM backend.",
                "The recommended metrics are answer support quality, verifier-triggered recovery rate, average retries, fallback success rate, trace completeness, node-level latency, and total end-to-end latency. Table 2 should hold the ablation comparisons, and Figure 2 should visualize node-level trace flow and retry outcomes across audited case examples. This protocol turns the current method paper into a testable orchestration study rather than a descriptive architecture note.",
            ],
            "tables": [
                {
                    "caption": "Table 2. Recommended ablation matrix for the orchestration paper.",
                    "headers": ["Workflow Variant", "Support Quality", "Recovery Rate", "Fallback Success", "Latency Cost"],
                    "rows": [
                        ["Full router-retriever-analyst-verifier workflow", "[fill after run]", "[fill]", "[fill]", "[fill]"],
                        ["Without verifier", "[fill after run]", "N/A", "[fill]", "[fill]"],
                        ["Without retry", "[fill after run]", "[fill]", "[fill]", "[fill]"],
                        ["Without query rewrite", "[fill after run]", "[fill]", "[fill]", "[fill]"],
                        ["Single-pass chain baseline", "[fill after run]", "N/A", "N/A", "[fill]"],
                    ],
                }
            ],
        },
    )
    paper_03_sections[9]["heading"] = "7. Discussion and Threats to Validity"
    paper_03_sections[10]["heading"] = "8. Artifact Availability"
    paper_03_sections[11]["heading"] = "9. Conclusion"

    experiment_01_sections = [
        {"heading": "1. 实验目标", "paragraphs": ["目标是把论文01从系统说明稿提升到具备外部评测、运行剖面和部署韧性的工程系统论文。实验必须同时回答三个问题：混合检索是否优于单路检索、回答是否真正 grounded、系统在生产式调用下是否稳定。"]},
        {"heading": "2. 数据与任务设计", "bullets": ["构建 300-500 条 ESG 问题，覆盖指标抽取、政策问答、年份对比、争议追踪、跨文档综合五类任务。", "按公司和年份切分评测集，避免同一公司相邻年份报告泄漏。", "每题保留 gold evidence span、可接受答案和任务标签。"]},
        {"heading": "3. 对照组与变量", "bullets": ["对照组一：dense-only retrieval。", "对照组二：BM25-only retrieval。", "对照组三：当前 hybrid retrieval。", "消融组：hybrid minus parent expansion、hybrid minus quality filter。"]},
        {"heading": "4. 指标与统计", "bullets": ["质量指标：Recall@k、MRR、answer exactness、grounded answer rate、human support score。", "系统指标：p95 latency、throughput、health-check success rate、degraded-mode success rate。", "统计：paired bootstrap for grounded answer rate and latency confidence interval。"]},
        {"heading": "5. 执行步骤", "bullets": ["先跑接口与检索 smoke tests，确认基线系统稳定。", "冻结评测问题集与 gold evidence。", "分别对四个检索配置批量调用系统并保存 JSON 输出。", "人工复核不少于 100 条样本，标记支持、部分支持、无支持。", "汇总表 2/表 3 与延迟剖面图。"]},
        {"heading": "6. 结果回填与交付", "bullets": ["回填主表：检索质量、回答 groundedness、人工支持率。", "回填系统表：延迟、吞吐、降级可用性。", "保存错误案例图和 deployment degradation 案例。"]},
    ]

    experiment_02_sections = [
        {"heading": "1. 实验目标", "paragraphs": ["目标是把论文02提升成可投稿的领域适配实验论文。实验要明确回答：LoRA 是否优于 untuned base model，性能来自适配还是提示格式，回答是否 faithful。"]},
        {"heading": "2. 数据与切分", "bullets": ["冻结 train/val 切分，记录样本来源、模板来源和领域标签。", "额外抽取 150-200 条高价值 ESG QA 样本做人审与事实核验。", "按问题类型建立 error taxonomy：数值抽取、政策解释、因果叙述、时间线。"]},
        {"heading": "3. 对照组与消融", "bullets": ["主模型：当前 LoRA r=16 alpha=32。", "对照一：untuned Qwen2.5-7B-Instruct。", "对照二：prompt-only adaptation。", "消融：低秩更小的 LoRA、无 chat template masking 的训练版本。"]},
        {"heading": "4. 指标与统计", "bullets": ["自动指标：ROUGE-L、BERTScore 或语义相似度。", "真实性指标：faithfulness review、evidence support rate。", "偏好评测：双盲 pairwise preference。", "统计：seed-level mean/std，paired significance for main rows。"]},
        {"heading": "5. 执行步骤", "bullets": ["先重跑 evaluate_model，锁定当前 LoRA row。", "在完全相同 prompts 上评估 untuned base model。", "按 3 seeds 重训主配置与低秩消融。", "对人审样本导出预测、reference、source evidence，执行双盲核验。", "汇总主表、faithfulness 表和错误案例表。"]},
        {"heading": "6. 结果回填与交付", "bullets": ["表 1 回填 LoRA / base / prompt baseline / low-rank ablation。", "输出错误案例表与 4 类 failure taxonomy。", "补数据卡与泄漏风险说明。"]},
    ]

    experiment_03_sections = [
        {"heading": "1. 实验目标", "paragraphs": ["目标是把论文03变成严格的 orchestration ablation paper。实验要回答 verifier、retry、query rewrite 和 fallback 是否真的带来质量收益或韧性收益。"]},
        {"heading": "2. 任务集设计", "bullets": ["构建 150-250 条 ESG QA 与分析任务，覆盖 direct QA、comparative analysis、multi-hop synthesis、low-support failure cases。", "每条任务保存 gold support evidence、task label 和是否需要 analysis stage。", "另外抽取 30-50 条 dependency degradation 样本专测 fallback。"]},
        {"heading": "3. 对照与消融", "bullets": ["Full workflow。", "Without verifier。", "Without retry。", "Without query rewrite。", "Single-pass chain baseline。"]},
        {"heading": "4. 指标与日志", "bullets": ["质量：support quality、human acceptability、grounded answer rate。", "控制指标：retry count、recovery rate、fallback success rate、trace completeness。", "性能：node latency、total latency、token cost。"]},
        {"heading": "5. 执行步骤", "bullets": ["固定 retriever 和 LLM backend。", "逐个 workflow 版本批量运行任务集并保存 trace。", "汇总 verifier 触发次数与 recovery 成功率。", "对 10-20 个典型案例做人审分析，展示 retry 或 verifier 的真实价值。", "生成 ablation 表、trace 图和 failure taxonomy。"]},
        {"heading": "6. 结果回填与交付", "bullets": ["表 2 回填各 ablation 质量、恢复率、延迟代价。", "图 2 回填节点级 trace 可视化。", "补 dependency degradation 与 fallback 专项小节。"]},
    ]

    paper_04_sections = [
        {
            "heading": "Abstract",
            "paragraphs": [
                "This paper develops a submission-ready experimental framework for OCR-free visual retrieval over ESG reports using ColPali-style late interaction. The study is motivated by a persistent weakness of text-only retrieval pipelines: charts, complex tables, scanned pages, and layout-heavy disclosures often lose evidence during OCR or text chunking. We therefore formulate ESG report retrieval as a page-level vision-language retrieval problem and specify a benchmark that compares OCR-free visual retrieval against text-only and OCR-plus-text baselines under company-level splits. The current repository does not yet contain a native ColPali implementation, so this version of the manuscript fixes the contribution boundary, benchmark design, annotation protocol, baselines, metrics, error taxonomy, reviewer-facing qualitative analysis, and ablation plan while leaving only the final experimental values to be inserted after implementation. The resulting draft is intended to be publication-ready in structure, related work, method definition, evaluation protocol, reporting logic, submission framing, and reviewer guidance before the numerical results are produced."
            ],
        },
        {"heading": "Keywords", "paragraphs": ["ColPali, OCR-free retrieval, ESG reports, document vision-language models, PDF retrieval"]},
        {
            "heading": "Highlights",
            "bullets": [
                "The paper reframes ESG report retrieval as a page-level OCR-free visual retrieval problem rather than a text-only RAG subtask.",
                "A high-spec benchmark is defined against text-only and OCR-plus-text baselines with company-level splits, page-level relevance, and answer-grounding review.",
                "The manuscript is written to near-submission quality before implementation so that only the final numerical results remain to be inserted.",
            ],
        },
        {
            "heading": "1. Introduction",
            "paragraphs": [
                "ESG reports are visually dense documents in which evidence may live inside charts, tables, scanned pages, or layout-dependent fragments that are not preserved cleanly by OCR-first pipelines. This creates a gap between text retrieval performance on clean paragraphs and evidence retrieval performance on real-world sustainability reports. OCR-free document retrieval is therefore attractive because it can preserve page layout and visual semantics instead of assuming that page content can be losslessly flattened into text [1], [2], [6], [7].",
                "This paper studies that problem in the ESG setting. The central question is whether ColPali-style visual late interaction can retrieve evidence-bearing pages more reliably than text-only or OCR-plus-text retrieval when the target information is visually grounded. The present version is experiment-ready rather than results-complete: it fixes the method boundary, benchmark protocol, baselines, and analysis plan so that final values can be inserted after the visual retriever is implemented.",
            ],
        },
        {
            "heading": "2. Related Work",
            "paragraphs": [
                "Late-interaction retrieval has become a strong paradigm for efficient yet expressive ranking. ColBERT showed that token-level late interaction can preserve fine-grained matching signals while staying tractable for retrieval workloads [3]. ColPali transfers that idea into the document-vision setting by representing pages with a compact visual-language retriever rather than relying on OCR-normalized text alone [1].",
                "Document understanding research provides the second pillar for this problem. PaliGemma offers a compact transfer-capable vision-language backbone [2], while Donut demonstrated that OCR-free document understanding can outperform OCR-dependent approaches on noisy visual inputs [6]. These models motivate a retrieval setting in which page images are first-class evidence objects rather than intermediate containers for OCR text only.",
                "Benchmark design is equally important. BEIR established the value of heterogeneous retrieval evaluation under realistic corpus diversity [5], and BM25 remains a meaningful lexical baseline when the evaluation target is flattened text [4]. For visually grounded document QA, datasets such as DocVQA, InfographicVQA, and ChartQA reveal that answer-bearing evidence is often inseparable from layout or chart structure [7]-[9].",
                "Multi-page document reasoning adds another challenge. Hierarchical multimodal transformers and later self-attention scoring methods for multi-page DocVQA show that page selection, cross-page evidence aggregation, and layout-aware ranking remain active problems [10], [11]. ScreenAI further indicates that UI-like and infographic-style layouts require visual grounding beyond plain OCR [12]. Our work draws from these literatures but narrows the focus to ESG report retrieval rather than full end-to-end document question answering.",
            ],
        },
        {
            "heading": "3. Problem Setting and Contribution Boundary",
            "paragraphs": [
                "The target task is page-level evidence retrieval for ESG reports. Given a user query, the system must rank report pages that contain the most relevant supporting evidence, even when the evidence is primarily visual or layout-dependent. Relevance is therefore defined at the page level first, with optional span- or region-level annotation for later analysis.",
                "The contribution boundary is deliberately narrow. The paper does not claim a new ESG QA model, a new OCR method, or a new general document benchmark. Instead, it contributes a high-spec experimental framework for testing whether OCR-free ColPali-style retrieval is a better evidence retriever for ESG reports than text-only or OCR-plus-text pipelines under the same corpus and split protocol.",
            ],
        },
        {
            "heading": "4. OCR-Free Visual Retrieval Architecture",
            "paragraphs": [
                "The proposed architecture renders each ESG report page into an image, encodes it with a ColPali-style visual retriever, and indexes page embeddings for retrieval-time late interaction [1]-[3]. Query encoding remains textual, but page relevance is computed directly in the joint visual-language space rather than after OCR flattening. This preserves chart legends, table structure, and page layout as part of the retrieval signal.",
                "The comparison baselines are text-only retrieval and OCR-plus-text retrieval. Text-only retrieval uses flattened chunks ranked by lexical or text-embedding methods [4], [5]. OCR-plus-text retrieval first runs OCR and then indexes the recognized content. The proposed method, by contrast, operates on page images and optional page metadata. Figure 1 placeholder. OCR-free visual retrieval pipeline from PDF rendering to page indexing, late interaction, and evidence review. Figure 2 placeholder. Benchmark comparison between text-only, OCR-plus-text, and ColPali-style page retrieval.",
            ],
        },
        {
            "heading": "5. Dataset Construction and Annotation Protocol",
            "paragraphs": [
                "The benchmark should include at least four strata: chart-centric pages, dense tables, scanned or low-quality pages, and complex layout pages combining narrative and graphics. Queries must be written against evidence that is visibly present on the page rather than recoverable only from OCR text. Company-level splits are recommended so that pages from the same issuer do not leak between train, validation, and test partitions.",
                "Each item should include the user query, the relevant page identifier, document metadata, a page-level relevance label, and an answer-support note. A second annotation layer can optionally mark gold visual regions for qualitative analysis. Table 1 summarizes the benchmark structure that should be frozen before implementation results are produced.",
            ],
            "tables": [
                {
                    "caption": "Table 1. Recommended ESG visual retrieval benchmark design.",
                    "headers": ["Stratum", "Minimum Count", "Primary Failure Mode", "Label Type"],
                    "rows": [
                        ["Charts and figures", "100+", "Legend or axis lost in OCR", "Page-level relevance + support note"],
                        ["Dense tables", "100+", "Cell structure collapse", "Page-level relevance + support note"],
                        ["Scanned pages", "80+", "Recognition noise", "Page-level relevance + support note"],
                        ["Complex layouts", "80+", "Reading-order ambiguity", "Page-level relevance + support note"],
                    ],
                }
            ],
        },
        {
            "heading": "6. Experimental Design and Baselines",
            "paragraphs": [
                "The main baselines should be text-only retrieval, OCR-plus-text retrieval, and the proposed OCR-free visual retrieval system. A stronger submission should additionally test a visual-only ablation without metadata, a hybrid visual-plus-text fusion condition, and one reduced-resolution ablation to measure the cost of page rendering choices. Metrics should include Recall@k, MRR, exact page hit rate, grounded answer rate on a reviewed subset, average latency, and GPU memory use.",
                "Table 2 defines the core comparison grid, and Table 3 fixes the reporting standard for both quality and systems metrics. This design keeps the paper honest: if visual retrieval wins, the improvement can be localized to page-level evidence access; if it does not, the benchmark still explains where OCR-free retrieval helps and where it does not.",
            ],
            "tables": [
                {
                    "caption": "Table 2. Planned baselines and ablations for the ColPali paper.",
                    "headers": ["System", "Input Modality", "Expected Role", "Status"],
                    "rows": [
                        ["Text-only retrieval", "Flattened text chunks", "Baseline", "To run"],
                        ["OCR-plus-text retrieval", "OCR text", "Baseline", "To run"],
                        ["ColPali-style visual retrieval", "Rendered page images", "Main method", "To implement and run"],
                        ["Visual retrieval without metadata", "Rendered page images", "Ablation", "To run"],
                        ["Visual-text fusion", "Images + text metadata", "Ablation", "To run"],
                    ],
                },
                {
                    "caption": "Table 3. Required metric and resource reporting for the visual retrieval paper.",
                    "headers": ["Metric Family", "Required Metric", "Purpose"],
                    "rows": [
                        ["Retrieval quality", "Recall@k, MRR, exact page hit rate", "Primary evaluation"],
                        ["Answer support", "Grounded answer rate on reviewed subset", "User-facing evidence quality"],
                        ["Systems cost", "Average latency, p95 latency, GPU memory", "Practical feasibility"],
                        ["Qualitative analysis", "Case-study pages with failure labels", "Error interpretation"],
                    ],
                },
            ],
        },
        {
            "heading": "7. Planned Results and Analysis Protocol",
            "paragraphs": [
                "The results section will compare OCR-free visual retrieval against the two text-centric baselines on the frozen benchmark. The main claim should rest on page-level retrieval quality and grounded answer support, not on isolated cherry-picked visual examples. Final numerical values will be inserted only after the visual retriever, benchmark annotations, and batch evaluation scripts are completed.",
                "Qualitative analysis should include at least three positive cases and three failure cases. Recommended failure labels are chart semantic miss, table structure miss, scan noise, and cross-page confusion. Figure 3 placeholder. Qualitative page-level retrieval cases comparing OCR-only errors against visual retrieval behavior.",
            ],
        },
        {
            "heading": "8. Threats to Validity and Implementation Roadmap",
            "paragraphs": [
                "The largest current threat is implementation absence: the repository does not yet include a native ColPali retriever, so the manuscript cannot claim completed empirical superiority. A second threat is benchmark construction bias, especially if query writers overfit to visually obvious cases. A third threat is resource sensitivity, because page rendering resolution and GPU memory constraints can materially affect practical performance [1], [2], [6].",
                "The implementation roadmap is therefore part of the study design. The required steps are page rendering, visual encoder integration, page-level indexing, batch evaluation, and qualitative review tooling. Once those components are complete, the paper can be finalized by populating the fixed tables and figures without rewriting the method, related work, or discussion structure.",
            ],
        },
        {
            "heading": "9. Conclusion",
            "paragraphs": [
                "This paper defines a publication-ready experimental framework for testing OCR-free visual retrieval on ESG reports. Its value lies in turning a broad intuition about visual evidence loss into a concrete benchmark, baseline suite, and analysis protocol. After implementation and runs are complete, the manuscript can mature into a full visual retrieval paper without changing its contribution boundary.",
            ],
        },
        {
            "heading": "Author Contributions",
            "paragraphs": [
                "The author contributed to conceptualization, benchmark design, method specification, experimental planning, and writing of the manuscript."
            ],
        },
        {
            "heading": "Data Availability Statement",
            "paragraphs": [
                "The manuscript-generation assets, current repository scan results, and supporting text-retrieval baseline components are available within the project repository. The native visual retrieval implementation and benchmark annotations are not yet complete and must be added before final experimental values can be released. External ESG reports remain subject to their original access conditions and licensing terms."
            ],
        },
        {
            "heading": "Conflict of Interest",
            "paragraphs": [
                "The author declares no conflict of interest."
            ],
        },
        {
            "heading": "References",
            "numbered_items": paper_04_references,
        },
    ]

    manual_04_sections = [
        {"heading": "1. 论文定位", "paragraphs": ["论文04现在从蓝图升级为实验待完成的正式视觉检索论文。主线不是整个 ESG 平台，而是 ColPali 式 OCR-free visual retrieval 是否优于 text-only 与 OCR+text baselines。"]},
        {"heading": "2. 现有可直接写入结果", "bullets": ["可直接写入：研究问题、benchmark 结构、page-level 标注协议、company-level split、baseline 与 ablation 设计。", "可直接写入：当前仓库没有 native ColPali implementation，这一点必须诚实保留。", "可直接写入：论文的表 1、表 2、表 3、图 1、图 2、图 3 都已固定用途。"]},
        {"heading": "3. 投稿前必须补的实验", "bullets": ["补页面渲染、视觉 embedding、page-level index、batch retrieval evaluator。", "补 text-only、OCR+text、visual retrieval 三组主对照。", "补 chart/table/scan/layout 四类页面的标注集与 reviewed subset。", "补 Recall@k、MRR、exact page hit、grounded answer rate、latency、GPU memory。"]},
        {"heading": "4. 实验步骤", "bullets": ["先整理 ESG 报告 PDF，按公司和年份建立 page-level corpus。", "对每页渲染图像并保存 page metadata。", "搭建 text-only 与 OCR+text baseline，导出统一 query set。", "接入 ColPali-style encoder，建立 page embedding 和 late interaction 检索。", "批量评估并抽样人工复核支持性。"]},
        {"heading": "5. 结果回填与交付", "bullets": ["表 1 回填 benchmark 统计与 split 说明。", "表 2 回填 baseline 与 ablation 的主结果。", "表 3 回填 latency、memory 与 qualitative failure taxonomy。", "图 3 回填正反案例各至少 3 组。"]},
    ]

    experiment_04_sections = [
        {"heading": "1. 实验目标", "paragraphs": ["目标是以高规格协议验证 ColPali 式视觉检索是否在 ESG 报告场景中优于 text-only 与 OCR+text 检索。实验必须覆盖图表、表格、扫描页和复杂布局四类页面，并同时报告质量与资源成本。"]},
        {"heading": "2. 数据与标注", "bullets": ["按公司维度切分 train/val/test，避免同公司不同年份泄漏。", "至少构建 360 条 query，其中四类页面每类不少于 80-100 条。", "每条 query 保留 relevant page、support note、failure category。"]},
        {"heading": "3. 对照与消融", "bullets": ["主对照：text-only retrieval。", "主对照：OCR+text retrieval。", "主方法：ColPali-style visual retrieval。", "消融：visual without metadata、visual-text fusion、reduced-resolution rendering。"]},
        {"heading": "4. 指标与统计", "bullets": ["主指标：Recall@1/5/10、MRR、exact page hit rate。", "支持性指标：grounded answer rate。", "系统指标：average latency、p95 latency、GPU memory。", "统计：按页面 strata 分层汇报，并给出 bootstrap interval。"]},
        {"heading": "5. 执行步骤", "bullets": ["冻结 query set 和标注协议。", "跑 text-only baseline。", "跑 OCR+text baseline。", "实现并运行 ColPali-style visual retrieval。", "汇总主表、分层表和 qualitative cases。"]},
        {"heading": "6. 结果回填与交付", "bullets": ["主表：三种主系统的检索质量。", "分层表：chart/table/scan/layout 四类页面结果。", "资源表：latency 与 memory。", "案例图：至少 6 个正反案例。"]},
    ]

    paper_05_sections = [
        {
            "heading": "Abstract",
            "paragraphs": [
                "This study investigates whether ESG-derived signals can improve risk-adjusted performance in a reinforcement learning equity strategy when the signal is aligned to disclosure timing and routed through regime-aware control. The paper upgrades a historical prototype into a finance-first experimental design centered on soft actor-critic trading, annual-report ESG features, and a routing layer that adjusts policy emphasis under changing macro-financial conditions. The formal protocol uses a fixed 20-stock universe, train-validation-test splits of 2022-2023, 2024, and 2025, and a frozen 96-run matrix spanning baselines, ESG variants, formulas, ablations, and random seeds. Current repository assets already define data lineage, time-alignment rules, preflight gates, and contribution-report generation for submission-grade experimentation. The manuscript therefore fixes the contribution boundary, statistical plan, and result-table structure before final numbers are inserted. The central question is whether ESG information improves Sharpe ratio, drawdown control, and stability, or whether low-frequency report signals mainly expose the limits of annual ESG disclosures for short-horizon trading."
            ],
        },
        {
            "heading": "Keywords",
            "paragraphs": [
                "ESG investing, reinforcement learning, regime-aware policy routing, quantitative equity trading, risk-adjusted return"
            ],
        },
        {
            "heading": "Highlights",
            "bullets": [
                "The paper evaluates ESG-derived signals in a reinforcement learning equity strategy under a fixed 2022-2025 protocol and a frozen 20-stock universe.",
                "A regime-aware routing layer is assessed against no-ESG, no-routing, and reward-observation ablations within a 96-run experimental matrix.",
                "The manuscript is structured so that final run-matrix values can be inserted without changing the contribution boundary, statistical tests, or submission narrative.",
            ],
        },
        {
            "heading": "1. Introduction",
            "paragraphs": [
                "Deep reinforcement learning has become a standard framework for sequential trading decisions, but the empirical literature still shows that strategy quality depends heavily on reward design, market frictions, and feature discipline rather than on the learning algorithm alone [1]-[6]. In parallel, the asset-pricing and sustainable-finance literatures have shown that ESG characteristics can matter for valuation, capital allocation, and downside protection, although the sign and size of the effect vary across data vendors and portfolio construction assumptions [11]-[14].",
                "This paper studies that intersection. It asks whether annual-report ESG signals, converted into time-aligned daily features, can improve risk-adjusted return in a soft actor-critic trading system and whether regime-aware policy routing adds stability beyond a single-policy baseline [2]-[5]. The goal is not to repackage the ESG scoring pipeline or the broader AI platform as a main contribution. Instead, the manuscript is written as a finance-first experimental paper whose claims live at the level of trading performance, ablation evidence, and robustness under a pre-registered protocol.",
            ],
        },
        {
            "heading": "2. Related Work",
            "paragraphs": [
                "The first relevant literature concerns reinforcement learning for trading. Early work on direct reinforcement highlighted the importance of risk-sensitive reward shaping in financial markets [6], while deep reinforcement learning later provided scalable function approximation for sequential decision making [1]. Recent domain toolkits and expert-system studies have adapted these ideas into reproducible algorithmic-trading workflows centered on soft actor-critic and related off-policy methods [2]-[5].",
                "A second literature concerns regime sensitivity and policy specialization. Markov-switching models remain a canonical approach to nonstationary market states [7], and mixture-of-experts research provides a principled basis for routing inputs toward different local decision rules when data-generating conditions shift [8]. In practical trading pipelines, lightweight tabular regime detectors and constrained hyperparameter search are natural complements to this idea because they preserve interpretability and protocol discipline under limited sample sizes [9], [10].",
                "The third literature concerns ESG and financial performance. Meta-analytic evidence suggests that ESG information can affect return and risk in economically meaningful ways, but the mechanism is rarely universal across sectors and horizons [11]. Multi-factor asset-pricing work provides a baseline for separating ESG effects from generic factor exposures [12], while portfolio studies on the ESG-efficient frontier suggest that risk-adjusted improvement can occur even when headline return gains are modest [13].",
                "A final literature addresses ESG signal construction and disagreement. Cross-provider divergence is substantial, which means that any finance paper using ESG information must document score construction and timing rather than treating ESG as a plug-and-play feature [14]. In this project, report-level evidence is converted into a house score through a documented evidence chain, conceptually related to retrieval-augmented evidence assembly but used here only as an upstream signal-generation step rather than as the paper's primary methodological contribution [15].",
            ],
        },
        {
            "heading": "3. Problem Formulation and Contribution Boundary",
            "paragraphs": [
                "The target problem is daily allocation over a fixed 20-stock universe. At each trading day, the agent observes market features, macro-regime variables, and ESG-derived variables that become available only after a report-specific effective date. The learning objective is to maximize risk-adjusted portfolio performance after transaction costs while avoiding leakage from future disclosures.",
                "The manuscript makes three bounded contributions. First, it formalizes a paper-grade protocol in which ESG-derived features are time aligned, neutral when missing, and evaluated under a frozen train-validation-test split. Second, it introduces regime-aware policy routing as an experimental control-plane addition on top of a soft actor-critic trading core rather than as a separate open-ended agent system. Third, it pre-specifies ablation, robustness, and statistical-significance tables so that final claims depend on completed runs rather than on post hoc narrative selection.",
                "The contribution boundary deliberately excludes the broader ESG question-answering platform, orchestration layer, vector-store implementation, lexical retrieval stack, and deployment plumbing. Those assets may explain where an ESG score comes from, but the present paper is evaluated only on quantitative trading outcomes and protocol-level financial evidence.",
            ],
        },
        {
            "heading": "4. ESG Signal Construction and Time Alignment",
            "paragraphs": [
                "The upstream ESG pipeline produces daily features from report-level evidence. The formal trading datasets include house_score_v2, house_score_v2_1, esg_delta, esg_delta_v2_1, esg_confidence, esg_staleness_days, esg_missing_flag, and sector_relative_esg. The main manuscript centers on house_score_v2_1 because the calibrated variant preserves the same evidence-only foundation while adding sector-year normalization inside the score-construction stage.",
                "Time alignment follows a strict disclosure rule: an ESG value becomes tradable only on published_date + 1 trading day. Before the first effective report for a company, the daily row uses neutral missing handling with house_score_v2_1 = 50, esg_confidence = 0, and esg_missing_flag = 1. This design prevents time leakage and keeps missing reports neutral rather than treating them as implicitly poor ESG performance.",
                "The main experimental table uses sample_full_2022_2025 and formula_v2_1. The robustness layer then compares formula_v2 against formula_v2_1 and contrasts full_2022_2025 with post_esg_effective, where the latter removes the pre-effective neutral window and therefore stresses the question of whether annual ESG signals still add value once timing frictions are stripped away.",
                "Figure 1 placeholder. Data lineage from ESG reports to house scores, effective-date alignment, daily merged datasets, and paper-run artifacts.",
            ],
        },
        {
            "heading": "5. Trading Environment and Reward Design",
            "paragraphs": [
                "The trading environment is built around daily bars, transaction costs, and risk-sensitive reward shaping. The soft actor-critic core operates on market prices, ESG features, and regime inputs while tracking annual return, Sharpe ratio, Sortino ratio, Calmar ratio, turnover, win rate, and maximum drawdown as paper-level outputs [2]-[6]. The no-ESG baseline keeps the same market environment so that any incremental effect can be attributed to ESG information rather than to unrelated infrastructure changes.",
                "Reward design separates portfolio-level return from the incremental ESG contribution. The repository already supports no-ESG, ESG-enabled, no-ESG-observation, no-ESG-reward, and no-regime configurations inside the formal paper-run matrix. This separation is essential because a positive headline Sharpe change can come either from better state information or from a reward bonus that changes behavior even when the state is unchanged.",
            ],
        },
        {
            "heading": "6. Regime-Aware Policy Routing",
            "paragraphs": [
                "Regime-aware policy routing is introduced as a constrained control layer rather than as a free-form multi-agent architecture. The route decision uses macro-financial context, recent volatility conditions, and ESG-confidence information to adjust how strongly the policy relies on ESG-sensitive actions versus market-only behavior. The design is conceptually related to mixture-of-experts routing under nonstationary conditions [7], [8], while practical implementation choices remain intentionally lightweight so that the final paper can isolate routing value from modeling complexity [9].",
                "The ablation group 6c_no_regime removes regime-sensitive routing while preserving the rest of the environment. This makes the routing contribution falsifiable. If the final results show that routing improves stability or downside control, the paper can attribute that gain to policy specialization under shifting regimes. If the final results do not show a routing benefit, the same design still provides a clean negative result on the limited value of regime-aware switching under annual-report ESG features.",
                "Figure 2 placeholder. Regime-aware policy routing from macro state and ESG confidence to policy emphasis, ablation controls, and execution outputs.",
            ],
        },
        {
            "heading": "7. Experimental Protocol",
            "paragraphs": [
                "The formal protocol is fixed by the runbook and hardening tests. The stock universe contains 20 equities across technology, financials, energy, consumer, and healthcare sectors. Time splits are frozen as train 2022-2023, validation 2024, and test 2025, and the 2025 test set is excluded from Optuna-style model selection or manual tuning [10].",
                "The complete paper-run matrix contains 2 samples x 2 formulas x 8 groups x 3 seeds = 96 expected runs. Samples are full_2022_2025 and post_esg_effective. Formula modes are v2 and v2_1. Groups are B1_buyhold, B2_macd, B3_sac_noesg, B4_sac_esg, OURS_full, 6a_no_esg_obs, 6b_no_esg_reward, and 6c_no_regime. Every expected run must produce metrics.json, equity_curve.csv, run_status.json, run.log, and the corresponding group log before the matrix is considered complete.",
                "The primary submission table will report formula_v2_1 + sample_full_2022_2025 as the main setting. Formula_v2 and sample_post_esg_effective are retained for robustness analysis rather than headline reporting. Statistical comparison will be based on paired daily-return tests, bootstrap Sharpe intervals, seed-stability summaries, and the contribution report generated from the completed matrix.",
            ],
        },
        {
            "heading": "8. Main Results",
            "paragraphs": [
                "The main results section is reserved for the completed formula_v2_1 + sample_full_2022_2025 matrix. The narrative boundary is already fixed: the paper will compare OURS_full against B1_buyhold, B2_macd, B3_sac_noesg, and B4_sac_esg on risk-adjusted metrics and drawdown-aware behavior. Final cells are populated only after the expected-run manifest verifies completeness for all required seeds and groups.",
                "Table 1 is the headline paper table. It should report annual return, Sharpe ratio, Sortino ratio, Calmar ratio, maximum drawdown, turnover, and win rate. Figure 3 placeholder. Equity-curve comparison for B3_sac_noesg, B4_sac_esg, and OURS_full under formula_v2_1 + sample_full_2022_2025.",
            ],
            "tables": [
                {
                    "caption": "Table 1. Main test-set results under formula_v2_1 and sample_full_2022_2025.",
                    "headers": ["Group", "Annual Return", "Sharpe Ratio", "Sortino Ratio", "Calmar Ratio", "Max Drawdown", "Turnover", "Win Rate"],
                    "rows": [
                        ["B1_buyhold", "[fill after 96-run completion]", "[fill]", "[fill]", "[fill]", "[fill]", "[fill]", "[fill]"],
                        ["B2_macd", "[fill after 96-run completion]", "[fill]", "[fill]", "[fill]", "[fill]", "[fill]", "[fill]"],
                        ["B3_sac_noesg", "[fill after 96-run completion]", "[fill]", "[fill]", "[fill]", "[fill]", "[fill]", "[fill]"],
                        ["B4_sac_esg", "[fill after 96-run completion]", "[fill]", "[fill]", "[fill]", "[fill]", "[fill]", "[fill]"],
                        ["OURS_full", "[fill after 96-run completion]", "[fill]", "[fill]", "[fill]", "[fill]", "[fill]", "[fill]"],
                    ],
                },
                {
                    "caption": "Table 2. Statistical significance and interval estimates for the main comparisons.",
                    "headers": ["Comparison", "Paired Return Test", "Bootstrap Sharpe Interval", "Effect Direction", "Interpretation"],
                    "rows": [
                        ["OURS_full vs B3_sac_noesg", "[fill]", "[fill]", "[fill]", "ESG incremental value"],
                        ["OURS_full vs B4_sac_esg", "[fill]", "[fill]", "[fill]", "Routing incremental value"],
                        ["B4_sac_esg vs B3_sac_noesg", "[fill]", "[fill]", "[fill]", "Pure ESG feature effect"],
                    ],
                },
            ],
        },
        {
            "heading": "9. Ablation and Robustness Analysis",
            "paragraphs": [
                "Ablation analysis should isolate where any observed gain comes from. Group 6a_no_esg_obs tests whether ESG information matters mainly through the observation vector, 6b_no_esg_reward tests whether the ESG-linked reward term is doing the work, and 6c_no_regime tests whether routing contributes beyond a single-policy design. These ablations must be read alongside B3_sac_noesg, B4_sac_esg, and OURS_full rather than in isolation.",
                "Robustness analysis then checks whether the headline finding survives protocol variation. The first robustness slice swaps formula_v2_1 for formula_v2. The second swaps full_2022_2025 for post_esg_effective. The third reports seed stability and contribution-report summaries. Together, these views test whether the main result is driven by calibration choices, timing windows, or random initialization rather than by a stable ESG contribution.",
            ],
            "tables": [
                {
                    "caption": "Table 3. Ablation analysis for ESG observation, ESG reward, and regime routing.",
                    "headers": ["Group", "Key Removal", "Sharpe Change vs OURS_full", "Drawdown Change vs OURS_full", "Interpretation"],
                    "rows": [
                        ["6a_no_esg_obs", "Remove ESG observations", "[fill]", "[fill]", "[fill]"],
                        ["6b_no_esg_reward", "Remove ESG reward term", "[fill]", "[fill]", "[fill]"],
                        ["6c_no_regime", "Remove regime-aware routing", "[fill]", "[fill]", "[fill]"],
                    ],
                },
                {
                    "caption": "Table 4. Robustness checks across formulas, samples, and seed stability.",
                    "headers": ["Setting", "Primary Metric", "Stability Signal", "Status"],
                    "rows": [
                        ["formula_v2 + sample_full_2022_2025", "[fill]", "[fill]", "Robustness"],
                        ["formula_v2_1 + sample_post_esg_effective", "[fill]", "[fill]", "Robustness"],
                        ["Seed stability summary", "[fill]", "[fill]", "Robustness"],
                    ],
                },
            ],
        },
        {
            "heading": "10. Discussion",
            "paragraphs": [
                "The discussion is intentionally dual-path. If the final results show that OURS_full improves Sharpe ratio, reduces drawdown, or stabilizes cross-seed performance, the paper can interpret ESG-derived information as an economically useful state variable whose contribution survives against both no-ESG and no-routing controls. Under that outcome, the strongest claim is not that ESG guarantees higher raw return, but that time-aligned ESG information improves risk-adjusted decision quality under a constrained trading protocol.",
                "If the final results do not show a statistically meaningful improvement, the paper still supports a publishable conclusion. In that case, the negative result would indicate that annual-report ESG signals are too low frequency, too delayed, or too noisy to create short-horizon alpha after realistic timing control. That finding would still matter for the literature because it separates the usefulness of ESG for long-horizon allocation from its usefulness for daily reinforcement-learning trading and grounds the interpretation in a complete ablation matrix rather than in anecdotal failure.",
                "In either outcome, the completed paper should compare the observed effect with the broader ESG-finance literature on risk adjustment, vendor disagreement, and multi-factor exposures [11]-[14]. The submission value lies in turning those debates into a controlled trading experiment with explicit timing discipline, not in claiming that ESG always dominates traditional market signals.",
            ],
        },
        {
            "heading": "11. Limitations",
            "paragraphs": [
                "The first limitation is data frequency. ESG reports are sparse relative to daily trading decisions, so even a well-aligned signal may operate more as a slow-moving preference modifier than as a rapid alpha driver. The second limitation is universe size: a fixed 20-stock set makes the experiment interpretable and feasible, but it does not establish broad cross-market generality. The third limitation is that the current design focuses on annual reports and evidence-derived house scores rather than on alternative higher-frequency ESG proxies such as news or controversy feeds.",
                "A related limitation is that the routing layer is deliberately constrained. This is a strength for causal interpretation, but it also means that the paper does not exhaust the design space of adaptive trading control. The submission should therefore present the regime-aware router as a disciplined experimental mechanism, not as a universal solution to nonstationary markets.",
            ],
        },
        {
            "heading": "12. Conclusion",
            "paragraphs": [
                "This paper reframes the existing ESG trading prototype as a submission-oriented financial experiment. The contribution is a fixed protocol for testing whether time-aligned ESG features and regime-aware policy routing improve risk-adjusted equity-trading performance under a transparent 2022-2025 design. Once the 96-run matrix is complete, the manuscript can be finalized by filling the planned tables and significance tests without changing its core claim structure.",
            ],
        },
        {
            "heading": "Author Contributions",
            "paragraphs": [
                "The author contributed to conceptualization, experimental design, ESG signal engineering, reinforcement learning workflow design, statistical analysis planning, and writing of the manuscript."
            ],
        },
        {
            "heading": "Data Availability Statement",
            "paragraphs": [
                "The code, configuration, runbook, validation tests, and manuscript-generation assets that support this study are available within the project repository. Formal paper-run artifacts are organized under the paper-run namespace, including scripts for preflight checking, expected-run verification, contribution reporting, and dataset construction. External ESG reports and source-company disclosures remain subject to their original access conditions and licensing terms."
            ],
        },
        {
            "heading": "Conflict of Interest",
            "paragraphs": [
                "The author declares no conflict of interest."
            ],
        },
        {
            "heading": "References",
            "numbered_items": paper_05_references,
        },
    ]

    manual_05_sections = [
        {
            "heading": "1. 论文定位",
            "paragraphs": [
                f"这篇论文05现在不是蓝图，而是主投稿版金融实验论文。它继承历史稿《{historical_title}》的研究问题，但统一收口到 ESWA 风格：主线只讨论 ESG 信号是否改善 risk-adjusted return，以及 regime-aware policy routing 是否带来稳定增益。",
                "写作时要坚持 finance-first experimental paper 的边界，不再把 RAG 系统、Qdrant、BM25、RRF、LangGraph、多智能体工程写成主贡献。那些内容最多只保留为 ESG score 来源背景。",
            ],
        },
        {
            "heading": "2. 现有可直接写入结果",
            "bullets": [
                "可直接写入正式协议：20 只股票固定池，Train 2022-2023 / Validation 2024 / Test 2025。",
                "可直接写入时间对齐规则：published_date + 1 trading day，且缺失 ESG 采用 neutral missing handling。",
                "可直接写入正式矩阵：2 samples x 2 formulas x 8 groups x 3 seeds = 96 expected runs。",
                "可直接写入主结果口径：主表默认使用 formula_v2_1 + sample_full_2022_2025，formula_v2 与 sample_post_esg_effective 放入稳健性层。",
                "可直接写入现有产物链：runbook、paper hardening test、expected-run manifest、contribution report、preflight、quant service 与 API 路由均已存在。",
            ],
        },
        {
            "heading": "3. 投稿前必须补的实验",
            "bullets": [
                "补完 96-run 主矩阵，并确认每个 expected run 都产出 metrics.json、equity_curve.csv、run_status.json、run.log 与 group log。",
                "回填 Table 1 主结果、Table 2 显著性检验、Table 3 消融、Table 4 稳健性，不要在正文里临时改贡献口径。",
                "补 paired daily-return test、bootstrap Sharpe interval、seed stability summary，以及 contribution report 的最终导出。",
                "补 equity curve 图与最少一组 negative-result interpretation 所需的汇总图，确保正负结果都能直接接入现有 Discussion。",
                "补 data lineage appendix，用于说明样本、公式、组别、种子、effective-date 对齐和缺失 ESG 处理。",
            ],
        },
        {
            "heading": "4. 建议执行命令",
            "bullets": [
                "python scripts/run_esg_rl_2022_2025_pipeline.py --run-namespace paper-run --sample full_2022_2025 --corpus-root esg_reports --build-datasets",
                "python scripts/run_esg_rl_2022_2025_pipeline.py --run-namespace paper-run --sample post_esg_effective --corpus-root esg_reports --build-datasets",
                "python scripts/quant_rl_paper_preflight.py --namespace paper-run --sample full_2022_2025 --require-cuda",
                "python scripts/quant_rl_paper_preflight.py --namespace paper-run --sample post_esg_effective --require-cuda",
                "bash scripts/autodl_run_paper_experiments.sh",
                "python scripts/quant_rl_esg_contribution_report.py --run-namespace paper-run --sample full_2022_2025 --formula-mode v2_1",
            ],
        },
        {
            "heading": "5. 结果回填清单",
            "bullets": [
                "主表：B1_buyhold、B2_macd、B3_sac_noesg、B4_sac_esg、OURS_full 的 annual return、Sharpe、Sortino、Calmar、MDD、turnover、win rate。",
                "显著性：OURS_full vs B3_sac_noesg、OURS_full vs B4_sac_esg、B4_sac_esg vs B3_sac_noesg 的 paired test 与 bootstrap interval。",
                "消融：6a_no_esg_obs、6b_no_esg_reward、6c_no_regime 相对 OURS_full 的指标差值与解释。",
                "稳健性：formula_v2、sample_post_esg_effective、seed stability、contribution report。",
                "图表：主 equity curve、ablation summary、seed stability 或 contribution 图至少各一张。",
            ],
        },
        {
            "heading": "6. 与其他论文的边界",
            "bullets": [
                "不要和 01 重复混合检索、服务部署、API contract、Railway 部署内容。",
                "不要和 02 重复 LoRA 微调、ROUGE 评测、Qwen2.5 适配内容。",
                "不要和 03 重复 typed state、retry、fallback、LangGraph orchestration 内容。",
                "05 只讨论 quant/RL protocol、ESG signal timing、regime-aware routing、ablation、robustness、statistical tests。",
            ],
        },
    ]

    experiment_05_sections = [
        {
            "heading": "1. 实验目标",
            "paragraphs": [
                "目标是把论文05推进到主投稿所需的 submission-grade finance experiment。实验必须同时回答三个问题：ESG signal 是否改善 risk-adjusted return，regime-aware routing 是否带来额外稳定增益，以及这些增益在统计上是否稳健。",
            ],
        },
        {
            "heading": "2. 数据与协议冻结",
            "bullets": [
                "固定 20-stock universe，不在主表阶段临时扩池。",
                "固定时间切分：Train 2022-2023 / Validation 2024 / Test 2025。",
                "固定时间对齐：published_date + 1 trading day，缺失 ESG 采用 neutral missing handling。",
                "固定主表口径：formula_v2_1 + sample_full_2022_2025；formula_v2 与 sample_post_esg_effective 进入稳健性层。",
            ],
        },
        {
            "heading": "3. 主矩阵与对照",
            "bullets": [
                "跑完 2 samples x 2 formulas x 8 groups x 3 seeds = 96 expected runs。",
                "主对照必须包含 B1_buyhold、B2_macd、B3_sac_noesg、B4_sac_esg、OURS_full。",
                "消融必须包含 6a_no_esg_obs、6b_no_esg_reward、6c_no_regime。",
                "稳健性必须覆盖 formula_v2 与 post_esg_effective sample。",
            ],
        },
        {
            "heading": "4. 指标与统计",
            "bullets": [
                "主指标：annual return、Sharpe、Sortino、Calmar、MDD、turnover、win rate。",
                "统计：paired daily-return test、bootstrap Sharpe interval、seed stability summary。",
                "解释层：ESG / reward / routing contribution report，以及主组与消融组的指标差值表。",
                "负结果包：若未显著 outperform，则仍需导出风险改善、稳定性、失败模式与限制解释。 ",
            ],
        },
        {
            "heading": "5. 执行步骤",
            "bullets": [
                "先运行 dataset build 与 preflight，确认 full_2022_2025 和 post_esg_effective 两个 sample 均通过检查。",
                "使用 expected-run manifest 冻结 96-run 组合，并在 AutoDL driver 中执行主矩阵。",
                "收集每个 run 的 metrics.json、equity_curve.csv、run_status.json、run.log 和 group log。",
                "生成 paired test、bootstrap Sharpe、seed stability 与 contribution report。",
                "回填主表、显著性表、消融表、稳健性表，并导出主 equity curve 与 contribution 图。",
            ],
        },
        {
            "heading": "6. 结果回填与交付",
            "bullets": [
                "Table 1 回填主结果，重点比较 B3、B4 与 OURS_full。",
                "Table 2 回填 paired daily-return test 与 bootstrap Sharpe interval。",
                "Table 3 回填 ESG / reward / routing 拆分后的 ablation 结果。",
                "Table 4 回填 formula、sample、seed stability 和 contribution report 的稳健性摘要。",
                "附录交付：data lineage、expected-run manifest、negative-result-safe output package。",
            ],
        },
    ]

    _extend_section_paragraphs(
        paper_01_sections,
        "1. Introduction",
        [
            "The research question is therefore not whether any single retrieval primitive is new, but whether a deployment-facing ESG QA system can keep retrieval quality, answer support, service contracts, and degraded-mode behavior within one jointly verifiable surface. The paper should be read as a systems study in which integration choices are justified only when they can be tested end to end rather than as a catalog of reusable components.",
        ],
    )
    _extend_section_paragraphs(
        paper_01_sections,
        "System Overview and Contribution Boundary",
        [
            "The contribution boundary is intentionally narrow. The paper does not claim a new rank fusion algorithm, a new vector index, or a new agentic workflow. Its claim is that hybrid retrieval, application wiring, and service-level validation are treated as one systems problem, so novelty must be judged by benchmarked answer support, retrieval robustness, and deployment resilience under the same executable stack.",
        ],
    )
    _extend_section_paragraphs(
        paper_01_sections,
        "Discussion and Threats to Validity",
        [
            "If the final study shows that hybrid retrieval improves retrieval recall without materially changing answer quality, the contribution should be interpreted as a systems robustness result rather than a direct QA-quality improvement. In that case the paper remains publishable by emphasizing evidence coverage, degraded-mode stability, and engineering reproducibility, while explicitly narrowing any claim about downstream answer gains.",
        ],
    )

    _extend_section_paragraphs(
        paper_02_sections,
        "Task Definition and Dataset Description",
        [
            "Leakage control is essential because ESG question answering often reuses repetitive disclosure language. The submission protocol therefore freezes train-validation boundaries by source document and reporting period, removes near-duplicate prompt templates where possible, and records question type, answer form, and document provenance so that later gains are not mistaken for template memorization.",
        ],
    )
    _extend_section_paragraphs(
        paper_02_sections,
        "Baseline and Evaluation Gap",
        [
            "The most important missing comparison is a matched untuned base model evaluated on exactly the same prompts, decoding settings, and validation split. ROUGE-L alone is insufficient for a high-factuality ESG task, so the current result should be presented as a reproducible snapshot rather than as evidence of superiority until faithfulness review, answer support assessment, and baseline-matched comparisons are added.",
        ],
    )
    _extend_section_paragraphs(
        paper_02_sections,
        "Threats to Validity",
        [
            "A second validity risk is that overlap-based metrics may reward answer style conformity rather than factual correctness. If future human review or evidence-grounding checks diverge from ROUGE-L, the paper should prioritize the groundedness-oriented findings and explicitly recast ROUGE-L as a supporting metric instead of a decision criterion.",
        ],
    )

    _extend_section_paragraphs(
        paper_03_sections,
        "Workflow Definition and Typed State",
        [
            "The workflow is motivated by the limits of single-pass prompt chains. When retrieval quality is uncertain, verification may fail, or dependencies may degrade, a monolithic chain provides little control over recovery behavior. Typed shared state is treated here as a methodological device for making routing conditions, retries, and fallback transitions explicit enough to measure rather than as a purely implementation-level convenience.",
        ],
    )
    _extend_section_paragraphs(
        paper_03_sections,
        "Control Flow, Retry, and Fallback Execution",
        [
            "Each control mechanism has a distinct role in the planned attribution study: query rewriting aims to improve evidence reachability before analysis, verifier-triggered retry addresses answer-support failures after an initial pass, and fallback execution preserves workflow continuity when the graph runtime or a dependency becomes unavailable. The final evaluation should therefore report their contributions separately instead of treating them as one bundled orchestration effect.",
        ],
    )
    _extend_section_paragraphs(
        paper_03_sections,
        "Validation Protocol and Planned Ablations",
        [
            "The ablations are not optional extensions but the core evidentiary requirement for publication. Without matched comparisons against without verifier, without retry, without query rewrite, and a single-pass chain baseline, the paper would remain a design note rather than an orchestration study.",
        ],
    )
    _extend_section_paragraphs(
        paper_03_sections,
        "Discussion and Threats to Validity",
        [
            "If the final results show stronger traceability and recovery behavior but only small answer-quality gains, the paper should narrow its contribution to controllability and runtime robustness. That framing is preferable to overstating quality gains and still fits the stated research boundary of auditable orchestration under practical service uncertainty.",
        ],
    )

    _extend_section_paragraphs(
        paper_04_sections,
        "Problem Setting and Contribution Boundary",
        [
            "The contribution should therefore be evaluated as a benchmark-and-protocol paper with a concrete method framework, not as a claim that the final visual retriever has already outperformed all baselines. This honesty is important because the experimental design itself must survive review before native ColPali-style implementation is complete.",
        ],
    )
    _extend_section_paragraphs(
        paper_04_sections,
        "Experimental Design and Baselines",
        [
            "Benchmark fairness is a first-class concern. The text-only and OCR-plus-text baselines should receive the same corpus coverage, company-level split discipline, and query set as the visual retriever, and the paper should explain that OCR-plus-text is included precisely because a strong layout-aware textual pipeline is the most credible alternative explanation for any observed gains.",
        ],
    )
    _extend_section_paragraphs(
        paper_04_sections,
        "Threats to Validity and Implementation Roadmap",
        [
            "A negative-result-safe interpretation is also necessary. If OCR-free retrieval only improves chart or scan pages while matching rather than surpassing OCR-plus-text on ordinary pages, the paper can still contribute by identifying where visual retrieval is actually needed and where text recovery remains sufficient, provided that the benchmark strata and annotation protocol are fixed before implementation.",
        ],
    )

    _extend_section_paragraphs(
        paper_05_sections,
        "Problem Formulation and Contribution Boundary",
        [
            "The core claim must separate three effects that are easy to conflate: ESG-derived observation features, ESG-aware reward shaping, and regime-aware policy routing. The paper should only claim a routing benefit when OURS_full is compared directly against B4_sac_esg and the no-regime ablation under the same sample, formula, and seed protocol.",
        ],
    )
    _extend_section_paragraphs(
        paper_05_sections,
        "Experimental Protocol",
        [
            "External-validity risk is addressed by freezing the 20-stock universe and preventing any 2025 information from entering feature design, tuning, or qualitative screening. The goal is not to claim universal market coverage, but to make one clean, auditable finance experiment in which 2025 remains a true held-out year and every design choice can be traced through the expected-run manifest and preflight checks.",
        ],
    )
    _extend_section_paragraphs(
        paper_05_sections,
        "Main Results",
        [
            "The main-results narrative should explicitly distinguish absolute return, risk-adjusted return, and stability. If ESG improves Sharpe ratio or drawdown control without maximizing raw return, the paper should present that outcome as a finance-relevant trade-off rather than as a failure to improve performance.",
        ],
    )
    _extend_section_paragraphs(
        paper_05_sections,
        "Ablation and Robustness Analysis",
        [
            "Ablation interpretation should focus on source attribution rather than leaderboard framing. In particular, if B4_sac_esg and OURS_full are close, the analysis must state whether routing primarily improves stability, recovery under regime shifts, or seed variance rather than forcing a weak average-return superiority claim.",
        ],
    )
    _extend_section_paragraphs(
        paper_05_sections,
        "Discussion",
        [
            "If the final results show statistically supported gains, the discussion should attribute them carefully across ESG signal timing, reward design, and routing control. If the final results do not show clear outperformance, the same section should pivot to what the experiment reveals about the low frequency of annual ESG disclosures, the limits of a 20-stock setting, and the conditions under which annual-report signals may still improve downside control or stability.",
        ],
    )
    _extend_section_paragraphs(
        paper_05_sections,
        "Limitations",
        [
            "The main external-validity constraint is that a fixed 20-stock universe cannot stand in for the entire equity market. The paper should therefore present its findings as a tightly controlled experimental study, with later expansion to broader universes treated as future work rather than assumed generalization.",
        ],
    )

    paper_run_state = paper_run_context["result_state"]
    main_report = paper_run_context.get("reports", {}).get("main") or {}
    robust_formula_report = paper_run_context.get("reports", {}).get("robust_formula_v2") or {}
    robust_post_report = paper_run_context.get("reports", {}).get("robust_post_effective") or {}
    progress_display = f"{paper_run_context['progress_label']} ({paper_run_context['progress_pct']:.1f}%)" if paper_run_context.get("expected_run_count") else paper_run_context["progress_label"]
    latest_metric_labels = ", ".join(paper_run_context.get("latest_metrics")[:3]) if paper_run_context.get("latest_metrics") else "none yet"

    def _group_metric(report: dict[str, Any], group: str, metric: str) -> Any:
        return (((report.get("grouped") or {}).get(group) or {}).get("metrics") or {}).get(metric, {}).get("mean")

    def _group_metric_std(report: dict[str, Any], group: str, metric: str) -> Any:
        return (((report.get("grouped") or {}).get(group) or {}).get("metrics") or {}).get(metric, {}).get("std")

    def _comparison(report: dict[str, Any], key: str) -> dict[str, Any]:
        return ((report.get("generated_comparisons") or {}).get(key) or {})

    def _curve_comparison(report: dict[str, Any], key: str) -> dict[str, Any]:
        return ((report.get("generated_equity_comparisons") or {}).get(key) or {})

    def _humanize_interpretation(code: str | None) -> str:
        mapping = {
            "positive_esg_contribution": "Positive",
            "positive_curve_contribution": "Positive",
            "negative_or_inconclusive_esg_contribution": "Negative or inconclusive",
            "negative_or_inconclusive_curve_contribution": "Negative or inconclusive",
            "inconclusive": "Inconclusive",
            None: "[pending]",
        }
        return mapping.get(code, str(code))

    def _return_diff_cell(curve_payload: dict[str, Any]) -> str:
        if not curve_payload:
            return "[pending]"
        diff = _fmt_percent(curve_payload.get("annualized_return_diff"))
        interval = _fmt_interval(curve_payload.get("annualized_bootstrap_ci95"), percent=True)
        if "[pending]" in {diff, interval}:
            return "[pending]"
        return f"ΔAnn={diff}; CI95 {interval}"

    def _sharpe_diff_cell(pair_payload: dict[str, Any]) -> str:
        if not pair_payload:
            return "[pending]"
        diff = _fmt_decimal(pair_payload.get("mean_diff"))
        interval = _fmt_interval(pair_payload.get("bootstrap_ci95"))
        if "[pending]" in {diff, interval}:
            return "[pending]"
        return f"ΔSharpe={diff}; CI95 {interval}"

    def _primary_outcome_sentence() -> str:
        if paper_run_state != "complete_results":
            return "The final main readout remains gated on local completion of the frozen matrix and verified statistical summaries."
        esg_pair = _comparison(main_report, "b4_vs_b3_sharpe")
        esg_curve = _curve_comparison(main_report, "b4_vs_b3_curve")
        routing_pair = _comparison(main_report, "ours_vs_b4_sharpe")
        routing_curve = _curve_comparison(main_report, "ours_vs_b4_curve")

        esg_positive = esg_pair.get("interpretation") == "positive_esg_contribution" or esg_curve.get("interpretation") == "positive_curve_contribution"
        routing_positive = routing_pair.get("interpretation") == "positive_esg_contribution" or routing_curve.get("interpretation") == "positive_curve_contribution"
        esg_mixed = esg_pair.get("interpretation") == "inconclusive" or esg_curve.get("interpretation") == "inconclusive"
        routing_mixed = routing_pair.get("interpretation") == "inconclusive" or routing_curve.get("interpretation") == "inconclusive"

        if esg_positive and routing_positive:
            return "The completed main matrix indicates that the ESG-enabled SAC baseline improves the primary risk-adjusted comparisons relative to the no-ESG SAC baseline, and OURS_full adds a further routing-oriented gain over B4_sac_esg."
        if esg_positive and (routing_mixed or not routing_positive):
            return "The strongest completed effect comes from ESG signal inclusion relative to the no-ESG SAC baseline, while the incremental routing effect remains mixed or statistically inconclusive."
        if routing_positive and (esg_mixed or not esg_positive):
            return "The completed matrix suggests that routing stabilizes the ESG-enabled strategy more clearly than the underlying annual ESG signal improves the no-ESG baseline."
        return "The completed matrix does not support a broad positive-performance claim and instead points to mixed or statistically inconclusive effects under annual-report timing constraints."

    table_1_rows = [
        [
            group,
            _fmt_percent(_group_metric(main_report, group, "annual_return")),
            _fmt_decimal(_group_metric(main_report, group, "sharpe_ratio")),
            _fmt_decimal(_group_metric(main_report, group, "sortino_ratio")),
            _fmt_decimal(_group_metric(main_report, group, "calmar_ratio")),
            _fmt_percent(_group_metric(main_report, group, "max_drawdown")),
            _fmt_percent(_group_metric(main_report, group, "turnover_rate")),
            _fmt_percent(_group_metric(main_report, group, "win_rate")),
        ]
        for group in ("B1_buyhold", "B2_macd", "B3_sac_noesg", "B4_sac_esg", "OURS_full")
    ]
    table_2_rows = [
        [
            "OURS_full vs B3_sac_noesg",
            _return_diff_cell(_curve_comparison(main_report, "ours_vs_b3_curve")),
            _sharpe_diff_cell(_comparison(main_report, "ours_vs_b3_sharpe")),
            _humanize_interpretation(_comparison(main_report, "ours_vs_b3_sharpe").get("interpretation")),
            "Joint ESG + routing effect under the primary setting.",
        ],
        [
            "OURS_full vs B4_sac_esg",
            _return_diff_cell(_curve_comparison(main_report, "ours_vs_b4_curve")),
            _sharpe_diff_cell(_comparison(main_report, "ours_vs_b4_sharpe")),
            _humanize_interpretation(_comparison(main_report, "ours_vs_b4_sharpe").get("interpretation")),
            "Incremental routing effect beyond ESG-enabled SAC.",
        ],
        [
            "B4_sac_esg vs B3_sac_noesg",
            _return_diff_cell(_curve_comparison(main_report, "b4_vs_b3_curve")),
            _sharpe_diff_cell(_comparison(main_report, "b4_vs_b3_sharpe")),
            _humanize_interpretation(_comparison(main_report, "b4_vs_b3_sharpe").get("interpretation")),
            "Pure ESG signal effect before regime-aware routing.",
        ],
    ]
    table_3_rows = [
        [
            "6a_no_esg_obs",
            "Remove ESG observations",
            _sharpe_diff_cell(_comparison(main_report, "6a_vs_ours_sharpe")),
            _sharpe_diff_cell(_comparison(main_report, "6a_vs_ours_mdd")).replace("ΔSharpe", "ΔMDD"),
            "Observation-channel contribution relative to OURS_full.",
        ],
        [
            "6b_no_esg_reward",
            "Remove ESG reward term",
            _sharpe_diff_cell(_comparison(main_report, "6b_vs_ours_sharpe")),
            _sharpe_diff_cell(_comparison(main_report, "6b_vs_ours_mdd")).replace("ΔSharpe", "ΔMDD"),
            "Reward-shaping contribution relative to OURS_full.",
        ],
        [
            "6c_no_regime",
            "Remove regime-aware routing",
            _sharpe_diff_cell(_comparison(main_report, "6c_vs_ours_sharpe")),
            _sharpe_diff_cell(_comparison(main_report, "6c_vs_ours_mdd")).replace("ΔSharpe", "ΔMDD"),
            "Routing contribution relative to OURS_full.",
        ],
    ]
    table_4_rows = [
        [
            "formula_v2 + sample_full_2022_2025",
            _fmt_decimal(_group_metric(robust_formula_report, "OURS_full", "sharpe_ratio")),
            _fmt_decimal(_group_metric_std(robust_formula_report, "OURS_full", "sharpe_ratio")),
            _humanize_interpretation((robust_formula_report.get("paper_readout") or {}).get("primary_result")),
        ],
        [
            "formula_v2_1 + sample_post_esg_effective",
            _fmt_decimal(_group_metric(robust_post_report, "OURS_full", "sharpe_ratio")),
            _fmt_decimal(_group_metric_std(robust_post_report, "OURS_full", "sharpe_ratio")),
            _humanize_interpretation((robust_post_report.get("paper_readout") or {}).get("primary_result")),
        ],
        [
            "Seed stability summary (main setting)",
            _fmt_decimal(_group_metric(main_report, "OURS_full", "sharpe_ratio")),
            _fmt_decimal(_group_metric_std(main_report, "OURS_full", "sharpe_ratio")),
            "Seed-level robustness anchor for the primary readout.",
        ],
    ]

    if paper_run_state == "no_results":
        paper_05_abstract = (
            "This paper studies whether time-aligned ESG signals improve risk-adjusted performance in a soft actor-critic equity strategy and whether regime-aware routing adds stability beyond an ESG-enabled baseline. The study is framed as a finance-first experiment rather than an AI systems paper: the protocol freezes a 20-stock universe, disclosure-aware feature timing, 2022-2023/2024/2025 splits, and a 96-run matrix that separates no-ESG, ESG-enabled, routing, reward, and observation ablations. The current artifact already fixes the contribution boundary, the main comparison tables, the statistical plan, and the data-lineage contract, so final claims cannot drift after the runs complete. The paper therefore asks not only whether ESG improves Sharpe ratio, drawdown control, and stability, but also whether an annual-report signal remains informative once timing frictions, neutral missing handling, and regime controls are imposed. This framing keeps both positive and negative outcomes publishable under a controlled quantitative-finance protocol and allows final values to be inserted without reopening the study design."
        )
        main_results_paragraphs = [
            "The main results section is still intentionally locked because no local paper-run metrics have been synchronized into the repository result tree. The primary reporting target remains formula_v2_1 + sample_full_2022_2025, and the final narrative will compare OURS_full against B1_buyhold, B2_macd, B3_sac_noesg, and B4_sac_esg only after the expected-run manifest and statistical summaries are complete.",
            "Table 1 and Table 2 therefore remain frozen in submission-ready structure rather than in value. This is deliberate: the paper is already explicit about what will count as a positive ESG effect, what will count as a routing effect, and what will count as a mixed or negative result. Figure 3 placeholder. Equity-curve comparison for B3_sac_noesg, B4_sac_esg, and OURS_full under formula_v2_1 + sample_full_2022_2025.",
        ]
        discussion_paragraphs = [
            "The discussion remains protocol-driven until the local paper-run matrix is populated. If the final results show that ESG-enabled SAC improves the primary risk-adjusted comparisons and that OURS_full adds an additional routing benefit, the paper will interpret that finding as evidence that time-aligned annual ESG information can matter under a tightly controlled trading protocol.",
            "If the final results do not show a statistically meaningful improvement, the same protocol still supports a publishable conclusion: annual-report ESG signals may be too low frequency, too delayed, or too noisy for short-horizon alpha, even when they remain useful for downside control, stability, or long-horizon allocation debates. The value of the paper is that the negative result would rest on a complete ablation matrix and strict timing discipline rather than on anecdotal failure.",
        ]
    elif paper_run_state == "partial_results":
        paper_05_abstract = (
            f"This paper studies whether time-aligned ESG signals improve risk-adjusted performance in a soft actor-critic equity strategy and whether regime-aware routing adds stability beyond an ESG-enabled baseline. The protocol is pre-registered around a fixed 20-stock universe, disclosure-aware feature timing, 2022-2023/2024/2025 splits, and a 96-run matrix that separates no-ESG, ESG-enabled, routing, reward, and observation effects. As of the current local sync, the repository contains {progress_display} worth of paper-run result files, which is sufficient to refresh provisional tables and contribution summaries but not yet sufficient for final inferential claims. The manuscript therefore keeps the contribution boundary and statistical logic fixed while treating the current readout as provisional. This state-aware framing prevents premature conclusions and preserves both the positive-result path and the negative-result interpretation under the same finance-first experimental design."
        )
        main_results_paragraphs = [
            f"The local paper-run namespace is partially synchronized, with {progress_display} worth of completed result files currently visible under {paper_run_context['root']}. Available groups so far include {', '.join(paper_run_context.get('available_groups') or ['none'])}, and the latest synced metric paths indicate {latest_metric_labels}.",
            "This version therefore shows provisional values where the local result tree already supports them, but it does not promote a final headline claim until the expected-run manifest verifies the frozen matrix. Table 1 and Table 2 should be read as progress-aware draft tables rather than as final submission tables. Figure 3 placeholder. Equity-curve comparison for B3_sac_noesg, B4_sac_esg, and OURS_full under formula_v2_1 + sample_full_2022_2025.",
        ]
        discussion_paragraphs = [
            "The partial local sync is already useful for checking whether the tables, comparison logic, and contribution report align with the paper narrative, but it is not yet sufficient for a final publication claim. Interim differences can still change as the missing seeds, groups, or robustness settings are synchronized.",
            "The paper therefore stays conservative at this stage. If the final results show a stable ESG or routing benefit, the discussion will narrow that gain to the completed paired tests and seed-level summaries. If the final results do not show a stable advantage, the negative-result path remains available without rewriting the contribution boundary or the core tables.",
        ]
    else:
        paper_05_abstract = (
            f"This paper studies whether time-aligned ESG signals improve risk-adjusted performance in a soft actor-critic equity strategy and whether regime-aware routing adds stability beyond an ESG-enabled baseline. Under a completed 96-run protocol with a fixed 20-stock universe, disclosure-aware feature timing, and a held-out 2025 test year, the local paper-run results support the following bounded readout: {_primary_outcome_sentence()} The study is framed as a finance-first experiment rather than an AI systems paper, so the contribution rests on paired return tests, bootstrap Sharpe intervals, ablations on ESG observation/reward/routing, and robustness checks across formulas and post-effective samples. The resulting manuscript remains publishable under either a positive or mixed outcome because the final interpretation is tied to timing discipline, seed stability, and downside-aware evidence rather than to a single headline return number."
        )
        main_results_paragraphs = [
            f"The completed primary setting is formula_v2_1 + sample_full_2022_2025, with {progress_display} runs verified against the frozen paper-run matrix. {_primary_outcome_sentence()}",
            "Table 1 reports the main test-set metrics, and Table 2 reports the paired-return and Sharpe-difference evidence used to interpret ESG signal value and routing value separately. Figure 3 placeholder. Equity-curve comparison for B3_sac_noesg, B4_sac_esg, and OURS_full under formula_v2_1 + sample_full_2022_2025.",
        ]
        discussion_paragraphs = [
            f"The completed matrix now supports a result-driven discussion rather than a purely protocol-driven one. {_primary_outcome_sentence()} The interpretation remains intentionally narrow: the evidence applies to a fixed 20-stock universe, annual-report ESG signals, and the disclosed timing rule rather than to ESG trading in general.",
            "The completed ablation and robustness views matter as much as the headline table. They show whether the effect is attributable to ESG observation channels, ESG reward shaping, or regime-aware routing, and they prevent the paper from collapsing into a single aggregated performance claim. Even a mixed outcome remains informative because it clarifies the limits of annual ESG disclosures for daily reinforcement-learning trading.",
        ]

    _replace_section_content(
        paper_05_sections,
        "Abstract",
        paragraphs=[paper_05_abstract],
    )
    _replace_section_content(
        paper_05_sections,
        "1. Introduction",
        paragraphs=[
            "Quantitative trading papers that use deep reinforcement learning often succeed or fail on protocol discipline rather than on algorithm choice alone. Reward design, transaction costs, held-out evaluation, and leakage control typically matter as much as the policy network itself [1]-[6]. ESG-based trading claims face an additional complication: annual-report information is sparse, provider disagreement is substantial, and many papers do not enforce disclosure timing rigorously enough to separate a real signal from hindsight contamination [11]-[15].",
            "This manuscript focuses on that gap. It asks whether an evidence-derived ESG score, converted into daily features under published_date + 1 trading day timing, can improve risk-adjusted outcomes in a soft actor-critic trading system, and whether regime-aware routing adds value beyond an ESG-enabled SAC baseline [2]-[5], [7]-[10]. The study is explicitly finance-first: it does not claim a new retrieval system, a new orchestration framework, or a general ESG platform contribution.",
            "The paper is built around three testable hypotheses. H1 asks whether ESG signal inclusion improves the no-ESG SAC baseline. H2 asks whether regime-aware routing improves the ESG-enabled SAC baseline. H3 asks whether any apparent gain survives observation, reward, and routing ablations together with formula-level and sample-level robustness checks. Those hypotheses are evaluated only under the frozen 2022-2025 paper-run protocol.",
        ],
    )
    _replace_section_content(
        paper_05_sections,
        "2. Related Work",
        paragraphs=[
            "The first relevant literature concerns reinforcement learning for trading. Direct reinforcement and later deep RL studies showed that sequential portfolio control depends heavily on the interaction between reward shaping, frictions, and state construction rather than on policy optimization in isolation [1], [2], [6]. More recent quantitative-finance toolkits and expert-system studies moved these ideas into reproducible SAC-centered workflows, which makes baseline discipline especially important when evaluating incremental signals [3]-[5].",
            "The second literature concerns regime sensitivity and policy specialization. Markov-switching models provide a canonical language for nonstationary market states [7], while mixture-of-experts research motivates routing or specialization when a single policy may not behave uniformly across regimes [8]. In practical trading systems, lightweight regime logic and constrained hyperparameter search are attractive because they keep the protocol interpretable under limited sample sizes [9], [10].",
            "The third literature concerns ESG and financial performance. Meta-analytic evidence suggests that ESG can matter for return, risk, and downside protection, but the sign and size of the effect depend on horizon, sector, portfolio construction, and data provider choice [11]-[13]. For a trading paper, this means that an ESG result must be tied to a precise timing rule and a narrow contribution boundary rather than presented as a universal finance claim.",
            "A final literature addresses ESG signal construction itself. Cross-provider disagreement is large enough that score provenance and timing discipline become part of the experimental contract [14]. In this project, report-level evidence is transformed into a house score through a documented evidence chain, but that upstream machinery is treated only as signal-generation context; the paper is evaluated on trading outcomes, ablations, and robustness under a pre-registered finance protocol [15].",
        ],
    )
    _replace_section_content(
        paper_05_sections,
        "3. Problem Formulation and Contribution Boundary",
        paragraphs=[
            "The target problem is daily allocation over a fixed 20-stock universe under transaction costs and held-out 2025 evaluation. At each trading day, the agent receives market features, macro-regime signals, and ESG-derived variables that become available only after a report-specific effective date. The objective is to improve risk-adjusted portfolio behavior while preserving a strict barrier between disclosed information and future returns.",
            "The study makes three bounded contributions. First, it formalizes a disclosure-aware paper protocol in which ESG-derived features are time aligned, neutral when missing, and separated cleanly from the no-ESG baseline. Second, it evaluates regime-aware policy routing as a constrained control-layer addition on top of an ESG-enabled SAC baseline rather than as an open-ended agentic architecture. Third, it pre-specifies the main, ablation, robustness, and significance tables so that final claims are tied to the completed matrix instead of to post hoc narrative selection.",
            "The core empirical logic is explicitly decomposed. H1 is the ESG signal effect measured by B4_sac_esg versus B3_sac_noesg. H2 is the routing effect measured by OURS_full versus B4_sac_esg. H3 is the decomposition of that gain across ESG observation, ESG reward shaping, and regime routing under 6a, 6b, and 6c. Everything outside this chain—including the broader ESG QA platform, vector stores, orchestration, and deployment plumbing—is out of scope for the present paper.",
        ],
    )
    _replace_section_content(
        paper_05_sections,
        "8. Main Results",
        paragraphs=main_results_paragraphs,
        tables=[
            {
                "caption": "Table 1. Main test-set results under formula_v2_1 and sample_full_2022_2025.",
                "headers": ["Group", "Annual Return", "Sharpe Ratio", "Sortino Ratio", "Calmar Ratio", "Max Drawdown", "Turnover", "Win Rate"],
                "rows": table_1_rows,
            },
            {
                "caption": "Table 2. Statistical significance and interval estimates for the main comparisons.",
                "headers": ["Comparison", "Paired Return Test", "Bootstrap Sharpe Interval", "Effect Direction", "Interpretation"],
                "rows": table_2_rows,
            },
        ],
    )
    _replace_section_content(
        paper_05_sections,
        "9. Ablation and Robustness Analysis",
        paragraphs=[
            "Ablation analysis is designed to explain where any observed gain actually comes from rather than to decorate the main table. Group 6a_no_esg_obs removes ESG observations, 6b_no_esg_reward removes the ESG-linked reward term, and 6c_no_regime removes regime-aware routing while preserving the rest of the environment. These groups must be read directly against OURS_full so that the paper can distinguish signal-channel value from control-layer value.",
            "Robustness analysis then asks whether the main readout survives protocol variation. The first slice swaps formula_v2_1 for formula_v2. The second swaps full_2022_2025 for post_esg_effective. The third inspects seed stability under the primary setting. Together, these checks determine whether the headline outcome is stable or merely an artifact of calibration, timing-window construction, or initialization.",
        ],
        tables=[
            {
                "caption": "Table 3. Ablation analysis for ESG observation, ESG reward, and regime routing.",
                "headers": ["Group", "Key Removal", "Sharpe Change vs OURS_full", "Drawdown Change vs OURS_full", "Interpretation"],
                "rows": table_3_rows,
            },
            {
                "caption": "Table 4. Robustness checks across formulas, samples, and seed stability.",
                "headers": ["Setting", "Primary Metric", "Stability Signal", "Status"],
                "rows": table_4_rows,
            },
        ],
    )
    _replace_section_content(
        paper_05_sections,
        "10. Discussion",
        paragraphs=discussion_paragraphs,
    )
    _replace_section_content(
        paper_05_sections,
        "11. Limitations",
        paragraphs=[
            "The first limitation is data frequency. Annual ESG disclosures move much more slowly than daily market conditions, so even a well-timed ESG feature may act more like a slow-moving preference modifier than like a high-frequency alpha source. This matters especially when the final readout is mixed or weak, because the negative result may reflect signal frequency rather than the irrelevance of ESG altogether.",
            "The second limitation is scope. A fixed 20-stock universe is appropriate for a tightly controlled experiment but not for a universal market claim, and formula_v2_1 plus sample_full_2022_2025 is only one primary readout inside a larger robustness envelope. The paper should therefore present its findings as bounded experimental evidence, with broader universes, higher-frequency ESG proxies, and extended horizon studies left for follow-on work.",
        ],
    )

    _replace_section_content(
        manual_05_sections,
        "2. 现有可直接写入结果",
        bullets=[
            "可直接写入正式协议：20 只股票固定池，Train 2022-2023 / Validation 2024 / Test 2025。",
            "可直接写入时间对齐规则：published_date + 1 trading day，且缺失 ESG 采用 neutral missing handling。",
            f"当前本地 paper-run 状态：{paper_run_state}，进度 {progress_display}。",
            "主表固定使用 formula_v2_1 + sample_full_2022_2025；formula_v2 与 sample_post_esg_effective 只进入稳健性层。",
            f"当前可直接引用的结果侧资产：expected-run manifest、frozen inputs、contribution report、latest metric labels={latest_metric_labels}。",
        ],
    )
    _replace_section_content(
        manual_05_sections,
        "3. 投稿前必须补的实验",
        bullets=[
            "先确保 96-run 主矩阵在本地同步完整，并通过 expected-run manifest 验证。",
            "主表必须基于 formula_v2_1 + sample_full_2022_2025 完整结果，不要用 partial run 提前写最终结论。",
            "必须同时补 paired daily-return test、bootstrap Sharpe interval、seed stability 和 contribution report。",
            "必须保留 negative-result-safe interpretation，不要因为结果 mixed 就临时改研究问题。",
            "若本地仍是 partial_results，只允许写 provisional readout，不允许写 definitive claim。",
        ],
    )
    _replace_section_content(
        manual_05_sections,
        "4. 建议执行命令",
        bullets=[
            "python scripts/quant_rl_expected_run_manifest.py verify --manifest-path storage/quant/rl-experiments/paper-run/protocol/expected_run_manifest.json --report-path storage/quant/rl-experiments/paper-run/summary/expected_run_verification.json",
            "python scripts/quant_rl_esg_contribution_report.py --run-namespace paper-run --sample full_2022_2025 --formula-mode v2_1",
            "python scripts/quant_rl_esg_contribution_report.py --run-namespace paper-run --sample full_2022_2025 --formula-mode v2",
            "python scripts/quant_rl_esg_contribution_report.py --run-namespace paper-run --sample post_esg_effective --formula-mode v2_1",
            "python scripts/generate_scipaper_bundle.py",
        ],
    )
    _replace_section_content(
        manual_05_sections,
        "5. 结果回填清单",
        bullets=[
            "Table 1 只回填 B1/B2/B3/B4/OURS_full 在 formula_v2_1 + sample_full_2022_2025 下的主指标。",
            "Table 2 只回填 OURS vs B3、OURS vs B4、B4 vs B3 的 paired daily-return 与 bootstrap Sharpe。",
            "Table 3 只回填 6a/6b/6c 相对 OURS_full 的 Sharpe 与 drawdown 变化解释。",
            "Table 4 只回填 formula_v2、post_esg_effective、seed stability 三层稳健性摘要。",
            "Discussion 只能在 Table 2 与 Table 4 都可用之后写最终结论；否则只写 provisional / pending 状态。",
        ],
    )

    manual_05_sections.extend(
        [
            {
                "heading": "7. 结果同步后先检查什么",
                "bullets": [
                    f"先看 expected-run manifest 是否仍然是 96 runs，当前本地进度是否为 {progress_display}。",
                    "再看 formula_v2_1/sample_full_2022_2025/results 是否同时有 metrics.json、equity_curve.csv、run_status.json。",
                    "再看 contribution report 是否生成主 JSON、paper_result_tables.md、negative_result note、seed stability、data lineage。",
                    "如果这些资产不齐，正文只能保持 result pending / provisional，不进入最终结论态。",
                ],
            },
            {
                "heading": "8. 哪些值直接进 Table 1-4",
                "bullets": [
                    "Table 1：直接取 grouped metrics 的 mean 值，不手算、不手抄。",
                    "Table 2：直接取 paired equity bootstrap 与 paired Sharpe comparison，不把普通均值差当显著性证据。",
                    "Table 3：直接取 6a/6b/6c 相对 OURS_full 的 paired comparison，不写主观解释替代数值。",
                    "Table 4：直接取 robustness report 的 OURS_full 指标与主设置 seed-level stability 摘要。",
                ],
            },
            {
                "heading": "9. 哪些结论必须等统计检验后才能写",
                "bullets": [
                    "“ESG signal improves risk-adjusted return” 这句必须等 B4 vs B3 的 paired/interval 结果出来后才能写。",
                    "“routing adds incremental value” 这句必须等 OURS vs B4 的 paired/interval 结果出来后才能写。",
                    "“results are robust” 这句必须等 formula_v2、post_esg_effective、seed stability 三层都可用后才能写。",
                    "如果统计仍是 inconclusive，只能写 mixed / inconclusive / downside-control-oriented interpretation。",
                ],
            },
            {
                "heading": "10. 审稿人最可能抓住的 5 个风险点",
                "bullets": [
                    "20-stock universe 太小，是否只是一组示范性结果。",
                    "2025 test 是否真的完全隔离，是否参与过调参或人工筛选。",
                    "ESG / reward shaping / routing 三者是否真的拆分清楚。",
                    "若 Sharpe 升但 raw return 不升，是否仍有金融意义。",
                    "如果结果不显著，这篇是否还能作为 finance-first experimental paper 成立。",
                ],
            },
        ]
    )

    _replace_section_content(
        experiment_05_sections,
        "5. 执行步骤",
        bullets=[
            "先验收 expected-run manifest，确认仍是 96-run matrix，且 sample / formula / group / seed 维度没有漂移。",
            "同步 AutoDL 结果到本地既有 storage/quant/rl-experiments/paper-run/ 路径。",
            "先跑 main setting 的 contribution report，再跑 formula_v2 和 post_esg_effective 的 robustness summaries。",
            "只有在 Table 1-4 所需资产都齐后，才允许把 05 生成到 complete_results 状态。",
            "结果未齐时，生成器只写 provisional readout，禁止手工补写 definitive claim。",
        ],
    )
    _replace_section_content(
        experiment_05_sections,
        "6. 结果回填与交付",
        bullets=[
            "交付顺序固定为：manifest verification -> contribution report -> 重新生成 05 文稿。",
            "正文必须优先回填 Table 1-4，再更新 Abstract 和 Discussion，不要反过来先写结论。",
            "partial run 只允许更新进度、已完成组别和 provisional tables；full run 才允许更新主结论。",
            "negative-result-safe output package 必须和正结果路径一起保留，不因为结果方向改变而删掉。",
            "最终交付应同时包含主图 equity curve 与至少一组 contribution / stability 图。",
        ],
    )
    experiment_05_sections.append(
        {
            "heading": "7. Partial 与 Full Run 的处理规则",
            "bullets": [
                "no_results：只保留 submission-ready skeleton，不写任何结果性句子。",
                "partial_results：允许显示进度、已完成组别和 provisional summary，但不写 final claim。",
                "complete_results：只有在 96-run matrix 验收通过后，才自动回填摘要结果句、Discussion 和 status note。",
            ],
        }
    )

    _append_manual_review_sections(
        manual_01_sections,
        reviewer_focus=[
            "审稿人会首先判断这篇是不是“系统研究问题”，还是把现有组件拼成了一个可用平台。",
            "他们会追问 hybrid retrieval 是否稳定优于 dense-only / BM25-only，而不是只在少数 query 上偶然有效。",
            "他们还会看 contract、runtime、deployment 验证能否真正连接到 answer quality 和 groundedness。",
        ],
        response_points=[
            "先用 benchmark + grounded answer rate 回应“工程说明稿”质疑，再补 service-level profiling 证明系统验证不是空话。",
            "明确主贡献不是新的检索算法，而是可验证的 hybrid retrieval system with service-level validation。",
            "若 hybrid 主要提升 retrieval 而非最终 answer quality，就把结论收紧为 evidence coverage 与 deployment resilience 的系统贡献。",
        ],
        top_experiments=[
            "Hybrid vs dense-only vs BM25-only vs minus-parent-expansion。",
            "Groundedness + human support review。",
            "Latency / throughput / degraded-mode profiling。",
        ],
    )
    _append_manual_review_sections(
        manual_02_sections,
        reviewer_focus=[
            "审稿人最容易质疑的是：只有 ROUGE-L，为什么就敢说领域适配有效。",
            "他们会追问有没有 untuned base model、prompt-only baseline，以及数据构造是否存在泄漏。",
            "他们还会盯住 faithfulness，因为 ESG QA 是高事实性任务，不是普通生成任务。",
        ],
        response_points=[
            "优先把 claim 收紧成 reproducible domain adaptation study，而不是 superiority claim。",
            "补 base model、faithfulness、人审和数据卡，先回答最致命的可比性与真实性问题。",
            "如果 ROUGE 与人审不一致，正文以 groundedness / faithfulness 结论为主，ROUGE 退回支持性指标。",
        ],
        top_experiments=[
            "Untuned base model 同集对照。",
            "Faithfulness / groundedness / 人审。",
            "Reduced-rank LoRA 或 prompt-only baseline。",
        ],
    )
    _append_manual_review_sections(
        manual_03_sections,
        reviewer_focus=[
            "审稿人会先问：为什么一定要多节点编排，单链路或普通 prompt chain 不行吗。",
            "他们会追问 verifier、retry、query rewrite、fallback 四个机制到底各自贡献了什么。",
            "如果没有严格 ablation，这篇很容易被看成设计漂亮但实验不足的 architecture note。",
        ],
        response_points=[
            "把 without verifier / retry / rewrite / single-pass chain 固定成主对照，而不是可选附加实验。",
            "单独报告 recovery rate、fallback success rate、trace completeness 和 latency cost，避免只讲质量不讲代价。",
            "若质量增益有限，就把结论收紧到 controllability、traceability 和 runtime robustness。",
        ],
        top_experiments=[
            "Without verifier / without retry / without rewrite / single-pass chain。",
            "Node-level trace + latency cost。",
            "Fallback degradation 专项实验。",
        ],
    )
    _append_manual_review_sections(
        manual_04_sections,
        reviewer_focus=[
            "审稿人最在意这篇会不会只是一个尚未完成实现的 proposal。",
            "他们会重点看 benchmark 是否公平，尤其是 OCR+text baseline 是否足够强。",
            "他们还会问视觉检索到底在哪些页面 strata 有优势，而不是只给一个总体均值。",
        ],
        response_points=[
            "先把 benchmark、标注协议、公司级 split、统计口径冻结，再去做实现和结果。",
            "明确贡献是 benchmark + method framework + protocol，不假装已经完成所有实验。",
            "即使总体不显著，也要准备好 chart/table/scan/layout 分层结果和 negative-result-safe discussion。",
        ],
        top_experiments=[
            "Benchmark 冻结。",
            "Text-only / OCR+text / visual retrieval 三主对照。",
            "分层页面评测 + reviewed subset groundedness。",
        ],
    )
    _append_manual_review_sections(
        manual_05_sections,
        reviewer_focus=[
            "主投稿最容易被追问的是 ESG signal、reward shaping、regime-aware routing 三者贡献是否真的拆开了。",
            "金融评审会盯住 20-stock universe、2025 test isolation、paired test 与 seed stability 是否足够硬。",
            "如果结果不显著，审稿人会直接问这篇是否还能成立，以及你是否准备好了 negative-result-safe interpretation。",
        ],
        response_points=[
            "把 protocol、ablation、significance、seed stability 放在所有论证之前，不要先讲系统平台背景。",
            "直接用 B3 / B4 / OURS / 6a / 6b / 6c 的拆分回答收益来源，避免泛泛说“ESG 有帮助”。",
            "若 Sharpe 提升而 raw return 不升，按 risk-adjusted improvement 写结论；若差异不显著，就写 annual ESG signal 的边界与下行风险控制价值。",
        ],
        top_experiments=[
            "跑完 96-run 主矩阵。",
            "Paired test + bootstrap Sharpe + seed stability。",
            "ESG / reward / routing 的贡献拆分。",
        ],
    )

    _append_experiment_priority_section(
        experiment_01_sections,
        [
            "这组实验最值当先跑，因为它直接回答“是不是工程拼装”和“hybrid 是否真有效”这两个最核心审稿问题。",
            "Groundedness、人审和 profiling 一起补，能把系统稿从平台说明推到论文级证据。",
        ],
    )
    _append_experiment_priority_section(
        experiment_02_sections,
        [
            "这组实验最值当先跑，因为 untuned base model 和 faithfulness 是当前最致命缺口，补上后整篇稿子的可信度会立刻上一个台阶。",
            "先补可比性和真实性，再谈更复杂的 ablation，能最大幅度降低审稿风险。",
        ],
    )
    _append_experiment_priority_section(
        experiment_03_sections,
        [
            "这组实验最值当先跑，因为没有四组 ablation，就无法证明多节点控制机制真的有用。",
            "Trace 与 latency 一起做，可以同时回答“是否更好”和“是否值得”这两个审稿问题。",
        ],
    )
    _append_experiment_priority_section(
        experiment_04_sections,
        [
            "这组实验最值当先跑，因为在实现之前最该先锁住的是 benchmark 和对照公平性，而不是先追求漂亮结果。",
            "一旦 benchmark、标注和三主对照固定，后续即使结果不显著，这篇论文也仍然有 protocol-level 发表价值。",
        ],
    )
    _append_experiment_priority_section(
        experiment_05_sections,
        [
            "这组实验最值当先跑，因为主矩阵和统计检验决定主投稿论文是否站得住，其它润色都不能替代硬结果。",
            "把贡献拆分和 negative-result-safe package 一起准备，可以避免结果出来后整篇重写。",
        ],
    )

    paper_05_status = f"generated_formal_primary_submission_{paper_run_state}"
    paper_05_priority = "active_primary_submission"
    paper_05_generation_mode = "result_aware_auto_fill"
    paper_05_status_note = "\n".join(
        [
            "# Status Note",
            "",
            "This folder contains the active primary submission package for Paper 05.",
            "",
            f"- Status: {paper_05_status}",
            "- Target venue: Expert Systems with Applications (ESWA)",
            f"- Historical source title: {historical_title}",
            f"- Priority: {paper_05_priority}",
            f"- Generation mode: {paper_05_generation_mode}",
            f"- Local paper-run state: {paper_run_state}",
            f"- Local progress: {progress_display}",
            f"- Expected manifest: {paper_run_context.get('manifest_path') or 'missing'}",
            f"- Root contribution summary: {paper_run_context.get('root_summary_path') or 'missing'}",
            f"- Frozen inputs: latest={paper_run_context.get('frozen_input_paths', {}).get('latest') or 'missing'}, full={paper_run_context.get('frozen_input_paths', {}).get('full') or 'missing'}, post={paper_run_context.get('frozen_input_paths', {}).get('post') or 'missing'}",
            f"- Available groups: {', '.join(paper_run_context.get('available_groups') or ['none'])}",
            f"- Latest synced metrics: {latest_metric_labels}",
            f"- Main readout: {_primary_outcome_sentence()}",
            "- Current package type: formal paper + experiment manual + experiment steps + evidence manifest",
            "- 01-04 are maintained as secondary reference tracks while 05 is the active primary submission.",
            "- Historical files retained in place: 论文1初稿.docx, 实验操作手册.docx",
            "- Legacy blueprint file may remain as archival reference, but it is no longer the primary deliverable.",
        ]
    ) + "\n"
    paper_05_evidence_manifest = {
        "paper_id": "05",
        "status": paper_05_status,
        "priority": paper_05_priority,
        "generation_mode": paper_05_generation_mode,
        "result_state": paper_run_state,
        "target_venue": "Expert Systems with Applications",
        "bundle_focus": "Finance-first ESG-augmented reinforcement learning trading experiment",
        "primary_protocol": {
            "universe_size": 20,
            "train_window": "2022-01-01 to 2023-12-31",
            "validation_window": "2024-01-01 to 2024-12-31",
            "test_window": "2025-01-01 to 2025-12-31",
            "matrix_shape": "2 samples x 2 formulas x 8 groups x 3 seeds = 96 expected runs",
            "main_setting": "formula_v2_1 + sample_full_2022_2025",
            "robustness_settings": ["formula_v2", "sample_post_esg_effective"],
        },
        "paper_run_sync": {
            "root": paper_run_context["root"],
            "manifest_path": paper_run_context.get("manifest_path"),
            "root_summary_path": paper_run_context.get("root_summary_path"),
            "frozen_input_paths": paper_run_context.get("frozen_input_paths"),
            "result_state": paper_run_state,
            "progress": progress_display,
            "metrics_count": paper_run_context.get("metrics_count"),
            "run_status_count": paper_run_context.get("run_status_count"),
            "equity_curve_count": paper_run_context.get("equity_curve_count"),
            "available_groups": paper_run_context.get("available_groups"),
            "available_formulas": paper_run_context.get("available_formulas"),
            "available_samples": paper_run_context.get("available_samples"),
        },
        "artifacts": [
            {"path": "docs/ESG_RL_2022_2025_RUNBOOK.md", "role": "Formal experiment runbook"},
            {"path": "docs/quant_rl/ARCHITECTURE.md", "role": "Quant RL architecture note"},
            {"path": "quant_rl/service/quant_service.py", "role": "Quant RL environment and service logic"},
            {"path": "api/routes_quant_rl.py", "role": "Quant API surface"},
            {"path": "gateway/api/routers/quant_rl.py", "role": "Gateway quant router"},
            {"path": "scripts/run_esg_rl_2022_2025_pipeline.py", "role": "Dataset and paper-run pipeline"},
            {"path": "scripts/autodl_run_paper_experiments.sh", "role": "Formal AutoDL execution driver"},
            {"path": "scripts/quant_rl_paper_preflight.py", "role": "Preflight validation"},
            {"path": "scripts/quant_rl_expected_run_manifest.py", "role": "Expected-run matrix verification"},
            {"path": "scripts/quant_rl_esg_contribution_report.py", "role": "Contribution and appendix report"},
            {"path": "tests/test_quant_rl_paper_hardening.py", "role": "Paper hardening regression suite"},
            {"path": "tests/test_quant_rl_api.py", "role": "Quant API validation"},
            {"path": "tests/test_quant_rl_core.py", "role": "Quant core validation"},
        ],
        "summary_artifacts": [
            _relative_path(Path(item)) if Path(item).is_absolute() else str(item)
            for item in [
                paper_run_context.get("root_summary_path"),
                main_report.get("json_path"),
                main_report.get("paper_tables_markdown_path"),
                main_report.get("negative_result_template_path"),
                ((main_report.get("seed_stability") or {}).get("csv_path")),
                ((main_report.get("data_lineage_appendix") or {}).get("csv_path")),
                (paper_run_context.get("frozen_input_paths") or {}).get("latest"),
                (paper_run_context.get("frozen_input_paths") or {}).get("full"),
                (paper_run_context.get("frozen_input_paths") or {}).get("post"),
            ]
            if item
        ],
        "notes": [
            "The paper is finance-first and does not treat the RAG platform or LangGraph workflow as a primary contribution.",
            "The generator now treats 05 as the active primary submission and reads local paper-run assets before writing manuscript conclusions.",
            f"Primary outcome summary: {_primary_outcome_sentence()}",
        ],
    }

    return [
        {
            "id": "01",
            "folder_name": "01_系统综述_ESG智能问答框架",
            "status": "generated_formal",
            "kind": "formal_paper",
            "priority": "secondary_reference_track",
            "title": "Hybrid Retrieval and Service-Level Validation for an ESG Question Answering System",
            "target_venue": "IEEE Access",
            "paper_filename": "01_Paper_Draft_EN.docx",
            "manual_filename": "01_实验手册_ZH.docx",
            "experiment_filename": "01_实验步骤_ZH.docx",
            "paper_document": _build_document(
                title="Hybrid Retrieval and Service-Level Validation for an ESG Question Answering System",
                subtitle=None,
                sections=paper_01_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "manual_document": _build_document(
                title="论文01 实验手册",
                subtitle="系统综述型：混合检索 + 服务化集成 + 部署验证",
                sections=manual_01_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "experiment_document": _build_document(
                title="论文01 实验步骤",
                subtitle="高规格系统评测：质量、延迟、groundedness、部署韧性",
                sections=experiment_01_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "evidence_manifest": {
                "paper_id": "01",
                "status": "generated_formal",
                "target_venue": "IEEE Access",
                "bundle_focus": "System-level ESG QA architecture",
                "artifacts": [
                    {"path": "gateway/main.py", "role": "FastAPI entrypoint"},
                    {"path": "gateway/api/factory.py", "role": "Application wiring"},
                    {"path": "gateway/rag/retriever.py", "role": "Hybrid retrieval"},
                    {"path": "gateway/rag/indexer.py", "role": "Qdrant persistence"},
                    {"path": "gateway/rag/rag_main.py", "role": "Runtime RAG initialization"},
                    {"path": "railway.toml", "role": "Railway deployment descriptor"},
                    {"path": "tests/test_api_contracts.py", "role": "API contract validation"},
                    {"path": "tests/test_quant_api.py", "role": "Quant interface validation"},
                    {"path": "tests/test_rag_retriever_runtime.py", "role": "Retriever validation"},
                ],
                "notes": [
                    "No fabricated benchmark scores were added.",
                    "This paper remains system-centric and does not treat four-agent orchestration as a primary contribution.",
                ],
            },
        },
        {
            "id": "02",
            "folder_name": "02_LoRA_Qwen25_ESG领域适配",
            "status": "generated_formal",
            "kind": "formal_paper",
            "priority": "secondary_reference_track",
            "title": "LoRA-Based Domain Adaptation of Qwen2.5-7B for ESG Question Answering",
            "target_venue": "Applied Intelligence",
            "paper_filename": "02_Paper_Draft_EN.docx",
            "manual_filename": "02_实验手册_ZH.docx",
            "experiment_filename": "02_实验步骤_ZH.docx",
            "paper_document": _build_document(
                title="LoRA-Based Domain Adaptation of Qwen2.5-7B for ESG Question Answering",
                subtitle=None,
                sections=paper_02_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "manual_document": _build_document(
                title="论文02 实验手册",
                subtitle="LoRA 微调型：Qwen2.5-7B ESG 领域适配",
                sections=manual_02_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "experiment_document": _build_document(
                title="论文02 实验步骤",
                subtitle="高规格领域适配实验：基线、faithfulness、错误分析",
                sections=experiment_02_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "evidence_manifest": {
                "paper_id": "02",
                "status": "generated_formal",
                "target_venue": "Applied Intelligence",
                "bundle_focus": "LoRA fine-tuning and ESG QA evaluation",
                "metrics": {
                    "num_samples": lora_summary["num_samples"],
                    "avg_rougeL": lora_summary["avg_rougeL"],
                },
                "artifacts": [
                    {"path": "training/finetune.py", "role": "Training entrypoint"},
                    {"path": "training/evaluate_model.py", "role": "Evaluation entrypoint"},
                    {"path": "data/rag_training_data/train.jsonl", "role": "Training dataset"},
                    {"path": "data/rag_training_data/val.jsonl", "role": "Validation dataset"},
                    {"path": lora_summary["path"], "role": "ROUGE evaluation summary"},
                    {"path": "model-serving/checkpoint/adapter_config.json", "role": "Saved LoRA configuration"},
                    {"path": "model-serving/checkpoint/README.md", "role": "Checkpoint metadata"},
                    {"path": "tests/test_training_finetune.py", "role": "Training script smoke coverage"},
                    {"path": "tests/test_training_evaluate_model.py", "role": "Evaluation script smoke coverage"},
                ],
                "notes": [
                    "The current repository does not contain a base-model comparison artifact.",
                    "Groundedness, faithfulness, and matched baselines must be added before submission.",
                ],
            },
        },
        {
            "id": "03",
            "folder_name": "03_多智能体_LangGraph金融文档编排",
            "status": "generated_formal",
            "kind": "formal_paper",
            "priority": "secondary_reference_track",
            "title": "Typed-State Multi-Agent Orchestration for ESG Question Answering with Retry and Fallback Control",
            "target_venue": "Electronics",
            "paper_filename": "03_Paper_Draft_EN.docx",
            "manual_filename": "03_实验手册_ZH.docx",
            "experiment_filename": "03_实验步骤_ZH.docx",
            "paper_document": _build_document(
                title="Typed-State Multi-Agent Orchestration for ESG Question Answering with Retry and Fallback Control",
                subtitle=None,
                sections=paper_03_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "manual_document": _build_document(
                title="论文03 实验手册",
                subtitle="多智能体型：状态机编排 + Retry/Fallback",
                sections=manual_03_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "experiment_document": _build_document(
                title="论文03 实验步骤",
                subtitle="高规格编排实验：ablation、trace、recovery、fallback",
                sections=experiment_03_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "evidence_manifest": {
                "paper_id": "03",
                "status": "generated_formal",
                "target_venue": "Electronics",
                "bundle_focus": "Stateful orchestration and robust execution control",
                "artifacts": [
                    {"path": "gateway/agents/graph.py", "role": "Typed state and conditional routing"},
                    {"path": "gateway/agents/router_agent.py", "role": "Task routing"},
                    {"path": "gateway/agents/retriever_agent.py", "role": "Evidence retrieval and caching"},
                    {"path": "gateway/agents/analyst_agent.py", "role": "Analysis stage"},
                    {"path": "gateway/agents/verifier_agent.py", "role": "Verification and retry"},
                    {"path": "tests/test_graph_runtime.py", "role": "Fallback graph validation"},
                    {"path": "tests/test_retriever_agent_runtime.py", "role": "Retriever state validation"},
                ],
                "notes": [
                    "This paper is control-plane centric and treats retrieval and deployment as supporting context only.",
                    "Ablation experiments against simplified chains are still required for submission.",
                ],
            },
        },
        {
            "id": "04",
            "folder_name": "04_视觉检索_ColPali蓝图",
            "status": "generated_formal_experiment_pending",
            "kind": "formal_paper",
            "priority": "secondary_reference_track",
            "title": "OCR-Free Visual Retrieval for ESG Reports with ColPali: A Submission-Ready Experimental Framework",
            "target_venue": "Multimedia Tools and Applications",
            "paper_filename": "04_Paper_Draft_EN.docx",
            "manual_filename": "04_实验手册_ZH.docx",
            "status_filename": "status_note.md",
            "status_note": "\n".join(
                [
                    "# Status Note",
                    "",
                    "This folder now contains a formal experiment-pending submission draft for Paper 04.",
                    "",
                    "- Status: generated_formal_experiment_pending",
                    "- Target venue: Multimedia Tools and Applications",
                    "- Current scope: ColPali-style OCR-free visual retrieval for ESG reports",
                    "- Repository state: benchmark and paper structure are fixed, native visual retriever implementation still pending",
                    "- Legacy blueprint file may remain as archival reference inside this folder.",
                ]
            )
            + "\n",
            "paper_document": _build_document(
                title="OCR-Free Visual Retrieval for ESG Reports with ColPali: A Submission-Ready Experimental Framework",
                subtitle=None,
                sections=paper_04_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "manual_document": _build_document(
                title="论文04 实验手册",
                subtitle="视觉检索论文：投稿边界、缺口与结果回填要求",
                sections=manual_04_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "experiment_filename": "04_实验步骤_ZH.docx",
            "experiment_document": _build_document(
                title="论文04 实验步骤",
                subtitle="高规格视觉检索实验：ColPali、OCR+Text、Text-only 对照",
                sections=experiment_04_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "evidence_manifest": {
                "paper_id": "04",
                "status": "generated_formal_experiment_pending",
                "target_venue": "Multimedia Tools and Applications",
                "bundle_focus": "OCR-free visual retrieval experiment design for ESG reports",
                "repo_scan": {
                    "term": "ColPali",
                    "hit_count": colpali_hit_count,
                    "hits": colpali_hits,
                },
                "artifacts": [
                    {"path": "scripts/esg_corpus_pipeline.py", "role": "Current ESG report corpus preparation baseline"},
                    {"path": "gateway/rag/retriever.py", "role": "Current text-retrieval baseline reference"},
                    {"path": "gateway/rag/indexer.py", "role": "Current index pipeline reference"},
                    {"path": "tests/test_rag_retriever_runtime.py", "role": "Current text retrieval validation surface"},
                ],
                "notes": [
                    "The paper is structurally optimized to near-submission quality, but native ColPali implementation and benchmark runs are still pending.",
                    "No retrieval results are fabricated in this draft.",
                ],
            },
        },
        {
            "id": "05",
            "folder_name": "05_ESG金融量化升级蓝图",
            "status": paper_05_status,
            "kind": "formal_paper",
            "priority": paper_05_priority,
            "generation_mode": paper_05_generation_mode,
            "result_state": paper_run_state,
            "title": "ESG-Augmented Reinforcement Learning with Regime-Aware Policy Routing for Quantitative Equity Trading",
            "target_venue": "Expert Systems with Applications",
            "paper_filename": "05_Paper_Draft_EN.docx",
            "manual_filename": "05_实验手册_ZH.docx",
            "status_filename": "status_note.md",
            "status_note": paper_05_status_note,
            "preserved_files": ["论文1初稿.docx", "实验操作手册.docx"],
            "paper_document": _build_document(
                title="ESG-Augmented Reinforcement Learning with Regime-Aware Policy Routing for Quantitative Equity Trading",
                subtitle=None,
                sections=paper_05_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "manual_document": _build_document(
                title="论文05 实验手册",
                subtitle="主投版：ESG 信号 + RL 交易 + Regime-Aware Policy Routing",
                sections=manual_05_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "experiment_filename": "05_实验步骤_ZH.docx",
            "experiment_document": _build_document(
                title="论文05 实验步骤",
                subtitle="主投实验：96-run 主矩阵、显著性检验、贡献拆分、负结果安全包",
                sections=experiment_05_sections,
                latin_font="Times New Roman",
                east_asia_font="Microsoft YaHei",
            ),
            "evidence_manifest": paper_05_evidence_manifest,
        },
    ]


def _blueprint_bundle_specs(context: dict[str, Any]) -> list[dict[str, Any]]:
    historical_title = context["historical_paper_title"]
    colpali_hits = context["colpali_hits"]
    colpali_hit_count = len(colpali_hits)

    blueprint_04_sections = [
        {
            "heading": "1. 为什么这里只有蓝图",
            "paragraphs": [
                "本文件夹不生成正式英文论文初稿，因为当前仓库扫描结果未发现 ColPali 相关实现证据。严格证据策略下，不能把不存在的模块写成已完成实验。",
                f"本次生成时，对 {', '.join(COLPALI_SCAN_DIRS)} 目录进行关键词扫描，共发现 {colpali_hit_count} 处 ColPali 命中。",
            ],
        },
        {
            "heading": "2. 目标论文定位",
            "paragraphs": [
                "建议题目方向：ColPali 在 ESG 报告视觉理解中的应用：无需 OCR 的 PDF 检索框架。",
                "核心研究问题：视觉 token 检索是否能在图表、版面结构、扫描页和复杂表格场景下优于传统 OCR 文本检索。",
            ],
        },
        {
            "heading": "3. 当前仓库缺口",
            "bullets": [
                "缺视觉 PDF 渲染与 page-image 数据管线。",
                "缺 ColPali 或等价视觉检索模型加载器。",
                "缺视觉 embedding 存储与检索接口。",
                "缺视觉检索基准集、标注集和评估脚本。",
            ],
        },
        {
            "heading": "4. 需要你做的事情",
            "bullets": [
                "新增 PDF 渲染步骤，把每页转换为图像并建立 page-level metadata。",
                "接入视觉检索模型，建立 page embedding 与 chunk/page 对齐逻辑。",
                "设计 3 组 baseline：纯文本 RAG、OCR+文本 RAG、视觉检索 RAG。",
                "构建 ESG 报告视觉问答集，至少覆盖图表、表格、扫描页、复杂布局四类页面。",
            ],
        },
        {
            "heading": "5. 实验矩阵",
            "bullets": [
                "指标：Recall@k、MRR、Exact Match、Grounded Answer Rate、平均延迟、显存占用。",
                "数据切分：公司维度划分，避免同一公司不同年份泄漏。",
                "对照一：文本块检索。",
                "对照二：OCR 后文本检索。",
                "方法组：视觉检索或视觉-文本混合检索。",
            ],
        },
        {
            "heading": "6. 交付清单",
            "bullets": [
                "视觉检索训练或推理脚本。",
                "评估脚本与结果表。",
                "至少 2 张定性案例图。",
                "最终再生成正式英文论文正文。",
            ],
        },
        {
            "heading": "7. 与其他论文的边界",
            "paragraphs": [
                "这篇蓝图完成后，应该与系统综述稿区分为“视觉检索子系统论文”，而不是重复描述整个 ESG 平台。",
            ],
        },
    ]

    blueprint_05_sections = [
        {
            "heading": "1. 为什么这里只有升级蓝图",
            "paragraphs": [
                f"现有历史稿标题为：{historical_title}。该稿已经实质覆盖 ESG + 金融量化交易方向，因此本次不再新写一篇撞题正文。",
                "严格去重策略下，最合理的做法是保留历史稿，并给出升级与分化路线，避免形成内容重合的新稿。",
            ],
        },
        {
            "heading": "2. 当前重合点",
            "bullets": [
                "都以 ESG 信号进入投资或交易决策为核心问题。",
                "都使用 RAG / ESG score 构建作为上游信号来源。",
                "都涉及 regime-aware 或 multi-agent 决策逻辑。",
            ],
        },
        {
            "heading": "3. 推荐升级路线",
            "bullets": [
                "补强 baseline：Buy-and-Hold、技术规则、无 ESG 的 RL、无 regime routing 的 RL。",
                "补强 ablation：去 ESG observation、去 ESG reward、去 regime routing。",
                "补强稳健性：不同股票池、不同时间窗、不同交易成本。",
                "补统计检验：daily return paired t-test、bootstrap Sharpe interval。",
            ],
        },
        {
            "heading": "4. 你需要做的事情",
            "bullets": [
                "继续使用现有论文1，不要新开一个相似题目。",
                "把历史稿中的占位符结果全部补完，并统一引用真实实验表。",
                "把与系统论文、LoRA 论文、多智能体论文重复的系统描述删减到背景级别。",
                "把贡献聚焦在 ESG 信号如何改善 risk-adjusted return。",
            ],
        },
        {
            "heading": "5. 必须重写或删减的部分",
            "bullets": [
                "不要大篇幅展开 Qdrant、BM25、RRF 的系统工程细节。",
                "不要把 LangGraph 四节点流程写成主贡献。",
                "不要与系统综述稿共享整段方法描述或图。",
            ],
        },
        {
            "heading": "6. 投稿注意事项",
            "bullets": [
                "结果表和图必须与新生成的系统稿、LoRA 稿、多智能体稿区分开。",
                "若复用 ESG score 数据，需在写作中声明复用的是数据资产而非同一论文文本。",
                "最安全的做法是把这篇定位为金融实验论文，其余三篇定位为 AI 系统与方法论文。",
            ],
        },
    ]

    return []


def _build_overview_markdown(managed_bundles: list[dict[str, Any]]) -> str:
    primary_bundle = next((bundle for bundle in managed_bundles if bundle.get("priority") == "active_primary_submission"), None)
    lines = [
        "# SCI Paper Bundle Overview",
        "",
        "本目录保留原有历史稿，并新增 5 个按方向拆分的交付文件夹。",
        "",
        "## 历史参考文件",
    ]
    for filename in HISTORICAL_REFERENCE_FILES:
        lines.append(f"- `05_ESG金融量化升级蓝图/{filename}`")
    lines.extend(["", "## 新增五文件夹状态"])
    for bundle in managed_bundles:
        lines.append(f"- `{bundle['id']}` `{bundle['folder_name']}`: `{bundle['status']}`")
    lines.extend(
        [
            "",
            "## 说明",
            "- `01` 到 `05` 均为受生成器管理的论文交付文件夹。",
            "- `01` 到 `05` 均包含正式英文论文初稿 + 中文实验手册 + 证据清单，并统一配套实验步骤文档。",
            "- `05` 是当前 active primary submission；`01` 到 `04` 为 secondary reference tracks / reserve papers。",
            "- `04` 已从蓝图升级为实验待完成的正式视觉检索论文草稿。",
            "- `05` 已升级为主投稿版金融实验论文，目标期刊为 `Expert Systems with Applications`。",
            "- `05` 现已补齐 `05_实验步骤_ZH.docx`，并与 `01` 到 `04` 形成一致的“正文 + 投稿手册 + 实验步骤”结构。",
            f"- `05` 当前结果状态：`{primary_bundle.get('result_state') if primary_bundle else 'unknown'}`；生成模式：`{primary_bundle.get('generation_mode') if primary_bundle else 'static'}`。",
            "- `05` 文件夹中保留历史稿 `论文1初稿.docx` 与 `实验操作手册.docx` 作为来源参考。",
            "- 所有新文件均由 `scripts/generate_scipaper_bundle.py` 管理生成。",
        ]
    )
    return "\n".join(lines) + "\n"


def _bundle_manifest(output_dir: Path, managed_bundles: list[dict[str, Any]]) -> dict[str, Any]:
    payload = {
        "generated_at": _utc_timestamp(),
        "output_dir": _relative_path(output_dir),
        "historical_reference_files": HISTORICAL_REFERENCE_FILES,
        "managed_bundle_order": [bundle["folder_name"] for bundle in managed_bundles],
        "managed_bundles": [],
    }
    for bundle in managed_bundles:
        folder_path = output_dir / bundle["folder_name"]
        if bundle["kind"] == "formal_paper":
            files = [
                _relative_path(folder_path / bundle["paper_filename"]),
                _relative_path(folder_path / bundle["manual_filename"]),
                _relative_path(folder_path / "evidence_manifest.json"),
            ]
            if bundle.get("experiment_filename"):
                files.append(_relative_path(folder_path / bundle["experiment_filename"]))
            if bundle.get("status_filename"):
                files.append(_relative_path(folder_path / bundle["status_filename"]))
        else:
            files = [
                _relative_path(folder_path / bundle["docx_filename"]),
                _relative_path(folder_path / bundle["status_filename"]),
            ]
        payload["managed_bundles"].append(
            {
                "id": bundle["id"],
                "folder_name": bundle["folder_name"],
                "status": bundle["status"],
                "kind": bundle["kind"],
                "priority": bundle.get("priority"),
                "generation_mode": bundle.get("generation_mode"),
                "result_state": bundle.get("result_state"),
                "target_venue": bundle.get("target_venue"),
                "files": files,
                "preserved_files": bundle.get("preserved_files", []),
            }
        )
    return payload


def _document_snapshot(path: Path) -> dict[str, Any]:
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_texts: list[str] = []
    for table in document.tables:
        cells = []
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    cells.append(text)
        table_texts.append(" | ".join(cells))
    return {
        "paragraphs": paragraphs,
        "full_text": "\n".join(paragraphs).lower(),
        "table_count": len(document.tables),
        "table_texts": table_texts,
    }


def _find_heading_index(paragraphs: list[str], heading: str) -> int | None:
    target = heading.strip().lower()
    for index, text in enumerate(paragraphs):
        normalized = text.strip().lower()
        if normalized == target:
            return index
        if normalized.endswith(target) and re.match(r"^\d+\.\s+", normalized):
            return index
    return None


def _is_section_heading(text: str) -> bool:
    normalized = text.strip()
    lower = normalized.lower()
    if lower in {
        "abstract",
        "keywords",
        "highlights",
        "author contributions",
        "data availability statement",
        "conflict of interest",
        "references",
    }:
        return True
    return bool(re.match(r"^\d+\.\s+[A-Za-z].+", normalized))


def _extract_section_paragraphs(paragraphs: list[str], heading: str) -> list[str]:
    start = _find_heading_index(paragraphs, heading)
    if start is None:
        return []
    collected: list[str] = []
    for text in paragraphs[start + 1 :]:
        if _is_section_heading(text):
            break
        collected.append(text)
    return collected


def _extract_reference_entries(paragraphs: list[str]) -> list[str]:
    start = _find_heading_index(paragraphs, "references")
    if start is None:
        return []
    return [text for text in paragraphs[start + 1 :] if re.match(r"^\[\d+\]\s+", text)]


def _extract_citation_numbers(text: str) -> set[int]:
    numbers: set[int] = set()
    for left, right in re.findall(r"\[(\d+)\]\s*-\s*\[(\d+)\]", text):
        numbers.update(range(int(left), int(right) + 1))
    for token in re.findall(r"\[(.*?)\]", text):
        for part in token.split(","):
            piece = part.strip()
            if "-" in piece:
                left, right = [item.strip() for item in piece.split("-", 1)]
                if left.isdigit() and right.isdigit():
                    numbers.update(range(int(left), int(right) + 1))
            elif piece.isdigit():
                numbers.add(int(piece))
    return numbers


def _abstract_word_count(paragraphs: list[str]) -> int:
    return len(re.findall(r"\b[\w.-]+\b", " ".join(paragraphs)))


def _count_citations(text: str) -> int:
    return len(_extract_citation_numbers(text))


def _validate_experiment_steps_document(bundle: dict[str, Any], snapshot: dict[str, Any]) -> None:
    paragraphs = snapshot["paragraphs"]
    joined = "\n".join(paragraphs)
    if len(paragraphs) < 6:
        raise RuntimeError(f"{bundle['id']} experiment steps document appears too short")
    for required in ("实验目标", "执行步骤", "结果回填", "为什么先做"):
        if required not in joined:
            raise RuntimeError(f"{bundle['id']} experiment steps document is missing required section text: {required}")


def _validate_formal_bundle(bundle: dict[str, Any], paper_snapshot: dict[str, Any], manual_snapshot: dict[str, Any]) -> None:
    paragraphs = paper_snapshot["paragraphs"]
    full_text = paper_snapshot["full_text"]
    table_text = "\n".join(paper_snapshot["table_texts"])
    lower_paragraphs = [item.lower() for item in paragraphs]
    manual_paragraphs = "\n".join(manual_snapshot["paragraphs"])

    if not any("related work" in item for item in lower_paragraphs):
        raise RuntimeError(f"{bundle['id']} paper is missing Related Work heading")
    if _find_heading_index(paragraphs, "references") is None:
        raise RuntimeError(f"{bundle['id']} paper is missing References heading")

    ordered_headings = [
        "abstract",
        "keywords",
        "highlights",
        "introduction",
        "conclusion",
        "author contributions",
        "data availability statement",
        "conflict of interest",
        "references",
    ]
    heading_positions: list[int] = []
    for heading in ordered_headings:
        index = _find_heading_index(paragraphs, heading)
        if index is None:
            raise RuntimeError(f"{bundle['id']} paper is missing required heading: {heading}")
        heading_positions.append(index)
    if heading_positions != sorted(heading_positions):
        raise RuntimeError(f"{bundle['id']} paper submission metadata headings are out of order")

    abstract_paragraphs = _extract_section_paragraphs(paragraphs, "abstract")
    if len(abstract_paragraphs) != 1:
        raise RuntimeError(f"{bundle['id']} paper abstract must be a single paragraph")
    abstract_word_count = _abstract_word_count(abstract_paragraphs)
    abstract_ranges = {"01": (130, 160), "02": (150, 190), "03": (130, 160), "04": (150, 190), "05": (150, 190)}
    min_words, max_words = abstract_ranges[bundle["id"]]
    if not (min_words <= abstract_word_count <= max_words):
        raise RuntimeError(
            f"{bundle['id']} paper abstract word count {abstract_word_count} is outside the target range {min_words}-{max_words}"
        )
    abstract_text = " ".join(abstract_paragraphs).lower()
    if re.search(r"\[\d+", abstract_text):
        raise RuntimeError(f"{bundle['id']} paper abstract should not contain citations")
    for phrase in ("this positioning", "we position", "we therefore frame", "the manuscript is framed"):
        if phrase in abstract_text:
            raise RuntimeError(f"{bundle['id']} paper abstract still contains self-positioning phrase: {phrase}")

    reference_entries = _extract_reference_entries(paragraphs)
    if len(reference_entries) < 12:
        raise RuntimeError(f"{bundle['id']} paper has too few numbered reference entries")
    expected_reference_numbers = list(range(1, len(reference_entries) + 1))
    actual_reference_numbers = [int(match.group(1)) for entry in reference_entries if (match := re.match(r"^\[(\d+)\]\s+", entry))]
    if actual_reference_numbers != expected_reference_numbers:
        raise RuntimeError(f"{bundle['id']} paper reference numbering is not sequential")

    if "figure 1 placeholder" not in full_text:
        raise RuntimeError(f"{bundle['id']} paper is missing a figure placeholder")
    if "table 1" not in full_text:
        raise RuntimeError(f"{bundle['id']} paper is missing a table placeholder or caption")
    if "threats to validity" not in full_text and "limitations" not in full_text:
        raise RuntimeError(f"{bundle['id']} paper is missing threats to validity or limitations")

    for forbidden in ("this draft", "repository-backed", "evidence file:", "representative", "should be inserted", "should be added"):
        if forbidden in full_text:
            raise RuntimeError(f"{bundle['id']} paper still contains forbidden phrase: {forbidden}")

    keyword_paragraphs = _extract_section_paragraphs(paragraphs, "keywords")
    if len(keyword_paragraphs) != 1:
        raise RuntimeError(f"{bundle['id']} paper must contain exactly one keywords paragraph")
    keywords = [item.strip() for item in keyword_paragraphs[0].split(",") if item.strip()]
    if not (5 <= len(keywords) <= 7):
        raise RuntimeError(f"{bundle['id']} paper must contain 5 to 7 keywords")
    if len({item.lower() for item in keywords}) != len(keywords):
        raise RuntimeError(f"{bundle['id']} paper keywords contain duplicates")

    highlight_items = _extract_section_paragraphs(paragraphs, "highlights")
    if len(highlight_items) != 3:
        raise RuntimeError(f"{bundle['id']} paper must contain exactly three highlights")
    for item in highlight_items:
        if not item.endswith("."):
            raise RuntimeError(f"{bundle['id']} paper highlight is not written as a full sentence")
        if re.search(r"[\u4e00-\u9fff]", item):
            raise RuntimeError(f"{bundle['id']} paper highlights must be written in English")
        if "todo" in item.lower() or "placeholder" in item.lower():
            raise RuntimeError(f"{bundle['id']} paper highlights still contain placeholder language")

    author_contribution_paragraphs = _extract_section_paragraphs(paragraphs, "author contributions")
    if len(author_contribution_paragraphs) != 1 or "the author contributed to" not in author_contribution_paragraphs[0].lower():
        raise RuntimeError(f"{bundle['id']} paper author contributions are missing the single-author template")
    if any(token in author_contribution_paragraphs[0].lower() for token in ("author a", "author b", "authors contributed", "the authors")):
        raise RuntimeError(f"{bundle['id']} paper author contributions still contain multi-author placeholder language")

    data_availability_paragraphs = _extract_section_paragraphs(paragraphs, "data availability statement")
    if len(data_availability_paragraphs) != 1:
        raise RuntimeError(f"{bundle['id']} paper must contain one data availability statement")
    if "project repository" not in data_availability_paragraphs[0].lower():
        raise RuntimeError(f"{bundle['id']} paper data availability statement must mention the project repository")

    conflict_paragraphs = _extract_section_paragraphs(paragraphs, "conflict of interest")
    if len(conflict_paragraphs) != 1 or "no conflict of interest" not in conflict_paragraphs[0].lower():
        raise RuntimeError(f"{bundle['id']} paper conflict statement is missing or malformed")

    for required_manual_phrase in (
        "现有可直接写入结果",
        "投稿前必须补的实验",
        "审稿人真正关心什么",
        "你应该准备什么回应",
        "最值当先补的实验",
    ):
        if required_manual_phrase not in manual_paragraphs:
            raise RuntimeError(f"{bundle['id']} manual is missing required review-oriented section: {required_manual_phrase}")

    citation_numbers = _extract_citation_numbers("\n".join(paragraphs[: _find_heading_index(paragraphs, "references") or len(paragraphs)]))
    if len(citation_numbers) < 6:
        raise RuntimeError(f"{bundle['id']} paper does not contain enough inline numbered citations")
    if set(expected_reference_numbers) - citation_numbers:
        raise RuntimeError(f"{bundle['id']} paper has uncited reference entries")
    if citation_numbers - set(expected_reference_numbers):
        raise RuntimeError(f"{bundle['id']} paper cites references that are missing from the reference list")

    related_work_paragraphs = _extract_section_paragraphs(paragraphs, "related work")
    related_work_min_paragraphs = {"01": 3, "02": 4, "03": 4, "04": 4, "05": 4}
    if len(related_work_paragraphs) < related_work_min_paragraphs[bundle["id"]]:
        raise RuntimeError(f"{bundle['id']} paper Related Work has too few paragraphs")
    for paragraph in related_work_paragraphs:
        citation_count = _count_citations(paragraph)
        if citation_count == 0:
            raise RuntimeError(f"{bundle['id']} paper has a Related Work paragraph without citations")
        if citation_count > 6:
            raise RuntimeError(f"{bundle['id']} paper has an overly dense Related Work citation block")

    if bundle["id"] == "01":
        for paragraph in paragraphs:
            lower = paragraph.lower()
            if lower.startswith("contribution") and all(term in lower for term in ("router", "retriever", "analyst", "verifier")):
                raise RuntimeError("01 paper still presents router/retriever/analyst/verifier as a main contribution list")
        if not re.match(r"^\[1\]\s+[A-Z]\.\s", reference_entries[0]):
            raise RuntimeError("01 paper references do not look like IEEE initials-first style")
        if '"' not in reference_entries[0]:
            raise RuntimeError("01 paper references are missing IEEE-style quoted titles")
        if not any("doi:" in entry.lower() for entry in reference_entries):
            raise RuntimeError("01 paper references should include at least one DOI in IEEE style")
        if not any("vol." in entry.lower() for entry in reference_entries):
            raise RuntimeError("01 paper references should include at least one IEEE journal-style volume entry")

    if bundle["id"] == "02":
        if paper_snapshot["table_count"] < 1:
            raise RuntimeError("02 paper is missing the required evaluation table")
        if "3723" not in table_text or "0.3596" not in table_text:
            raise RuntimeError("02 paper table does not contain the required LoRA snapshot values")
        if not re.search(r"\(\d{4}\)", reference_entries[0]):
            raise RuntimeError("02 paper references do not match Springer year placement")
        if '"' in " ".join(reference_entries[:3]):
            raise RuntimeError("02 paper references should not use quoted titles")
        if not any("https://doi.org/" in entry.lower() for entry in reference_entries):
            raise RuntimeError("02 paper references should include DOI links where available")

    if bundle["id"] == "03":
        for forbidden in ("bm25", "rrf", "qdrant"):
            if forbidden in full_text:
                raise RuntimeError(f"03 paper still contains retrieval-stack term: {forbidden}")
        if ";" not in reference_entries[0]:
            raise RuntimeError("03 paper references do not look like MDPI ACS-style author lists")
        if not any(re.search(r"arxiv\s+\d{4},\s+arxiv:\d{4}\.\d{4,5}", entry.lower()) for entry in reference_entries):
            raise RuntimeError("03 paper references should include MDPI-style arXiv entries")

    if bundle["id"] == "04":
        for required_phrase in (
            "colpali",
            "ocr-plus-text",
            "company-level",
            "recall@k",
            "grounded answer rate",
            "figure 2 placeholder",
            "table 2",
            "table 3",
        ):
            if required_phrase not in full_text:
                raise RuntimeError(f"04 paper is missing required visual-retrieval phrase: {required_phrase}")
        if paper_snapshot["table_count"] < 3:
            raise RuntimeError("04 paper must contain at least three planned experiment tables")
        if not re.search(r"\(\d{4}\)", reference_entries[0]):
            raise RuntimeError("04 paper references do not match Springer year placement")
        if '"' in " ".join(reference_entries[:3]):
            raise RuntimeError("04 paper references should not use quoted titles")
        if not any("https://doi.org/" in entry.lower() for entry in reference_entries):
            raise RuntimeError("04 paper references should include DOI links where available")

    if bundle["id"] == "05":
        for required_heading in (
            "main results",
            "ablation and robustness analysis",
            "discussion",
            "limitations",
        ):
            if _find_heading_index(paragraphs, required_heading) is None:
                raise RuntimeError(f"05 paper is missing required section: {required_heading}")
        for required_phrase in (
            "house_score_v2_1",
            "published_date + 1 trading day",
            "96 expected runs",
            "full_2022_2025",
            "post_esg_effective",
            "figure 2 placeholder",
            "table 2",
            "table 3",
            "table 4",
        ):
            if required_phrase not in full_text:
                raise RuntimeError(f"05 paper is missing required protocol phrase: {required_phrase}")
        result_state = bundle.get("result_state")
        if result_state in {"no_results", "partial_results"}:
            for required_phrase in ("if the final results show", "if the final results do not show"):
                if required_phrase not in full_text:
                    raise RuntimeError(f"05 paper is missing required dual-path phrase for state {result_state}: {required_phrase}")
        if result_state == "complete_results":
            if "completed matrix" not in full_text and "completed 96-run protocol" not in full_text:
                raise RuntimeError("05 paper complete-results mode must mention the completed matrix or completed protocol")
            if "[pending]" in full_text or "[fill" in full_text:
                raise RuntimeError("05 paper complete-results mode must not retain pending placeholders")
        for forbidden in ("qdrant", "bm25", "rrf", "langgraph"):
            if forbidden in full_text:
                raise RuntimeError(f"05 paper still contains out-of-scope systems term: {forbidden}")
        if paper_snapshot["table_count"] < 4:
            raise RuntimeError("05 paper must contain four submission-oriented result tables")
        if '"' in " ".join(reference_entries[:3]):
            raise RuntimeError("05 paper references should not use quoted titles in Elsevier style")
        if not any("https://doi.org/" in entry.lower() for entry in reference_entries):
            raise RuntimeError("05 paper references should include DOI links where available")


def validate_bundle(output_dir: Path, managed_bundles: list[dict[str, Any]]) -> dict[str, Any]:
    for filename in HISTORICAL_REFERENCE_FILES:
        historical_path = _resolve_historical_reference_path(output_dir, filename)
        if not historical_path.exists():
            raise FileNotFoundError(f"Historical reference file missing: {historical_path}")

    expected_folder_order = [bundle["folder_name"] for bundle in managed_bundles]
    actual_folder_order = sorted([path.name for path in output_dir.iterdir() if path.is_dir()])
    if actual_folder_order != expected_folder_order:
        raise RuntimeError(
            f"Managed folder order mismatch. Expected {expected_folder_order}, got {actual_folder_order}"
        )

    doc_checks: list[dict[str, Any]] = []
    keyword_signatures: dict[str, tuple[str, ...]] = {}
    for bundle in managed_bundles:
        folder_path = output_dir / bundle["folder_name"]
        if not folder_path.exists():
            raise FileNotFoundError(f"Missing bundle folder: {folder_path}")
        if len(list(folder_path.iterdir())) < 2:
            raise RuntimeError(f"Bundle folder has fewer than two files: {folder_path}")

        if bundle["kind"] == "formal_paper":
            required_files = [
                folder_path / bundle["paper_filename"],
                folder_path / bundle["manual_filename"],
                folder_path / "evidence_manifest.json",
            ]
            if bundle.get("experiment_filename"):
                required_files.append(folder_path / bundle["experiment_filename"])
            if bundle.get("status_filename"):
                required_files.append(folder_path / bundle["status_filename"])
            for required_file in required_files:
                if not required_file.exists():
                    raise FileNotFoundError(f"Missing managed file: {required_file}")
            for preserved_file in bundle.get("preserved_files", []):
                preserved_path = folder_path / preserved_file
                if not preserved_path.exists():
                    raise FileNotFoundError(f"Missing preserved historical file: {preserved_path}")

            paper_snapshot = _document_snapshot(folder_path / bundle["paper_filename"])
            manual_snapshot = _document_snapshot(folder_path / bundle["manual_filename"])
            if len(paper_snapshot["paragraphs"]) < 10:
                raise RuntimeError(f"Formal paper appears too empty: {folder_path / bundle['paper_filename']}")
            if len(manual_snapshot["paragraphs"]) < 8:
                raise RuntimeError(f"Manual appears too empty: {folder_path / bundle['manual_filename']}")
            _validate_formal_bundle(bundle, paper_snapshot, manual_snapshot)
            experiment_snapshot = None
            if bundle.get("experiment_filename"):
                experiment_snapshot = _document_snapshot(folder_path / bundle["experiment_filename"])
                _validate_experiment_steps_document(bundle, experiment_snapshot)
            keywords_section = _extract_section_paragraphs(paper_snapshot["paragraphs"], "keywords")
            keyword_signatures[bundle["id"]] = tuple(
                sorted(item.strip().lower() for item in keywords_section[0].split(",") if item.strip())
            )

            doc_checks.append(
                {
                    "path": _relative_path(folder_path / bundle["paper_filename"]),
                    "nonempty_paragraphs": len(paper_snapshot["paragraphs"]),
                    "title": paper_snapshot["paragraphs"][0],
                    "table_count": paper_snapshot["table_count"],
                    "keywords": list(keyword_signatures[bundle["id"]]),
                }
            )
            doc_checks.append(
                {
                    "path": _relative_path(folder_path / bundle["manual_filename"]),
                    "nonempty_paragraphs": len(manual_snapshot["paragraphs"]),
                    "title": manual_snapshot["paragraphs"][0],
                    "table_count": manual_snapshot["table_count"],
                }
            )
            if experiment_snapshot is not None:
                doc_checks.append(
                    {
                        "path": _relative_path(folder_path / bundle["experiment_filename"]),
                        "nonempty_paragraphs": len(experiment_snapshot["paragraphs"]),
                        "title": experiment_snapshot["paragraphs"][0],
                        "table_count": experiment_snapshot["table_count"],
                    }
                )
        else:
            required_files = [
                folder_path / bundle["docx_filename"],
                folder_path / bundle["status_filename"],
            ]
            for required_file in required_files:
                if not required_file.exists():
                    raise FileNotFoundError(f"Missing managed file: {required_file}")

            blueprint_snapshot = _document_snapshot(folder_path / bundle["docx_filename"])
            if len(blueprint_snapshot["paragraphs"]) < 5:
                raise RuntimeError(f"Blueprint document appears too empty: {folder_path / bundle['docx_filename']}")
            doc_checks.append(
                {
                    "path": _relative_path(folder_path / bundle["docx_filename"]),
                    "nonempty_paragraphs": len(blueprint_snapshot["paragraphs"]),
                    "title": blueprint_snapshot["paragraphs"][0],
                    "table_count": blueprint_snapshot["table_count"],
                }
            )

    if len(set(keyword_signatures.values())) != len(keyword_signatures):
        raise RuntimeError("Formal papers must not reuse an identical keyword set")

    manifest_path = output_dir / "paper_bundle_manifest.json"
    overview_path = output_dir / "总览说明.md"
    if not manifest_path.exists():
        raise FileNotFoundError(f"Root manifest missing: {manifest_path}")
    if not overview_path.exists():
        raise FileNotFoundError(f"Overview file missing: {overview_path}")

    manifest = _load_json(manifest_path)
    for bundle in manifest.get("managed_bundles", []):
        for file_path in bundle.get("files", []):
            absolute_path = PROJECT_ROOT / Path(file_path)
            if not absolute_path.exists():
                raise FileNotFoundError(f"Manifest path does not exist: {absolute_path}")

    return {
        "validated_at": _utc_timestamp(),
        "managed_folder_order": expected_folder_order,
        "doc_checks": doc_checks,
    }


def generate_bundle(output_dir: Path) -> dict[str, Any]:
    output_dir.mkdir(parents=True, exist_ok=True)

    historical_paper_title = _load_historical_paper_title(
        _resolve_historical_reference_path(output_dir, "论文1初稿.docx")
    )
    colpali_hits = _scan_term_hits("ColPali", COLPALI_SCAN_DIRS)
    lora_summary = _load_lora_eval_summary()
    paper_run_context = _load_paper_run_context()
    context = {
        "historical_paper_title": historical_paper_title,
        "colpali_hits": colpali_hits,
        "lora_summary": lora_summary,
        "paper_run_context": paper_run_context,
    }

    managed_bundles = sorted(
        _formal_bundle_specs(context) + _blueprint_bundle_specs(context),
        key=lambda bundle: bundle["folder_name"],
    )

    for bundle in managed_bundles:
        folder_path = output_dir / bundle["folder_name"]
        folder_path.mkdir(parents=True, exist_ok=True)

        if bundle["kind"] == "formal_paper":
            _write_docx(folder_path / bundle["paper_filename"], bundle["paper_document"])
            _write_docx(folder_path / bundle["manual_filename"], bundle["manual_document"])
            _write_json(folder_path / "evidence_manifest.json", bundle["evidence_manifest"])
            if bundle.get("experiment_filename"):
                _write_docx(folder_path / bundle["experiment_filename"], bundle["experiment_document"])
            if bundle.get("status_filename"):
                _write_text(folder_path / bundle["status_filename"], bundle["status_note"])
        else:
            _write_docx(folder_path / bundle["docx_filename"], bundle["document"])
            _write_text(folder_path / bundle["status_filename"], bundle["status_note"])

    _write_json(output_dir / "paper_bundle_manifest.json", _bundle_manifest(output_dir, managed_bundles))
    _write_text(output_dir / "总览说明.md", _build_overview_markdown(managed_bundles))

    validation = validate_bundle(output_dir, managed_bundles)
    return {
        "generated_at": _utc_timestamp(),
        "output_dir": _relative_path(output_dir),
        "historical_reference_files": HISTORICAL_REFERENCE_FILES,
        "managed_bundles": [
            {
                "id": bundle["id"],
                "folder_name": bundle["folder_name"],
                "status": bundle["status"],
                "kind": bundle["kind"],
            }
            for bundle in managed_bundles
        ],
        "colpali_hits": colpali_hits,
        "historical_paper_title": historical_paper_title,
        "lora_summary": lora_summary,
        "paper_run_context": {
            "result_state": paper_run_context["result_state"],
            "progress_label": paper_run_context["progress_label"],
            "progress_pct": paper_run_context["progress_pct"],
            "metrics_count": paper_run_context["metrics_count"],
            "completed_run_count": paper_run_context["completed_run_count"],
        },
        "validation": validation,
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate the five-folder SCI paper bundle under scipaper.")
    parser.add_argument(
        "--output-dir",
        default=str(DEFAULT_OUTPUT_DIR),
        help="Target output directory. Defaults to PROJECT_ROOT/scipaper.",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    output_dir = Path(args.output_dir).expanduser()
    if not output_dir.is_absolute():
        output_dir = (PROJECT_ROOT / output_dir).resolve()

    summary = generate_bundle(output_dir)
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
