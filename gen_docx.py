# -*- coding: utf-8 -*-
"""
Generate ESG_RAG_代码精读.docx using python-docx.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"E:\项目夹\ESG Agentic RAG Copilot\ESG_RAG_代码精读.docx"


# ---------------------------------------------------------------------------
# Helper: apply light-gray shading to a paragraph
# ---------------------------------------------------------------------------
def set_paragraph_shading(paragraph, fill="F2F2F2"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


# ---------------------------------------------------------------------------
# Helper: add a code block (monospace, 9pt, gray background)
# ---------------------------------------------------------------------------
def add_code_block(doc, code_text):
    for line in code_text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)
        set_paragraph_shading(p, "F0F0F0")
        run = p.add_run(line if line else 「 」)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    # small gap after code block
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Helper: add explanation bullet lines
# ---------------------------------------------------------------------------
def add_explanations(doc, explanations):
    """explanations: list of (label, text) tuples"""
    for label, text in explanations:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(4)
        # Bold label
        run_label = p.add_run(label + 「  」)
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        # Normal explanation text
        run_text = p.add_run(text)
        run_text.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Data
# ---------------------------------------------------------------------------

CHAPTERS = [
    # -----------------------------------------------------------------------
    # Chapter 1
    # -----------------------------------------------------------------------
    {
        "title": "第一章：chunking.py — 文档切块",
        "intro": (
            "chunking.py 是整个 RAG 管线的第一步，负责把加载进来的原始文档切割成有层级关系的节点树。「
            」核心思路是三层嵌套切块：根节点（2048 token）→ 中间节点（512 token）→ 叶节点（128 token），「
            」同一段文字同时存在于三个粒度，为后续的 AutoMergingRetriever 提供上下文扩展的物质基础。"
        ),
        "code": """\
from llama_index.core.node_parser import HierarchicalNodeParser, get_leaf_nodes


def chunk_documents(documents: list) -> tuple[list, list]:
    \"\"\"
    Split documents into a parent-child node hierarchy.

    Returns:
        all_nodes  — every node (parent + child), needed for the docstore
        leaf_nodes — smallest nodes only, used to build the vector index
    \"\"\"
    parser = HierarchicalNodeParser.from_defaults(
        chunk_sizes=[2048, 512, 128]  # parent → mid → leaf
    )
    all_nodes = parser.get_nodes_from_documents(documents)
    leaf_nodes = get_leaf_nodes(all_nodes)
    return all_nodes, leaf_nodes""",
        "explanations": [
            ("第 1 行", "从 llama_index 核心模块导入两个工具。HierarchicalNodeParser 是父子层级切块解析器，能把一份文档切成有父子关系的节点树；get_leaf_nodes 是工具函数，从所有节点中筛选出最底层的叶节点。"),
            ("第 4 行", "定义函数 chunk_documents，接收文档列表，返回一个元组：(所有节点, 叶节点)。tuple[list, list] 是 Python 3.9+ 的类型注解写法，明确告诉调用方返回两个列表。"),
            ("第 12–14 行", "创建层级解析器。chunk_sizes=[2048, 512, 128] 定义了三个层级：根节点最大 2048 token（约 1500 字），中间节点 512 token（约 380 字），叶节点最小 128 token（约 96 字）。同一段文字会被切成三个粒度，形成父子关系。from_defaults 表示其余参数使用默认值。"),
            ("第 15 行", "调用解析器处理所有文档，返回包含三个层级全部节点的列表。每个节点对象内部记录了：节点文本、节点 ID、父节点 ID、子节点 ID 列表、来源文件元数据。"),
            ("第 16 行", "从全部节点中筛出叶节点（128 token 级别，无子节点的节点）。叶节点是最细粒度的文本单元，后续向量化和存入 Qdrant 的就是这些叶节点。"),
            ("第 17 行", "同时返回两个列表。all_nodes 给 ingestion.py 存入 docstore（AutoMergingRetriever 合并时需要父节点内容）；leaf_nodes 给 ingestion.py 存入 Qdrant 向量库（检索时在叶节点上做相似度查询）。"),
        ],
    },
    # -----------------------------------------------------------------------
    # Chapter 2
    # -----------------------------------------------------------------------
    {
        "title": "第二章：ingestion.py — 向量建库",
        "intro": (
            "ingestion.py 负责把切块后的节点写入持久化存储：叶节点向量化后存入 Qdrant 向量数据库，「
            」全部节点（含父节点）存入 SimpleDocumentStore（内存 docstore）。「
            」同时提供容错逻辑：若本地 Qdrant 不可用，自动降级为内存模式，保证开发环境下程序不崩溃。"
        ),
        "code": """\
import sys
from pathlib import Path
from dotenv import load_dotenv

load_dotenv(Path(__file__).resolve().parents[2] / ".env")
sys.path.insert(0, str(Path(__file__).resolve().parent))

from qdrant_client import QdrantClient
from llama_index.core import StorageContext, VectorStoreIndex
from llama_index.core.storage.docstore import SimpleDocumentStore
from llama_index.vector_stores.qdrant import QdrantVectorStore

QDRANT_URL = "http://localhost:6333"
COLLECTION_NAME = "esg_docs"


def _get_qdrant_client() -> QdrantClient:
    try:
        client = QdrantClient(url=QDRANT_URL, timeout=3)
        client.get_collections()
        print(f"Connected to Qdrant at {QDRANT_URL}")
        return client
    except Exception:
        print("Local Qdrant not reachable — using in-memory store.")
        return QdrantClient(":memory:")


def build_index(all_nodes: list, leaf_nodes: list) -> tuple[VectorStoreIndex, StorageContext]:
    client = _get_qdrant_client()

    vector_store = QdrantVectorStore(
        client=client,
        collection_name=COLLECTION_NAME,
    )
    docstore = SimpleDocumentStore()
    docstore.add_documents(all_nodes)

    storage_context = StorageContext.from_defaults(
        vector_store=vector_store,
        docstore=docstore,
    )

    index = VectorStoreIndex(
        leaf_nodes,
        storage_context=storage_context,
        show_progress=True,
    )
    print(f"Index built — {len(leaf_nodes)} leaf nodes in Qdrant, {len(all_nodes)} total nodes in docstore.")
    return index, storage_context""",
        "explanations": [
            ("第 1–3 行", "导入三个标准库。sys 用于修改 Python 模块搜索路径；Path 是面向对象的路径操作工具，比字符串拼接路径更安全；load_dotenv 从 .env 文件读取环境变量（如 OPENAI_API_KEY）。"),
            ("第 5 行", "定位并加载 .env 文件。__file__ 是当前脚本的绝对路径，.resolve() 转为真实路径，.parents[2] 向上跳两级目录到项目根目录，再拼接 \".env\"。这样无论从哪里运行脚本，都能正确找到 .env 文件。"),
            ("第 6 行", "把当前文件所在目录（gateway/rag/）加入 Python 的模块搜索路径。这样同目录下的 chunking.py、retriever.py 等文件可以直接用 import 导入，而不需要写完整包路径。"),
            ("第 8–11 行", "导入向量数据库相关组件。QdrantClient 是 Qdrant 的 Python 客户端；StorageContext 是 llama_index 的存储上下文容器；VectorStoreIndex 是建立在向量存储上的索引；SimpleDocumentStore 是内存文档存储（存节点对象）；QdrantVectorStore 是 Qdrant 的向量存储适配器。"),
            ("第 13–14 行", "定义两个常量。QDRANT_URL 是本地 Qdrant 服务的地址（Docker 容器默认端口 6333）；COLLECTION_NAME 是向量集合的名称，相当于数据库里的【表名】，所有 ESG 文档的向量都存在这个集合里。"),
            ("第 17–25 行", "定义获取 Qdrant 客户端的私有函数（函数名以 _ 开头表示仅内部使用）。try 块里先创建客户端对象，再立即调用 get_collections() 发送一次真实请求来验证连接是否可用（仅创建对象不代表连接成功）。timeout=3 表示 3 秒内没响应就超时。如果连接失败，except 捕获异常，降级使用内存模式 QdrantClient(\":memory:\")——数据存在内存里，程序结束后丢失，但不会崩溃，适合开发测试。"),
            ("第 28–56 行", "核心建库函数。接收 all_nodes（全部节点）和 leaf_nodes（叶节点），返回 (index, storage_context) 元组。"),
            ("第 30–33 行", "创建 Qdrant 向量存储适配器，指定使用哪个 client 和哪个 collection。这一步只是配置，还没有写入任何数据。"),
            ("第 34–35 行", "创建内存文档存储，把 all_nodes（含父节点、中间节点、叶节点）全部存入。docstore 存的是节点的完整对象：文本内容、层级关系、元数据。AutoMergingRetriever 做父节点合并时，就是从 docstore 里取出父节点的文本。"),
            ("第 37–40 行", "创建存储上下文，把 vector_store 和 docstore 打包在一起。StorageContext 是 llama_index 的统一存储管理器，index 通过它同时访问向量库和文档库。"),
            ("第 42–46 行", "创建向量索引。传入 leaf_nodes（只对叶节点做向量化），llama_index 会自动调用 OpenAI Embedding API 把每个叶节点的文本转成向量，存入 Qdrant。show_progress=True 显示进度条。这一步是最耗时的（30 份报告可能需要数分钟）。"),
        ],
    },
    # -----------------------------------------------------------------------
    # Chapter 3
    # -----------------------------------------------------------------------
    {
        "title": "第三章：indexer.py — 索引持久化与增量管理",
        "intro": (
            "indexer.py 解决向量库的生命周期管理问题：首次建库后如何持久化到磁盘、「
            」服务重启后如何快速恢复（无需重新 Embedding）、如何增量插入新文档、「
            」如何按来源文件删除节点，以及如何查询当前索引的统计信息。"
        ),
        "code": """\
import sys
from pathlib import Path
from dotenv import load_dotenv

load_dotenv(Path(__file__).resolve().parents[2] / ".env")
sys.path.insert(0, str(Path(__file__).resolve().parent))

from qdrant_client import QdrantClient
from llama_index.core import StorageContext, VectorStoreIndex
from llama_index.vector_stores.qdrant import QdrantVectorStore

from ingestion import _get_qdrant_client, COLLECTION_NAME

PERSIST_DIR = str(Path(__file__).resolve().parents[2] / "storage「 / 」docstore")


def collection_exists(client: QdrantClient) -> bool:
    existing = [c.name for c in client.get_collections().collections]
    return COLLECTION_NAME in existing


def persist_storage(storage_context: StorageContext) -> None:
    Path(PERSIST_DIR).mkdir(parents=True, exist_ok=True)
    storage_context.persist(persist_dir=PERSIST_DIR)
    print(f"DocStore persisted to {PERSIST_DIR}")


def load_index() -> tuple[VectorStoreIndex, StorageContext]:
    client = _get_qdrant_client()

    if not collection_exists(client):
        raise RuntimeError(
            f"Collection '{COLLECTION_NAME}' not found in Qdrant. 「
            」Run ingestion.build_index() first."
        )

    if not Path(PERSIST_DIR).exists():
        raise RuntimeError(
            f"DocStore not found at {PERSIST_DIR}. 「
            」Run persist_storage() after building the index."
        )

    vector_store = QdrantVectorStore(client=client, collection_name=COLLECTION_NAME)

    storage_context = StorageContext.from_defaults(
        vector_store=vector_store,
        persist_dir=PERSIST_DIR,
    )

    index = VectorStoreIndex.from_vector_store(
        vector_store,
        storage_context=storage_context,
    )
    print(f"Index loaded — Qdrant + DocStore from {PERSIST_DIR}")
    return index, storage_context


def insert_nodes(index, storage_context, new_all_nodes, new_leaf_nodes):
    storage_context.docstore.add_documents(new_all_nodes)
    index.insert_nodes(new_leaf_nodes)
    persist_storage(storage_context)
    print(f"Inserted {len(new_leaf_nodes)} leaf nodes, {len(new_all_nodes)} total nodes.")


def delete_nodes_by_source(index, storage_context, source_file):
    all_node_ids = list(storage_context.docstore.docs.keys())
    to_delete = [
        nid for nid in all_node_ids
        if storage_context.docstore.get_document(nid).metadata.get("file_name") == source_file
    ]

    if not to_delete:
        print(f"No nodes found for source: {source_file}")
        return

    for nid in to_delete:
        index.delete_nodes([nid])

    persist_storage(storage_context)
    print(f"Deleted {len(to_delete)} nodes from source '{source_file}'.")


def index_stats(client=None):
    client = client or _get_qdrant_client()
    if not collection_exists(client):
        return {"status": "collection not found"}

    info = client.get_collection(COLLECTION_NAME)
    return {
        "collection": COLLECTION_NAME,
        "vector_count": info.vectors_count,
        "status": str(info.status),
        "docstore_path": PERSIST_DIR,
        "docstore_exists": Path(PERSIST_DIR).exists(),
    }""",
        "explanations": [
            ("第 12 行", "从 ingestion.py 导入两个内容：_get_qdrant_client 函数（复用连接逻辑，不重复写）和 COLLECTION_NAME 常量（保持集合名一致）。以 _ 开头的函数虽然是「私有」惯例，但 Python 并不强制，仍可跨文件导入。"),
            ("第 14 行", "定义 docstore 持久化目录路径。parents[2] 跳到项目根目录，再拼接 storage/docstore。首次调用 persist_storage() 时会自动创建这个目录。"),
            ("第 17–19 行", "collection_exists 函数检查 Qdrant 里是否已存在目标集合。get_collections() 返回所有集合的列表，用列表推导式提取名称列表，再检查 COLLECTION_NAME 是否在其中。这是「建库还是加载」决策的依据。"),
            ("第 22–25 行", "persist_storage 函数将 docstore 持久化到磁盘。mkdir(parents=True, exist_ok=True) 递归创建目录，exist_ok=True 表示目录已存在时不报错。storage_context.persist() 是 llama_index 的内置方法，把 docstore 序列化为 JSON 文件存到 persist_dir。"),
            ("第 28–54 行", "load_index 是服务重启后的恢复函数，不重新 embed。先做两个前置检查：Qdrant collection 必须存在（否则说明从未建过库）；本地 docstore 目录必须存在（否则父节点结构丢失，AutoMerging 无法工作）。两个检查都通过后，用 StorageContext.from_defaults(persist_dir=...) 从磁盘 JSON 恢复 docstore，再用 VectorStoreIndex.from_vector_store() 从已有 Qdrant 数据恢复 index，整个过程不调用任何 Embedding API，速度极快。"),
            ("第 57–61 行", "insert_nodes 增量插入新文档。先把新文档的所有节点加入 docstore（保持父子结构完整），再把叶节点插入 Qdrant 向量库（触发 Embedding），最后调用 persist_storage 同步更新磁盘上的 docstore JSON。"),
            ("第 64–78 行", "delete_nodes_by_source 按来源文件删除节点。先从 docstore 里找出所有来自该文件的节点 ID（通过 metadata[\"file_name\"] 匹配），再逐一从 index 删除（同时会从 Qdrant 删除对应向量），最后持久化更新 docstore。"),
            ("第 81–91 行", "index_stats 返回当前向量库的统计信息，包括向量数量、状态、docstore 路径和是否存在。用于监控和调试，可在 main.py 的 /health 接口里调用。"),
        ],
    },
    # -----------------------------------------------------------------------
    # Chapter 4
    # -----------------------------------------------------------------------
    {
        "title": "第四章：retriever.py — 五级检索链",
        "intro": (
            "retriever.py 是整个 RAG 管线中技术密度最高的模块，构建了一条五级检索链：\n「
            」密集检索（向量语义）→ 稀疏检索（BM25 关键词）→ RRF 融合 → AutoMerging 父节点扩展 → BGE 精排。\n「
            」每一级都有明确的分工，最终把最相关的 5 个父节点上下文传给 LLM 生成答案。"
        ),
        "code": """\
from llama_index.core import StorageContext, VectorStoreIndex
from llama_index.core.retrievers import AutoMergingRetriever, QueryFusionRetriever
from llama_index.core.retrievers.fusion_retriever import FUSION_MODES
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.retrievers.bm25 import BM25Retriever
from llama_index.postprocessor.flag_embedding_reranker import FlagEmbeddingReranker


def build_query_engine(
    index: VectorStoreIndex,
    storage_context: StorageContext,
    leaf_nodes: list,
    similarity_top_k: int = 12,
    rerank_top_n: int = 5,
) -> RetrieverQueryEngine:

    # 1. Dense retriever
    vector_retriever = index.as_retriever(similarity_top_k=similarity_top_k)

    # 2. BM25 retriever
    bm25_retriever = BM25Retriever.from_defaults(
        nodes=leaf_nodes,
        similarity_top_k=similarity_top_k,
    )

    # 3. Reciprocal Rank Fusion
    fusion_retriever = QueryFusionRetriever(
        retrievers=[vector_retriever, bm25_retriever],
        similarity_top_k=similarity_top_k,
        num_queries=1,
        mode=FUSION_MODES.RECIPROCAL_RANK,
        use_async=True,
        verbose=True,
    )

    # 4. Parent node expansion
    auto_merging_retriever = AutoMergingRetriever(
        base_retriever=fusion_retriever,
        storage_context=storage_context,
        simple_ratio_thresh=0.4,
        verbose=True,
    )

    # 5. Rerank
    reranker = FlagEmbeddingReranker(
        model="BAAI/bge-reranker-base",
        top_n=rerank_top_n,
    )

    # 6. Query engine
    query_engine = RetrieverQueryEngine(
        retriever=auto_merging_retriever,
        node_postprocessors=[reranker],
    )

    return query_engine""",
        "explanations": [
            ("第 1–6 行", "导入检索链所需的所有组件。AutoMergingRetriever 负责父节点合并；QueryFusionRetriever 负责多路检索融合；FUSION_MODES 是枚举常量，定义融合算法类型；RetrieverQueryEngine 是最终的查询引擎壳；BM25Retriever 是关键词稀疏检索器；FlagEmbeddingReranker 是基于 BGE 模型的精排器。"),
            ("第 9–15 行", "函数签名定义了完整的检索链构建接口。index 和 storage_context 来自 ingestion.py；leaf_nodes 来自 chunking.py（BM25 需要在原始文本上建倒排索引）；similarity_top_k=12 表示每路检索各取 12 个候选节点；rerank_top_n=5 表示精排后最终保留 5 个节点传给 LLM。"),
            ("第 18 行", "创建密集检索器（向量语义检索）。index.as_retriever() 把向量索引包装成检索器接口，查询时把用户问题转成向量，在 Qdrant 里做余弦相似度搜索，返回最相近的 12 个叶节点。"),
            ("第 21–24 行", "创建 BM25 稀疏检索器（关键词检索）。BM25 是经典的 TF-IDF 改进算法，不需要向量化，直接在叶节点文本上建倒排索引，对精确关键词（如公司名、指标名）匹配效果好，弥补语义检索的不足。"),
            ("第 27–34 行", "创建 RRF 融合检索器，把向量检索和 BM25 检索的两路结果合并。num_queries=1 表示不额外生成查询变体，只做结果融合。FUSION_MODES.RECIPROCAL_RANK 指定使用 RRF 算法：对每个文档计算 Σ 1/(k+rank)，在两路结果中排名越靠前的文档得分越高，最终合并成一个统一排名的列表。use_async=True 让两路检索并行执行，减少等待时间。"),
            ("第 37–42 行", "创建父节点扩展检索器，包裹在 fusion_retriever 外层。工作流程：先让 fusion_retriever 找出 top-12 叶节点（128 token），然后检查每个父节点（512 token）下，被命中的叶节点占该父所有叶节点的比例是否超过 40%（simple_ratio_thresh=0.4）。超过则用整个父节点替换这些叶节点，返回更完整的上下文。"),
            ("第 45–48 行", "创建精排器。BAAI/bge-reranker-base 是北京智源研究院开源的 cross-encoder 模型，首次使用时会自动下载（约 1GB）。cross-encoder 会把"问题+候选文档"拼在一起输入模型，比向量相似度更准确地评估相关性，但速度慢，所以只对父节点扩展后的少量候选（约 5–15 个）做精排，最终保留 top_n=5 个。"),
            ("第 51–55 行", "创建最终的查询引擎。RetrieverQueryEngine 把检索链（retriever）和后处理器（node_postprocessors）打包，内置一个 response_synthesizer，负责把 top-5 父节点的文本拼成 prompt 发给 LLM，返回最终回答。调用方只需 query_engine.query(\"问题\") 一行代码即可完成全流程。"),
        ],
    },
    # -----------------------------------------------------------------------
    # Chapter 5
    # -----------------------------------------------------------------------
    {
        "title": "第五章：rag_main.py — 管线总调度",
        "intro": (
            "rag_main.py 是 RAG 子系统的统一入口，扮演「总导演」角色。「
            」它不实现任何具体的切块、建库或检索逻辑，而是按正确的顺序调用其余四个模块，「
            」并通过单例模式保证整个 FastAPI 进程生命周期内 RAG 引擎只初始化一次。"
        ),
        "code": """\
import sys
from pathlib import Path
from dotenv import load_dotenv

load_dotenv(Path(__file__).resolve().parents[2] / ".env")
sys.path.insert(0, str(Path(__file__).resolve().parent))

from llama_index.core import SimpleDirectoryReader
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.core.node_parser import get_leaf_nodes

from chunking  import chunk_documents
from ingestion import build_index
from indexer   import collection_exists, load_index, persist_storage, _get_qdrant_client
from retriever import build_query_engine

DATA_DIR = str(Path(__file__).resolve().parents[2] / "data「 / 」raw")

_query_engine: RetrieverQueryEngine | None = None


def get_query_engine(force_rebuild: bool = False) -> RetrieverQueryEngine:
    global _query_engine
    if _query_engine is not None and not force_rebuild:
        return _query_engine

    client = _get_qdrant_client()

    if not force_rebuild and collection_exists(client):
        print("[RAG] Existing index found — loading from Qdrant + docstore...")
        index, storage_context = load_index()
        leaf_nodes = _get_leaf_nodes_from_docstore(storage_context)
    else:
        print(f"[RAG] Building index from documents in {DATA_DIR} ...")
        documents = _load_documents()
        all_nodes, leaf_nodes = chunk_documents(documents)
        index, storage_context = build_index(all_nodes, leaf_nodes)
        persist_storage(storage_context)

    _query_engine = build_query_engine(index, storage_context, leaf_nodes)
    print("[RAG] Query engine ready.")
    return _query_engine


def _load_documents():
    if not Path(DATA_DIR).exists():
        raise FileNotFoundError(
            f"Data directory not found: {DATA_DIR}\\n「
            」Please put ESG reports (PDF/docx/txt) into data/raw/ first."
        )
    docs = SimpleDirectoryReader(
        DATA_DIR,
        required_exts=[".pdf", ".docx", ".txt", ".md"],
        recursive=True,
    ).load_data()

    if not docs:
        raise ValueError(f"No documents found in {DATA_DIR}")

    print(f"[RAG] Loaded {len(docs)} document(s) from {DATA_DIR}")
    return docs


def _get_leaf_nodes_from_docstore(storage_context) -> list:
    all_nodes = list(storage_context.docstore.docs.values())
    leaf_nodes = get_leaf_nodes(all_nodes)
    print(f"[RAG] Restored {len(leaf_nodes)} leaf nodes from docstore.")
    return leaf_nodes""",
        "explanations": [
            ("第 12–15 行", "导入其余四个模块的核心函数。这四行清晰地展示了 rag_main.py 的职责：它不实现任何具体逻辑，只是把四个模块的功能按正确顺序串联起来，是整个 RAG 流程的【总导演】。"),
            ("第 17 行", "用 Path 动态计算 data/raw 的绝对路径。这样无论项目放在哪个目录，路径都能正确解析，避免硬编码路径。"),
            ("第 20 行", "声明模块级全局变量 _query_engine，初始为 None。类型注解 RetrieverQueryEngine | None 是 Python 3.10+ 的写法，表示可能是引擎对象也可能是 None。这实现了单例模式：整个进程生命周期内，RAG 引擎只初始化一次。"),
            ("第 23–24 行", "单例检查。如果 _query_engine 已经初始化且不需要强制重建，直接返回缓存的引擎，避免重复加载向量库。global 关键字声明要修改模块级变量。"),
            ("第 28–31 行", "快速恢复路径。collection_exists() 检测到 Qdrant 已有数据，说明不是首次运行，直接调用 load_index() 从磁盘+Qdrant 恢复 index 和 storage_context，然后从 docstore 重建 leaf_nodes（BM25 需要）。整个过程不调用 Embedding API，通常在几秒内完成。"),
            ("第 32–37 行", "首次建库路径。加载 data/raw 里的所有文档，依次调用切块→建库→持久化，完成后 Qdrant 和磁盘都有了数据，下次重启走快速恢复路径。"),
            ("第 44–57 行", "_load_documents 先检查目录是否存在，再用 SimpleDirectoryReader 递归读取所有指定格式的文件。required_exts 过滤只读取 PDF/docx/txt/md，避免误读系统文件。recursive=True 允许 data/raw 下有子文件夹。"),
            ("第 60–63 行", "_get_leaf_nodes_from_docstore 从已持久化的 docstore 恢复叶节点列表。取出所有节点后调用 get_leaf_nodes()，它通过检查 child_nodes 是否为空来识别叶节点（无子节点的就是叶节点，即 128 token 级别的节点）。"),
        ],
    },
    # -----------------------------------------------------------------------
    # Chapter 6
    # -----------------------------------------------------------------------
    {
        "title": "第六章：evaluator.py — RAG 效果评估",
        "intro": (
            "evaluator.py 实现了对整条 RAG 管线的自动化评估框架，使用 GPT-4o-mini 作为评判模型，「
            」从三个维度对每个问题的回答打分：\n「
            」  • Faithfulness（忠实度）：答案是否完全来自检索内容，防止模型幻觉；\n「
            」  • Relevancy（检索相关性）：检索到的节点是否与问题相关；\n「
            」  • Answer Relevancy（回答相关性）：答案是否直接回应了问题。"
        ),
        "code": """\
import sys
import json
from pathlib import Path
from dotenv import load_dotenv

load_dotenv(Path(__file__).resolve().parents[2] / ".env")
sys.path.insert(0, str(Path(__file__).resolve().parent))

from llama_index.core.evaluation import (
    FaithfulnessEvaluator,
    RelevancyEvaluator,
    AnswerRelevancyEvaluator,
    BatchEvalRunner,
)
from llama_index.llms.openai import OpenAI
from llama_index.core.query_engine import RetrieverQueryEngine

_JUDGE_LLM = OpenAI(model="gpt-4o-mini", temperature=0)

DEFAULT_QUESTIONS = [
    "该公司的环境评分（Environmental Score）是多少？",
    "该公司在碳排放管理上采取了哪些措施？",
    "该公司的社会责任（Social）评分如何？",
    "董事会中女性成员的比例是多少？",
    "该公司是否有明确的净零排放承诺，目标年份是哪年？",
    "该公司的 ESG 综合评级相比同行处于什么水平？",
    "供应链中的劳工标准合规情况如何？",
    "该公司在水资源管理方面的具体指标是什么？",
]


def evaluate(query_engine, questions=None):
    questions = questions or DEFAULT_QUESTIONS

    faithfulness_eval     = FaithfulnessEvaluator(llm=_JUDGE_LLM)
    relevancy_eval        = RelevancyEvaluator(llm=_JUDGE_LLM)
    answer_relevancy_eval = AnswerRelevancyEvaluator(llm=_JUDGE_LLM)

    runner = BatchEvalRunner(
        evaluators={
            "faithfulness":     faithfulness_eval,
            "relevancy":        relevancy_eval,
            "answer_relevancy": answer_relevancy_eval,
        },
        workers=4,
        show_progress=True,
    )

    eval_results = runner.evaluate_queries(
        query_engine=query_engine,
        queries=questions,
    )

    results = []
    for i, q in enumerate(questions):
        faithfulness     = eval_results["faithfulness"][i]
        relevancy        = eval_results["relevancy"][i]
        answer_relevancy = eval_results["answer_relevancy"][i]

        results.append({
            "question":         q,
            "answer":           faithfulness.response or "",
            "faithfulness":     "YES" if faithfulness.passing     else "NO",
            "relevancy":        "YES" if relevancy.passing        else "NO",
            "answer_relevancy": "YES" if answer_relevancy.passing else "NO",
        })

    return results


def report(results):
    total = len(results)
    faith_pass  = sum(1 for r in results if r["faithfulness"]     == "YES")
    relev_pass  = sum(1 for r in results if r["relevancy"]        == "YES")
    answer_pass = sum(1 for r in results if r["answer_relevancy"] == "YES")

    summary = {
        "total_questions":       total,
        "faithfulness_pass":     f"{faith_pass}/{total}  ({faith_pass/total:.0%})",
        "relevancy_pass":        f"{relev_pass}/{total}  ({relev_pass/total:.0%})",
        "answer_relevancy_pass": f"{answer_pass}/{total}  ({answer_pass/total:.0%})",
        "details":               results,
    }
    return summary


def save_report(summary, output_path=None):
    path = Path(output_path) if output_path else (
        Path(__file__).resolve().parents[2] / "evaluation「 / 」rag_eval_report.json"
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)
    print(f"Report saved → {path}")
    return str(path)""",
        "explanations": [
            ("第 18 行", "创建评判 LLM 的单例。使用 GPT-4o-mini 而非 GPT-4o 的原因：评估任务是结构化判断（YES/NO），不需要顶级模型，4o-mini 成本约为 GPT-4o 的 1/15。temperature=0 关闭随机性，确保对同一问题每次给出相同判断，结果可重复。"),
            ("第 20–30 行", "定义 8 道 ESG 场景测试题，覆盖环境（E）、社会（S）、治理（G）三个维度。这些问题不需要标准答案，评估器通过 LLM 自动判断答案质量。"),
            ("第 35–37 行", "创建三个评估器，全部使用同一个 _JUDGE_LLM。FaithfulnessEvaluator 检查答案是否完全来自检索到的文本（防幻觉）；RelevancyEvaluator 检查检索到的节点是否与问题相关；AnswerRelevancyEvaluator 检查答案内容是否直接回应了问题（防答非所问）。"),
            ("第 39–46 行", "BatchEvalRunner 批量并行运行评估。workers=4 表示同时跑 4 个评估任务，加速评估过程。它内部会自动调用 query_engine 跑所有问题，并把每个问题的检索结果和答案传给对应评估器打分。"),
            ("第 48–51 行", "调用 evaluate_queries 启动批量评估。返回值是字典：{\"faithfulness\": [EvaluationResult, ...], \"relevancy\": [...], \"answer_relevancy\": [...]}，每个列表长度等于问题数量。"),
            ("第 53–65 行", "遍历每个问题，从 eval_results 中取出三个维度的评估结果。faithfulness.response 是 BatchEvalRunner 内部查询引擎时存下的回答文本，直接复用，不重复查询（避免额外 API 调用）。passing 是布尔值，True 表示通过，转成 \"YES\"/\"NO\" 方便阅读。"),
            ("第 68–79 行", "report 函数统计三个维度的通过率。用生成器表达式计算通过数量，格式化为 \"7/8 (88%)\" 的可读字符串。"),
            ("第 82–89 行", "save_report 将评估报告序列化为 JSON 文件。ensure_ascii=False 保留中文字符不转义；indent=2 美化缩进便于阅读。默认保存到 evaluation/ 目录，也可传自定义路径。"),
        ],
    },
]


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------
def build_document():
    doc = Document()

    # ---- Page margins ----
    section = doc.sections[0]
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ---- Main Title ----
    title_para = doc.add_heading("ESG Agentic RAG Copilot — RAG 模块代码精读", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # ---- Subtitle / intro ----
    intro_para = doc.add_paragraph(
        "本文档对 ESG Agentic RAG Copilot 项目中 RAG 子系统的六个核心 Python 文件进行逐行精读。「
        」整条管线的数据流如下：\n\n「
        」  ① chunking.py   — 三层层级切块，生成 all_nodes 与 leaf_nodes\n「
        」  ② ingestion.py  — 叶节点向量化写入 Qdrant，全节点存入 docstore\n「
        」  ③ indexer.py    — docstore 持久化、服务重启恢复、增量插入/删除\n「
        」  ④ retriever.py  — 向量+BM25 融合 → AutoMerging → BGE 精排，构建查询引擎\n「
        」  ⑤ rag_main.py   — 总调度入口，单例模式管理引擎生命周期\n「
        」  ⑥ evaluator.py  — 三维度自动化评估（忠实度、检索相关性、回答相关性）"
    )
    intro_para.paragraph_format.space_after = Pt(12)
    for run in intro_para.runs:
        run.font.size = Pt(10.5)

    doc.add_paragraph()  # spacer

    # ---- Chapters ----
    for ch in CHAPTERS:
        # Heading 1
        h1 = doc.add_heading(ch["title"], level=1)
        for run in h1.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            run.font.size = Pt(14)

        # Chapter intro
        intro = doc.add_paragraph(ch["intro"])
        intro.paragraph_format.space_after = Pt(8)
        for run in intro.runs:
            run.font.size = Pt(10.5)

        # Code block heading
        code_heading = doc.add_paragraph()
        run_ch = code_heading.add_run("▌ 完整源码")
        run_ch.bold = True
        run_ch.font.size = Pt(11)
        run_ch.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        code_heading.paragraph_format.space_after = Pt(2)

        # Code block
        add_code_block(doc, ch["code"])

        # Explanations heading
        h2 = doc.add_heading(「逐句解析」, level=2)
        for run in h2.runs:
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            run.font.size = Pt(12)

        # Explanations
        add_explanations(doc, ch["explanations"])

        # Chapter separator
        doc.add_paragraph()

    # ---- Save ----
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
